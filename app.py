'''
	******************************************************************************************
	    Assembly:                Bro
	    Filename:                app.py
	    Author:                  Terry D. Eppler
	    Created:                 05-31-2024
	
	    Last Modified By:        Terry D. Eppler
	    Last Modified On:        05-01-2025
	******************************************************************************************
	<copyright file="app.py" company="Terry D. Eppler">
	
	           Bro is a data analysis tool integrating various Generative GPT, Text-Processing, and
	           Machine-Learning algorithms for federal analysts.
	           Copyright ©  2023 Terry Eppler
	
	   Permission is hereby granted, free of charge, to any person obtaining a copy
	   of this software and associated documentation files (the “Software”),
	   to deal in the Software without restriction,
	   including without limitation the rights to use,
	   copy, modify, merge, publish, distribute, sublicense,
	   and/or sell copies of the Software,
	   and to permit persons to whom the Software is furnished to do so,
	   subject to the following conditions:
	
	   The above copyright notice and this permission notice shall be included in all
	   copies or substantial portions of the Software.
	
	   THE SOFTWARE IS PROVIDED “AS IS”, WITHOUT WARRANTY OF ANY KIND, EXPRESS OR IMPLIED,
	   INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
	   FITNESS FOR A PARTICULAR PURPOSE AND NON-INFRINGEMENT.
	   IN NO EVENT SHALL THE AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM,
	   DAMAGES OR OTHER LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE,
	   ARISING FROM, OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER
	   DEALINGS IN THE SOFTWARE.
	
	   You can contact me at:  terryeppler@gmail.com or eppler.terry@epa.gov
	
	</copyright>
	<summary>
	  app.py
	</summary>
	******************************************************************************************
'''
from __future__ import annotations

import base64
import hashlib
from io import BytesIO
import os
import re
import sqlite3
import time
from pathlib import Path
from typing import Any, Dict, List, Tuple
import numpy as np
import pandas as pd
import plotly.express as px
import streamlit as st
from llama_cpp import Llama
import config as cfg
from boogr import Error, Logger

try:
	import fitz
except ImportError:
	fitz = None

# ==============================================================================
# Model Path Resolution
# ==============================================================================
MODEL_PATH_OBJ = Path( cfg.MODEL_PATH )

def resolve_mmproj_path( ) -> Path | None:
	"""Resolves the explicitly configured Gemma 3 multimodal projector path.

	Purpose:
		Resolves only an explicitly configured multimodal projector so Bro cannot silently pair the
		active Gemma model with an unrelated mmproj GGUF found in the same directory.

	Returns:
		Path | None: Existing configured multimodal projector path when available.
	"""
	configured = str(
		getattr( cfg, 'MMPROJ_PATH', '' ) or
		getattr( cfg, 'MM_PROJ_PATH', '' ) or
		os.environ.get( 'BRO_MMPROJ_PATH', '' ) or
		os.environ.get( 'GEMMA_MMPROJ_PATH', '' )
	).strip( )
	if not configured:
		return None
	configured_path = Path( configured )
	return configured_path if configured_path.exists( ) else None


MMPROJ_PATH_OBJ = resolve_mmproj_path( )

def local_model_available( ) -> bool:
	"""Determines whether the configured local GGUF model path resolves to an existing model file before runtime loading is attempted.

	Purpose:
		Determines whether the configured local GGUF model path resolves to an existing
		model file before runtime loading is attempted.

	Returns:
		bool: Boolean result produced by the operation.
	"""
	try:
		return MODEL_PATH_OBJ.exists( )
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'local_model_available'
		exception.method = 'local_model_available(  ) -> bool'
		Logger( ).write( exception )
		return False

# ==============================================================================
# SESSION STATE INITIALIZATION
# ==============================================================================
if 'mode' not in st.session_state:
	st.session_state[ 'mode' ] = ''

if 'messages' not in st.session_state:
	st.session_state[ 'messages' ] = [ ]

if 'system_instructions' not in st.session_state:
	st.session_state[ 'system_instructions' ] = ''

if 'context_window' not in st.session_state:
	st.session_state[ 'context_window' ] = 0

if 'cpu_threads' not in st.session_state:
	st.session_state[ 'cpu_threads' ] = int( cfg.CORES )

if 'max_tokens' not in st.session_state:
	st.session_state[ 'max_tokens' ] = 0

if 'temperature' not in st.session_state:
	st.session_state[ 'temperature' ] = 0.0

if 'top_percent' not in st.session_state:
	st.session_state[ 'top_percent' ] = 0.0

if 'top_k' not in st.session_state:
	st.session_state[ 'top_k' ] = 0

if 'frequency_penalty' not in st.session_state:
	st.session_state[ 'frequency_penalty' ] = 0.0

if 'presense_penalty' not in st.session_state:
	st.session_state[ 'presense_penalty' ] = 0.0

if 'repeat_penalty' not in st.session_state:
	st.session_state[ 'repeat_penalty' ] = 0.0

if 'repeat_window' not in st.session_state:
	st.session_state[ 'repeat_window' ] = 0

if 'random_seed' not in st.session_state:
	st.session_state[ 'random_seed' ] = 0

if 'basic_docs' not in st.session_state:
	st.session_state[ 'basic_docs' ] = [ ]

if 'use_semantic' not in st.session_state:
	st.session_state[ 'use_semantic' ] = False

if 'is_grounded' not in st.session_state:
	st.session_state[ 'is_grounded' ] = False

if 'selected_prompt_id' not in st.session_state:
	st.session_state[ 'selected_prompt_id' ] = ''

if 'pending_system_prompt_name' not in st.session_state:
	st.session_state[ 'pending_system_prompt_name' ] = ''

# -------- TEXT GENERATION  ---------------------

if 'task_preset' not in st.session_state:
	st.session_state[ 'task_preset' ] = 'Chat'

if 'response_format' not in st.session_state:
	st.session_state[ 'response_format' ] = 'Markdown'

if 'use_chat_history' not in st.session_state:
	st.session_state[ 'use_chat_history' ] = True

if 'use_document_context' not in st.session_state:
	st.session_state[ 'use_document_context' ] = False

if 'reasoning_depth' not in st.session_state:
	st.session_state[ 'reasoning_depth' ] = 'Medium'

if 'answer_only' not in st.session_state:
	st.session_state[ 'answer_only' ] = False

if 'use_self_check' not in st.session_state:
	st.session_state[ 'use_self_check' ] = False

if 'deterministic_reasoning' not in st.session_state:
	st.session_state[ 'deterministic_reasoning' ] = False

if 'coding_language' not in st.session_state:
	st.session_state[ 'coding_language' ] = 'Python'

if 'coding_task' not in st.session_state:
	st.session_state[ 'coding_task' ] = 'Generate'

if 'coding_include_comments' not in st.session_state:
	st.session_state[ 'coding_include_comments' ] = True

if 'coding_editor_format' not in st.session_state:
	st.session_state[ 'coding_editor_format' ] = True

if 'coding_fenced_output' not in st.session_state:
	st.session_state[ 'coding_fenced_output' ] = True

if 'response_language' not in st.session_state:
	st.session_state[ 'response_language' ] = 'English'

if 'translation_source_language' not in st.session_state:
	st.session_state[ 'translation_source_language' ] = 'Auto Detect'

if 'translation_target_language' not in st.session_state:
	st.session_state[ 'translation_target_language' ] = 'English'

if 'writing_task' not in st.session_state:
	st.session_state[ 'writing_task' ] = 'Draft'

if 'writing_tone' not in st.session_state:
	st.session_state[ 'writing_tone' ] = 'Professional'

if 'writing_audience' not in st.session_state:
	st.session_state[ 'writing_audience' ] = 'General'

if 'classification_type' not in st.session_state:
	st.session_state[ 'classification_type' ] = 'Multi-Class'

if 'classification_confidence' not in st.session_state:
	st.session_state[ 'classification_confidence' ] = False

if 'classification_allow_unknown' not in st.session_state:
	st.session_state[ 'classification_allow_unknown' ] = True

if 'task_detail' not in st.session_state:
	st.session_state[ 'task_detail' ] = 'Standard'

if 'task_focus' not in st.session_state:
	st.session_state[ 'task_focus' ] = 'Balanced'

if 'translation_mode' not in st.session_state:
	st.session_state[ 'translation_mode' ] = 'Natural'

if 'translation_preserve_formatting' not in st.session_state:
	st.session_state[ 'translation_preserve_formatting' ] = True

if 'writing_length' not in st.session_state:
	st.session_state[ 'writing_length' ] = 'Standard'

if 'classification_explain' not in st.session_state:
	st.session_state[ 'classification_explain' ] = False

if 'response_length' not in st.session_state:
	st.session_state[ 'response_length' ] = 'Standard'

if 'response_include_headings' not in st.session_state:
	st.session_state[ 'response_include_headings' ] = True

if 'batch_size' not in st.session_state:
	st.session_state[ 'batch_size' ] = 512

if 'micro_batch_size' not in st.session_state:
	st.session_state[ 'micro_batch_size' ] = 128

# -------- IMAGE TO TEXT ---------------------

if 'vision_task' not in st.session_state:
	st.session_state[ 'vision_task' ] = 'Extract Visible Text'

if 'vision_detail' not in st.session_state:
	st.session_state[ 'vision_detail' ] = 'Detailed'

if 'vision_response_format' not in st.session_state:
	st.session_state[ 'vision_response_format' ] = 'Plain Text'

if 'vision_response_language' not in st.session_state:
	st.session_state[ 'vision_response_language' ] = 'English'

if 'vision_preserve_layout' not in st.session_state:
	st.session_state[ 'vision_preserve_layout' ] = True

if 'vision_include_visible_text' not in st.session_state:
	st.session_state[ 'vision_include_visible_text' ] = True

if 'vision_last_response' not in st.session_state:
	st.session_state[ 'vision_last_response' ] = ''

if 'active_prompt_caption' not in st.session_state:
	st.session_state[ 'active_prompt_caption' ] = ''

if 'preview_effective_prompt' not in st.session_state:
	st.session_state[ 'preview_effective_prompt' ] = False

if 'last_preview_input' not in st.session_state:
	st.session_state[ 'last_preview_input' ] = ''

# -------- DOCQNA ---------------------

if 'uploaded' not in st.session_state:
	st.session_state[ 'uploaded' ] = [ ]

if 'active_docs' not in st.session_state:
	st.session_state[ 'active_docs' ] = [ ]

if 'doc_bytes' not in st.session_state:
	st.session_state[ 'doc_bytes' ] = { }

if 'doc_source' not in st.session_state:
	st.session_state[ 'doc_source' ] = 'uploadlocal'

if 'docqna_vec_ready' not in st.session_state:
	st.session_state[ 'docqna_vec_ready' ] = False

if 'docqna_fingerprint' not in st.session_state:
	st.session_state[ 'docqna_fingerprint' ] = ''

if 'docqna_chunk_count' not in st.session_state:
	st.session_state[ 'docqna_chunk_count' ] = 0

if 'docqna_fallback_rows' not in st.session_state:
	st.session_state[ 'docqna_fallback_rows' ] = [ ]

# -------- DOCUMENT Q&A EXTENSIONS ---------------------

if 'retrieval_k' not in st.session_state:
	st.session_state[ 'retrieval_k' ] = 6

if 'retrieval_chunk_size' not in st.session_state:
	st.session_state[ 'retrieval_chunk_size' ] = 1200

if 'retrieval_chunk_overlap' not in st.session_state:
	st.session_state[ 'retrieval_chunk_overlap' ] = 200

if 'show_retrieved_chunks' not in st.session_state:
	st.session_state[ 'show_retrieved_chunks' ] = True

if 'require_grounding' not in st.session_state:
	st.session_state[ 'require_grounding' ] = True

if 'answer_from_excerpts_only' not in st.session_state:
	st.session_state[ 'answer_from_excerpts_only' ] = True

if 'prefer_sqlite_vec' not in st.session_state:
	st.session_state[ 'prefer_sqlite_vec' ] = True

if 'allow_similarity_fallback' not in st.session_state:
	st.session_state[ 'allow_similarity_fallback' ] = True

if 'docqna_action' not in st.session_state:
	st.session_state[ 'docqna_action' ] = 'Answer Question'

if 'ocr_enabled' not in st.session_state:
	st.session_state[ 'ocr_enabled' ] = False

if 'prefer_native_pdf_text' not in st.session_state:
	st.session_state[ 'prefer_native_pdf_text' ] = True

if 'include_page_markers' not in st.session_state:
	st.session_state[ 'include_page_markers' ] = False

if 'show_docqna_diagnostics' not in st.session_state:
	st.session_state[ 'show_docqna_diagnostics' ] = False

if 'show_ocr_status' not in st.session_state:
	st.session_state[ 'show_ocr_status' ] = True

if 'show_runtime_metadata' not in st.session_state:
	st.session_state[ 'show_runtime_metadata' ] = False

if 'docqna_last_retrieval' not in st.session_state:
	st.session_state[ 'docqna_last_retrieval' ] = [ ]

if 'docqna_inventory_rows' not in st.session_state:
	st.session_state[ 'docqna_inventory_rows' ] = [ ]

if 'grounding_failure_behavior' not in st.session_state:
	st.session_state[ 'grounding_failure_behavior' ] = 'State Insufficient Information'

if 'retrieval_backend' not in st.session_state:
	st.session_state[ 'retrieval_backend' ] = 'Automatic'

if 'docqna_rebuild_each_query' not in st.session_state:
	st.session_state[ 'docqna_rebuild_each_query' ] = False

if 'docqna_action_detail' not in st.session_state:
	st.session_state[ 'docqna_action_detail' ] = 'Standard'

if 'ocr_page_limit' not in st.session_state:
	st.session_state[ 'ocr_page_limit' ] = '5 Pages'

if 'docqna_include_semantic_context' not in st.session_state:
	st.session_state[ 'docqna_include_semantic_context' ] = True

if 'docqna_context_order' not in st.session_state:
	st.session_state[ 'docqna_context_order' ] = 'Retrieved First'

if 'docqna_ocr_cache' not in st.session_state:
	st.session_state[ 'docqna_ocr_cache' ] = { }

# -------- SEMANTIC SEARCH ---------------------

if 'semantic_context_buffer' not in st.session_state:
	st.session_state[ 'semantic_context_buffer' ] = [ ]

if 'semantic_chunk_size' not in st.session_state:
	st.session_state[ 'semantic_chunk_size' ] = 1200

if 'semantic_chunk_overlap' not in st.session_state:
	st.session_state[ 'semantic_chunk_overlap' ] = 200

if 'semantic_top_k' not in st.session_state:
	st.session_state[ 'semantic_top_k' ] = 8

if 'semantic_min_similarity' not in st.session_state:
	st.session_state[ 'semantic_min_similarity' ] = 0.0

if 'semantic_query_show_diagnostics' not in st.session_state:
	st.session_state[ 'semantic_query_show_diagnostics' ] = False

if 'semantic_clear_existing' not in st.session_state:
	st.session_state[ 'semantic_clear_existing' ] = True

if 'semantic_append_existing' not in st.session_state:
	st.session_state[ 'semantic_append_existing' ] = False

if 'semantic_show_diagnostics' not in st.session_state:
	st.session_state[ 'semantic_show_diagnostics' ] = True

if 'semantic_uploaded_names' not in st.session_state:
	st.session_state[ 'semantic_uploaded_names' ] = [ ]

if 'semantic_result_rows' not in st.session_state:
	st.session_state[ 'semantic_result_rows' ] = [ ]

if 'semantic_selected_rows' not in st.session_state:
	st.session_state[ 'semantic_selected_rows' ] = [ ]

if 'semantic_index_chunk_count' not in st.session_state:
	st.session_state[ 'semantic_index_chunk_count' ] = 0

if 'semantic_index_dim' not in st.session_state:
	st.session_state[ 'semantic_index_dim' ] = 0

if 'semantic_index_doc_count' not in st.session_state:
	st.session_state[ 'semantic_index_doc_count' ] = 0

if 'semantic_last_query' not in st.session_state:
	st.session_state[ 'semantic_last_query' ] = ''

# -------- PROMPT ENGINEERING EXTENSIONS ---------------------

if 'prompt_category' not in st.session_state:
	st.session_state[ 'prompt_category' ] = ''

if 'prompt_task' not in st.session_state:
	st.session_state[ 'prompt_task' ] = 'Chat'

if 'prompt_response_format' not in st.session_state:
	st.session_state[ 'prompt_response_format' ] = 'Markdown'

if 'pe_language' not in st.session_state:
	st.session_state[ 'pe_language' ] = 'English'

if 'pe_generator_goal' not in st.session_state:
	st.session_state[ 'pe_generator_goal' ] = ''

if 'pe_generator_constraints' not in st.session_state:
	st.session_state[ 'pe_generator_constraints' ] = ''

if 'pe_generator_style' not in st.session_state:
	st.session_state[ 'pe_generator_style' ] = 'Practical'

if 'pe_generated_template' not in st.session_state:
	st.session_state[ 'pe_generated_template' ] = ''

# -------- DATABASE  ---------------------

if 'dm_asset_sync_status' not in st.session_state:
	st.session_state[ 'dm_asset_sync_status' ] = ''

if 'dm_asset_counts' not in st.session_state:
	st.session_state[ 'dm_asset_counts' ] = { }

if 'dm_selected_asset_table' not in st.session_state:
	st.session_state[ 'dm_selected_asset_table' ] = 'documents'

if 'dm_register_uploaded_docs' not in st.session_state:
	st.session_state[ 'dm_register_uploaded_docs' ] = False

if 'dm_register_uploaded_images' not in st.session_state:
	st.session_state[ 'dm_register_uploaded_images' ] = False

# ==============================================================================
# UTILITIES
# ==============================================================================

def image_to_base64( path: str ) -> str:
	"""Reads an image file from disk and returns a Base64-encoded text representation for Streamlit or Markdown rendering workflows.

	Purpose:
		Reads an image file from disk and returns a Base64-encoded text representation for
		Streamlit or Markdown rendering workflows.

	Args:
		path: Filesystem path to read.

	Returns:
		str: Text produced by the operation.
	"""
	with open( path, "rb" ) as f:
		return base64.b64encode( f.read( ) ).decode( )

def cosine_similarity( a: np.ndarray, b: np.ndarray ) -> float:
	"""Computes cosine similarity between two numeric embedding vectors used by semantic retrieval and fallback vector search workflows.

	Purpose:
		Computes cosine similarity between two numeric embedding vectors used by semantic
		retrieval and fallback vector search workflows.

	Args:
		a: First numeric vector.
		b: Second numeric vector.

	Returns:
		float: Floating-point result produced by the operation.
	"""
	denom = np.linalg.norm( a ) * np.linalg.norm( b )
	return float( np.dot( a, b ) / denom ) if denom else 0.0

# -------- CHAT/TEXT UTILITIES --------------------

def normalize_text( text: str ) -> str:
	"""Normalizes user-provided text for prompt preparation, semantic comparison, and lightweight search workflows.

	Purpose:
		Normalizes user-provided text for prompt preparation, semantic comparison, and
		lightweight search workflows.

	Args:
		text: Text to process.

	Returns:
		str: Text produced by the operation.
	"""
	if not text:
		return ""
	
	# Lowercase
	text = text.lower( )
	
	# Remove punctuation except . ! ?
	text = re.sub( r"[^\w\s\.\!\?]", "", text )
	
	# Ensure single space after sentence delimiters
	text = re.sub( r"([.!?])\s*", r"\1 ", text )
	
	# Normalize whitespace
	text = re.sub( r"\s+", " ", text ).strip( )
	
	return text

def chunk_text( text: str, size: int = None, overlap: int = None ) -> List[ str ]:
	"""Splits text into overlapping chunks using explicit arguments or retrieval defaults stored in Streamlit session state.

	Purpose:
		Splits text into overlapping chunks using explicit arguments or retrieval defaults
		stored in Streamlit session state.

	Args:
		text: Text to process.
		size: size value used by this workflow.
		overlap: overlap value used by this workflow.

	Returns:
		List[str]: Result produced by the operation.
	"""
	if not text:
		return [ ]
	
	chunk_size = int(
		size if size is not None else st.session_state.get( 'retrieval_chunk_size', 1200 ) )
	chunk_overlap = int(
		overlap if overlap is not None else st.session_state.get( 'retrieval_chunk_overlap', 200 ) )
	
	if chunk_size <= 0:
		chunk_size = 1200
	
	if chunk_overlap < 0:
		chunk_overlap = 0
	
	if chunk_overlap >= chunk_size:
		chunk_overlap = max( 0, chunk_size // 4 )
	
	chunks: List[ str ] = [ ]
	i = 0
	step = max( 1, chunk_size - chunk_overlap )
	
	while i < len( text ):
		chunk = text[ i:i + chunk_size ]
		if chunk and chunk.strip( ):
			chunks.append( chunk )
		i += step
	
	return chunks

def convert_xml( text: str ) -> str:
	"""Converts XML-like prompt sections into Markdown headings for prompt editing and preview workflows.

	Purpose:
		Converts XML-like prompt sections into Markdown headings for prompt editing and
		preview workflows.

	Args:
		text: Text to process.

	Returns:
		str: Text produced by the operation.
	"""
	markdown_blocks: List[ str ] = [ ]
	for match in cfg.XML_BLOCK_PATTERN.finditer( text ):
		raw_tag: str = match.group( "tag" )
		body: str = match.group( "body" ).strip( )
		
		# Humanize tag name for Markdown heading
		heading: str = raw_tag.replace( "_", " " ).replace( "-", " " ).title( )
		markdown_blocks.append( f"## {heading}" )
		if body:
			markdown_blocks.append( body )
	return "\n\n".join( markdown_blocks )

def convert_markdown( text: Any ) -> str:
	"""Converts Markdown headings to XML-like tags, or XML-like heading tags back to Markdown, for prompt-template interoperability.

	Purpose:
		Converts Markdown headings to XML-like tags, or XML-like heading tags back to
		Markdown, for prompt-template interoperability.

	Args:
		text: Text to process.

	Returns:
		str: Text produced by the operation.
	"""
	if not isinstance( text, str ) or not text.strip( ):
		return ""
	
	# Normalize newlines
	src = text.replace( "\r\n", "\n" ).replace( "\r", "\n" )
	
	htag_pattern = re.compile( r"<h([1-6])>(.*?)</h\1>", flags=re.IGNORECASE | re.DOTALL )
	md_heading_pattern = re.compile( r"^(#{1,6})[ \t]+(.+?)[ \t]*$", flags=re.MULTILINE )
	
	# ------------------------------------------------------------------
	# Direction detection
	# ------------------------------------------------------------------
	contains_htags = bool( htag_pattern.search( src ) )
	
	# ------------------------------------------------------------------
	# XML-like heading tags -> Markdown headings
	# ------------------------------------------------------------------
	if contains_htags:
		def _htag_to_md( match: re.Match ) -> str:
			level = int( match.group( 1 ) )
			content = match.group( 2 ).strip( )
			
			# Preserve inner newlines safely by collapsing interior whitespace
			# while keeping content readable.
			content = re.sub( r"[ \t]+\n", "\n", content )
			content = re.sub( r"\n[ \t]+", "\n", content )
			
			return f"{'#' * level} {content}"
		
		out = htag_pattern.sub( _htag_to_md, src )
		return out.strip( )
	
	# ------------------------------------------------------------------
	# Markdown headings -> XML-like heading tags
	# ------------------------------------------------------------------
	def _md_to_htag( match: re.Match ) -> str:
		hashes = match.group( 1 )
		content = match.group( 2 ).strip( )
		level = len( hashes )
		return f"<h{level}>{content}</h{level}>"
	
	out = md_heading_pattern.sub( _md_to_htag, src )
	return out.strip( )

def inject_response_css( ) -> None:
	"""Injects response-specific CSS into the Streamlit page to style chat text, headings, and links.

	Purpose:
		Injects response-specific CSS into the Streamlit page to style chat text, headings,
		and links.
	"""
	st.markdown(
		"""
		<style>
		/* Chat message text */
		.stChatMessage p {
			color: rgb(220, 220, 220);
			font-size: 1rem;
			line-height: 1.6;
		}

		/* Headings inside chat responses */
		.stChatMessage h1 {
			color: rgb(0, 120, 252); /* DoD Blue */
			font-size: 1.6rem;
		}

		.stChatMessage h2 {
			color: rgb(0, 120, 252);
			font-size: 1.35rem;
		}

		.stChatMessage h3 {
			color: rgb(0, 120, 252);
			font-size: 1.15rem;
		}
		
		.stChatMessage a {
			color: rgb(0, 120, 252); /* DoD Blue */
			text-decoration: underline;
		}
		
		.stChatMessage a:hover {
			color: rgb(80, 160, 255);
		}

		</style>
		""", unsafe_allow_html=True )

def style_subheaders( ) -> None:
	"""Injects CSS that standardizes subheader colors across the main Streamlit interface and chat output areas.

	Purpose:
		Injects CSS that standardizes subheader colors across the main Streamlit interface
		and chat output areas.
	"""
	st.markdown(
		"""
		<style>
		div[data-testid="stMarkdownContainer"] h2,
		div[data-testid="stMarkdownContainer"] h3,
		div[data-testid="stChatMessage"] div[data-testid="stMarkdownContainer"] h2,
		div[data-testid="stChatMessage"] div[data-testid="stMarkdownContainer"] h3 {
			color: rgb(0, 120, 252) !important;
		}
		</style>
		""",
		unsafe_allow_html=True, )

def save_message( role: str, content: str ) -> None:
	"""Persists a chat message role and content pair into the local SQLite chat history table.

	Purpose:
		Persists a chat message role and content pair into the local SQLite chat history
		table.

	Args:
		role: Chat role to persist.
		content: Chat message content to persist.
	"""
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute( 'INSERT INTO chat_history (role, content) VALUES (?, ?)', (role, content) )

def load_history( ) -> List[ Tuple[ str, str ] ]:
	"""Loads persisted chat history from SQLite in insertion order for Text Generation startup state.

	Purpose:
		Loads persisted chat history from SQLite in insertion order for Text Generation
		startup state.

	Returns:
		List[Tuple[str, str]]: Result produced by the operation.
	"""
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		return conn.execute( 'SELECT role, content FROM chat_history ORDER BY id' ).fetchall( )

def clear_history( ) -> None:
	"""Deletes persisted chat history from the local SQLite database without altering prompt, model, or document state.

	Purpose:
		Deletes persisted chat history from the local SQLite database without altering
		prompt, model, or document state.
	"""
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute( "DELETE FROM chat_history" )

# -------- PROMPT ENGINEERING UTILITIES ----------------

PROMPT_CATEGORY_OPTIONS: List[ str ] = [
	'Business / Finance / Marketing',
	'Compliance / Legal / Budget',
	'Data Analytics & Governance',
	'Instruction/ Training / Planning',
	'Prompt Engineering',
	'Research / Academic',
	'Software Engineering',
	'Writing / Administrative',
	'Image Analysis',
	'Image Editing',
	'Image Generation',
	'Speech API',
	'Transcription API',
	'Translation API'
]

TEXT_PROMPT_CATEGORIES: List[ str ] = [
	'Writing / Administrative',
	'Research / Academic',
	'Data Analytics & Governance',
	'Software Engineering',
	'Business / Finance / Marketing',
	'Compliance / Legal / Budget',
	'Prompt Engineering',
	'Instruction/ Training / Planning'
]

DOCUMENT_PROMPT_CATEGORIES: List[ str ] = [
	'Research / Academic',
	'Data Analytics & Governance',
	'Business / Finance / Marketing',
	'Compliance / Legal / Budget',
	'Instruction/ Training / Planning',
	'Writing / Administrative'
]

VISION_PROMPT_CATEGORIES: List[ str ] = [ 'Image Analysis' ]

PROMPT_TASK_OPTIONS: List[ str ] = [
	'Chat', 'Analysis', 'Reasoning', 'Coding', 'Writing', 'Editing', 'Summarization',
	'Extraction', 'Classification', 'Translation', 'Comparison', 'Structured Output'
]

RESPONSE_FORMAT_OPTIONS: List[ str ] = [
	'Plain Text', 'Markdown', 'Bullet List', 'Numbered List', 'Markdown Table',
	'JSON', 'XML', 'YAML', 'CSV', 'Code'
]

CODING_LANGUAGE_OPTIONS: List[ str ] = [
	'Python', 'C', 'C++', 'C#', 'Java', 'JavaScript', 'TypeScript', 'SQL', 'VBA',
	'PowerShell', 'Bash', 'HTML', 'CSS', 'Markdown', 'JSON', 'YAML'
]

CODING_TASK_OPTIONS: List[ str ] = [
	'Generate', 'Complete', 'Refactor', 'Debug', 'Review', 'Explain', 'Optimize',
	'Convert', 'Test', 'Document', 'Design'
]

SPOKEN_LANGUAGE_OPTIONS: List[ str ] = [
	'Auto Detect', 'Arabic', 'Bengali', 'Chinese (Simplified)', 'Chinese (Traditional)',
	'Czech', 'Danish', 'Dutch', 'English', 'Finnish', 'French', 'German', 'Greek',
	'Hebrew', 'Hindi', 'Hungarian', 'Indonesian', 'Italian', 'Japanese', 'Korean',
	'Norwegian', 'Persian', 'Polish', 'Portuguese', 'Romanian', 'Russian', 'Spanish',
	'Swedish', 'Thai', 'Turkish', 'Ukrainian', 'Urdu', 'Vietnamese'
]

WRITING_TASK_OPTIONS: List[ str ] = [
	'Draft', 'Rewrite', 'Edit', 'Proofread', 'Expand', 'Condense', 'Reformat'
]

WRITING_TONE_OPTIONS: List[ str ] = [
	'Neutral', 'Professional', 'Formal', 'Conversational', 'Technical', 'Academic'
]

WRITING_AUDIENCE_OPTIONS: List[ str ] = [
	'General', 'Technical', 'Executive', 'Federal', 'Academic'
]

CLASSIFICATION_TYPE_OPTIONS: List[ str ] = [
	'Binary', 'Multi-Class', 'Multi-Label', 'Sentiment', 'Intent', 'Topic', 'Relevance'
]

TASK_DETAIL_OPTIONS: List[ str ] = [ 'Concise', 'Standard', 'Detailed' ]

TASK_FOCUS_OPTIONS: List[ str ] = [ 'Accuracy', 'Balanced', 'Creativity' ]

TRANSLATION_MODE_OPTIONS: List[ str ] = [
	'Natural', 'Literal', 'Formal', 'Technical', 'Localization'
]

RESPONSE_LENGTH_OPTIONS: List[ str ] = [ 'Concise', 'Standard', 'Detailed' ]

VISION_TASK_OPTIONS: List[ str ] = [
	'Extract Visible Text', 'Describe Image', 'Answer Questions', 'Analyze Screenshot',
	'Analyze Chart', 'Analyze Diagram', 'Extract Structured Data', 'Compare Images'
]

VISION_DETAIL_OPTIONS: List[ str ] = [ 'Concise', 'Standard', 'Detailed' ]

GROUNDING_FAILURE_OPTIONS: List[ str ] = [
	'State Insufficient Information', 'Return Retrieved Excerpts', 'Best Supported Answer'
]

RETRIEVAL_BACKEND_OPTIONS: List[ str ] = [ 'Automatic', 'sqlite-vec', 'Cosine Similarity' ]

DOC_ACTION_DETAIL_OPTIONS: List[ str ] = [ 'Concise', 'Standard', 'Detailed' ]

OCR_PAGE_LIMIT_OPTIONS: List[ str ] = [ '1 Page', '2 Pages', '5 Pages', '10 Pages', 'All Pages' ]

DOC_CONTEXT_ORDER_OPTIONS: List[ str ] = [ 'Retrieved First', 'Semantic First' ]


def throw_if( name: str, value: object ) -> None:
	"""Input guard.

	Purpose:
		Validates that a required argument contains a usable value before the surrounding workflow
		continues. This guard centralizes early validation so prompt and runtime routines fail with
		consistent, readable error messages.

	Args:
		name (str): Name value used by the operation.
		value (object): Value value used by the operation.

	Returns:
		None: This function performs its work through validation and does not return a value.
	"""
	if value is None or value == '':
		raise ValueError( f'Argument "{name}" cannot be empty!' )


def get_prompt_categories( ) -> List[ str ]:
	"""Returns the supported prompt categories used by Prompt Engineering controls.

	Purpose:
		Returns the supported prompt categories used by Prompt Engineering controls and the
		category-aware System Instructions selector.

	Returns:
		List[str]: Supported prompt category values.
	"""
	return PROMPT_CATEGORY_OPTIONS.copy( )


def get_prompt_task_types( ) -> List[ str ]:
	"""Returns the supported task types used by Prompt Engineering controls and Text Generation presets.

	Purpose:
		Returns the complete text-task set that the current Gemma 3 text runtime can execute through
		the shared chat-completion path.

	Returns:
		List[str]: Supported task values.
	"""
	return PROMPT_TASK_OPTIONS.copy( )


def get_response_formats( ) -> List[ str ]:
	"""Returns supported response-format values.

	Purpose:
		Provides a single bounded option source for every response-format selectbox used by the
		application.

	Returns:
		List[str]: Supported response-format values.
	"""
	return RESPONSE_FORMAT_OPTIONS.copy( )


def get_spoken_languages( include_auto_detect: bool = True ) -> List[ str ]:
	"""Returns supported human-language values for language selectboxes.

	Purpose:
		Provides bounded language options for translation and prompt-generation controls so users do
		not need to enter language names manually.

	Args:
		include_auto_detect (bool): True to include the Auto Detect source-language option.

	Returns:
		List[str]: Supported language values.
	"""
	if include_auto_detect:
		return SPOKEN_LANGUAGE_OPTIONS.copy( )
	return [ language for language in SPOKEN_LANGUAGE_OPTIONS if language != 'Auto Detect' ]


def fetch_prompt_categories( db_path: str ) -> List[ str ]:
	"""Retrieves distinct persisted prompt categories.

	Purpose:
		Returns the category values that actually exist in the Prompts table for Prompt Engineering and
		database-management workflows. The application does not rewrite or normalize stored values.

	Args:
		db_path (str): SQLite database path.

	Returns:
		List[str]: Sorted persisted category values.
	"""
	try:
		throw_if( 'db_path', db_path )
		with sqlite3.connect( db_path ) as conn:
			rows = conn.execute(
				"""SELECT DISTINCT Category
				   FROM Prompts
				   WHERE Category IS NOT NULL AND TRIM(Category) <> ''
				   ORDER BY Category;"""
			).fetchall( )
		return [ str( row[ 0 ] ) for row in rows if row and row[ 0 ] ]
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'fetch_prompt_categories'
		exception.method = 'fetch_prompt_categories( db_path: str ) -> List[str]'
		Logger( ).write( exception )
		return [ ]


def get_mode_prompt_category_policy( workflow: str ) -> List[ str ]:
	"""Returns the persisted category policy for a model-facing workflow.

	Purpose:
		Maps each execution mode to the existing project/database category names that are appropriate
		for that workflow without altering the Prompts table.

	Args:
		workflow (str): Text Generation, Document Q&A, or Image to Text workflow name.

	Returns:
		List[str]: Permitted persisted category values.
	"""
	throw_if( 'workflow', workflow )
	workflow_value = str( workflow ).strip( )
	if workflow_value == 'Document Q&A':
		return DOCUMENT_PROMPT_CATEGORIES.copy( )
	if workflow_value == 'Image to Text':
		return VISION_PROMPT_CATEGORIES.copy( ) if vision_runtime_available( ) else [ ]
	return TEXT_PROMPT_CATEGORIES.copy( )


def fetch_mode_prompt_categories( db_path: str, workflow: str ) -> List[ str ]:
	"""Retrieves populated prompt categories permitted for a model-facing workflow.

	Purpose:
		Intersects the workflow category policy with categories that actually contain usable prompt
		templates. This prevents empty or incompatible categories from appearing in model-facing
		selectors while preserving every database record unchanged.

	Args:
		db_path (str): SQLite database path.
		workflow (str): Model-facing workflow name.

	Returns:
		List[str]: Populated persisted categories valid for the selected workflow.
	"""
	try:
		throw_if( 'db_path', db_path )
		throw_if( 'workflow', workflow )
		policy = get_mode_prompt_category_policy( workflow )
		if not policy:
			return [ ]
		placeholders = ', '.join( [ '?' ] * len( policy ) )
		query = (
			'SELECT DISTINCT Category FROM Prompts '
			'WHERE Category IN (' + placeholders + ') '
			'AND Category IS NOT NULL AND TRIM(Category) <> \'\' '
			'AND Text IS NOT NULL AND TRIM(Text) <> \'\' '
			'ORDER BY Category;'
		)
		with sqlite3.connect( db_path ) as conn:
			rows = conn.execute( query, tuple( policy ) ).fetchall( )
		persisted = { str( row[ 0 ] ) for row in rows if row and row[ 0 ] }
		return [ category for category in policy if category in persisted ]
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'fetch_mode_prompt_categories'
		exception.method = 'fetch_mode_prompt_categories( db_path: str, workflow: str ) -> List[str]'
		Logger( ).write( exception )
		return [ ]


def is_prompt_category_allowed_for_workflow( category: str, workflow: str ) -> bool:
	"""Determines whether a persisted category can be applied to a model-facing workflow.

	Purpose:
		Validates Prompt Engineering cascade actions using the existing database taxonomy rather than
		invented replacement labels.

	Args:
		category (str): Persisted prompt category.
		workflow (str): Target workflow.

	Returns:
		bool: True when the category is permitted for the target workflow.
	"""
	if not category or not workflow:
		return False
	return str( category ).strip( ) in get_mode_prompt_category_policy( workflow )


def fetch_prompt_names( db_path: str, category: str = '' ) -> List[ str ]:
	"""Retrieves sorted prompt captions from the local Prompts table.

	Purpose:
		Retrieves prompt captions for the selected category while preserving the existing helper used
		by prompt-template controls.

	Args:
		db_path (str): SQLite database path.
		category (str): Optional persisted category used to filter prompt captions.

	Returns:
		List[str]: Sorted prompt captions.
	"""
	try:
		throw_if( 'db_path', db_path )
		with sqlite3.connect( db_path ) as conn:
			if category:
				rows = conn.execute(
					'''SELECT Caption FROM Prompts
					   WHERE Category = ? AND Caption IS NOT NULL
					   ORDER BY Caption, ID;''',
					(category,)
				).fetchall( )
			else:
				rows = conn.execute(
					'''SELECT Caption FROM Prompts
					   WHERE Caption IS NOT NULL ORDER BY Caption, ID;'''
				).fetchall( )
		return [ str( row[ 0 ] ) for row in rows if row and row[ 0 ] is not None ]
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'fetch_prompt_names'
		exception.method = 'fetch_prompt_names( db_path: str, category: str ) -> List[str]'
		Logger( ).write( exception )
		return [ ]


def fetch_prompt_options( db_path: str, category: str ) -> List[ Tuple[ int, str ] ]:
	"""Retrieves prompt identifiers and captions for a selected category.

	Purpose:
		Builds ID-backed template options so duplicate captions cannot cause the wrong prompt record to
		be loaded.

	Args:
		db_path (str): SQLite database path.
		category (str): Persisted prompt category.

	Returns:
		List[Tuple[int, str]]: Prompt ID and caption pairs ordered for display.
	"""
	try:
		throw_if( 'db_path', db_path )
		throw_if( 'category', category )
		with sqlite3.connect( db_path ) as conn:
			rows = conn.execute(
				'''SELECT ID, Caption FROM Prompts
				   WHERE Category = ? AND Text IS NOT NULL AND TRIM(Text) <> ''
				   ORDER BY Caption, ID;''',
				(category,)
			).fetchall( )
		return [ (int( row[ 0 ] ), str( row[ 1 ] or f'Prompt {row[ 0 ]}' )) for row in rows ]
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'fetch_prompt_options'
		exception.method = 'fetch_prompt_options( db_path: str, category: str ) -> List[Tuple[int, str]]'
		Logger( ).write( exception )
		return [ ]


def fetch_prompt_text( db_path: str, name: str ) -> str | None:
	"""Retrieves prompt template text for a selected caption.

	Purpose:
		Retains caption-based lookup for legacy callers while deterministic System Instructions controls
		use primary-key lookup.

	Args:
		db_path (str): SQLite database path.
		name (str): Prompt caption.

	Returns:
		str | None: Prompt text when available.
	"""
	try:
		throw_if( 'db_path', db_path )
		throw_if( 'name', name )
		with sqlite3.connect( db_path ) as conn:
			row = conn.execute(
				'SELECT Text FROM Prompts WHERE Caption = ? ORDER BY ID LIMIT 1;',
				(name,)
			).fetchone( )
		return str( row[ 0 ] ) if row and row[ 0 ] is not None else None
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'fetch_prompt_text'
		exception.method = 'fetch_prompt_text( db_path: str, name: str ) -> str | None'
		Logger( ).write( exception )
		return None


def fetch_prompts_df( ) -> pd.DataFrame:
	"""Builds a prompt-management DataFrame from the Prompts table.

	Purpose:
		Builds the Prompt Engineering data surface from the authoritative five-column Prompts schema and
		adds a selection column for Streamlit editing.

	Returns:
		pd.DataFrame: Prompt-management DataFrame.
	"""
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		df_prompts = pd.read_sql_query(
			'SELECT ID, Caption, Name, Category, Text FROM Prompts ORDER BY ID DESC;', conn )
	df_prompts.insert( 0, 'Selected', False )
	return df_prompts


def fetch_prompt_by_id( pid: int ) -> Dict[ str, Any ] | None:
	"""Retrieves one prompt record by primary key.

	Purpose:
		Retrieves the authoritative prompt record used by System Instructions and Prompt Engineering.

	Args:
		pid (int): Prompt primary key.

	Returns:
		Dict[str, Any] | None: Prompt record when found.
	"""
	try:
		throw_if( 'pid', pid )
		with sqlite3.connect( cfg.DB_PATH ) as conn:
			cur = conn.execute(
				'SELECT ID, Caption, Name, Category, Text FROM Prompts WHERE ID = ?;',
				(pid,)
			)
			row = cur.fetchone( )
			return dict( zip( [ column[ 0 ] for column in cur.description ], row ) ) if row else None
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'fetch_prompt_by_id'
		exception.method = 'fetch_prompt_by_id( pid: int ) -> Dict[str, Any] | None'
		Logger( ).write( exception )
		return None


def fetch_prompt_by_name( name: str ) -> Dict[ str, Any ] | None:
	"""Retrieves one prompt record by caption.

	Purpose:
		Retains caption lookup for legacy behavior while primary-key lookup remains authoritative for
		interactive selection.

	Args:
		name (str): Prompt caption.

	Returns:
		Dict[str, Any] | None: Prompt record when found.
	"""
	try:
		throw_if( 'name', name )
		with sqlite3.connect( cfg.DB_PATH ) as conn:
			cur = conn.execute(
				'''SELECT ID, Caption, Name, Category, Text FROM Prompts
				   WHERE Caption = ? ORDER BY ID LIMIT 1;''',
				(name,)
			)
			row = cur.fetchone( )
			return dict( zip( [ column[ 0 ] for column in cur.description ], row ) ) if row else None
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'fetch_prompt_by_name'
		exception.method = 'fetch_prompt_by_name( name: str ) -> Dict[str, Any] | None'
		Logger( ).write( exception )
		return None


def insert_prompt( data: Dict[ str, Any ] ) -> int:
	"""Inserts a prompt-template record into the local Prompts table.

	Purpose:
		Persists Caption, Name, Category, and Text while allowing SQLite to assign the immutable ID.

	Args:
		data (Dict[str, Any]): Prompt record fields to write.

	Returns:
		int: Newly generated prompt ID.
	"""
	throw_if( 'data', data )
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		cur = conn.execute(
			'INSERT INTO Prompts (Caption, Name, Category, Text) VALUES (?, ?, ?, ?);',
			(data[ 'Caption' ], data[ 'Name' ], data[ 'Category' ], data[ 'Text' ])
		)
		conn.commit( )
		return int( cur.lastrowid )


def update_prompt( pid: int, data: Dict[ str, Any ] ) -> None:
	"""Updates an existing prompt-template record by primary key.

	Purpose:
		Updates mutable prompt metadata without modifying the SQLite-assigned ID.

	Args:
		pid (int): Prompt primary key.
		data (Dict[str, Any]): Prompt record fields to write.

	Returns:
		None: This function performs its work through database side effects.
	"""
	throw_if( 'pid', pid )
	throw_if( 'data', data )
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute(
			'''UPDATE Prompts SET Caption = ?, Name = ?, Category = ?, Text = ?
			   WHERE ID = ?;''',
			(data[ 'Caption' ], data[ 'Name' ], data[ 'Category' ], data[ 'Text' ], pid)
		)
		conn.commit( )


def delete_prompt( pid: int ) -> None:
	"""Deletes a prompt-template record from the local Prompts table by primary key.

	Purpose:
		Deletes the selected prompt without relying on caption uniqueness.

	Args:
		pid (int): Prompt primary key.

	Returns:
		None: This function performs its work through database side effects.
	"""
	throw_if( 'pid', pid )
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute( 'DELETE FROM Prompts WHERE ID = ?;', (pid,) )
		conn.commit( )


def get_effective_system_instructions( ) -> str:
	"""Returns the active system-instruction text from Streamlit session state.

	Purpose:
		Returns the active instruction text used by every model-facing text execution path.

	Returns:
		str: Active system instructions.
	"""
	text = st.session_state.get( 'system_instructions', '' )
	return str( text ).strip( ) if text is not None else ''


def build_task_instruction_block( ) -> str:
	"""Builds task-specific instruction text from Text Generation controls.

	Purpose:
		Converts bounded UI selections into a single task contract consumed by the shared Gemma 3 chat
		path.

	Returns:
		str: Task instruction block.
	"""
	task_preset = str( st.session_state.get( 'task_preset', 'Chat' ) or 'Chat' ).strip( )
	response_format = str( st.session_state.get( 'response_format', 'Markdown' ) or 'Markdown' ).strip( )
	response_language = str( st.session_state.get( 'response_language', 'English' ) or 'English' ).strip( )
	reasoning_depth = str( st.session_state.get( 'reasoning_depth', 'Medium' ) or 'Medium' ).strip( )
	answer_only = bool( st.session_state.get( 'answer_only', False ) )
	use_self_check = bool( st.session_state.get( 'use_self_check', False ) )
	deterministic_reasoning = bool( st.session_state.get( 'deterministic_reasoning', False ) )
	coding_language = str( st.session_state.get( 'coding_language', 'Python' ) or 'Python' ).strip( )
	coding_task = str( st.session_state.get( 'coding_task', 'Generate' ) or 'Generate' ).strip( )
	coding_include_comments = bool( st.session_state.get( 'coding_include_comments', True ) )
	coding_editor_format = bool( st.session_state.get( 'coding_editor_format', True ) )
	coding_fenced_output = bool( st.session_state.get( 'coding_fenced_output', True ) )
	translation_source_language = str(
		st.session_state.get( 'translation_source_language', 'Auto Detect' ) or 'Auto Detect' ).strip( )
	translation_target_language = str(
		st.session_state.get( 'translation_target_language', 'English' ) or 'English' ).strip( )
	writing_task = str( st.session_state.get( 'writing_task', 'Draft' ) or 'Draft' ).strip( )
	writing_tone = str( st.session_state.get( 'writing_tone', 'Professional' ) or 'Professional' ).strip( )
	writing_audience = str( st.session_state.get( 'writing_audience', 'General' ) or 'General' ).strip( )
	classification_type = str(
		st.session_state.get( 'classification_type', 'Multi-Class' ) or 'Multi-Class' ).strip( )
	classification_confidence = bool( st.session_state.get( 'classification_confidence', False ) )
	classification_allow_unknown = bool( st.session_state.get( 'classification_allow_unknown', True ) )
	task_detail = str( st.session_state.get( 'task_detail', 'Standard' ) or 'Standard' ).strip( )
	task_focus = str( st.session_state.get( 'task_focus', 'Balanced' ) or 'Balanced' ).strip( )
	translation_mode = str( st.session_state.get( 'translation_mode', 'Natural' ) or 'Natural' ).strip( )
	translation_preserve_formatting = bool( st.session_state.get( 'translation_preserve_formatting', True ) )
	writing_length = str( st.session_state.get( 'writing_length', 'Standard' ) or 'Standard' ).strip( )
	classification_explain = bool( st.session_state.get( 'classification_explain', False ) )
	response_length = str( st.session_state.get( 'response_length', 'Standard' ) or 'Standard' ).strip( )
	response_include_headings = bool( st.session_state.get( 'response_include_headings', True ) )

	lines: List[ str ] = [ 'Task Preset:', f'- Active Task: {task_preset}',
		f'- Task Detail: {task_detail}', f'- Task Focus: {task_focus}',
		f'- Response Format: {response_format}', f'- Response Language: {response_language}',
		f'- Response Length: {response_length}' ]
	if response_include_headings:
		lines.append( '- Use descriptive headings when they improve readability.' )
	if bool( st.session_state.get( 'is_grounded', False ) ):
		lines.append( '- Ground claims in supplied context and state when the context is insufficient.' )

	if task_preset in ( 'Analysis', 'Reasoning' ):
		lines.append( f'- Analysis Depth: {reasoning_depth}' )
		lines.append( '- Analyze the request carefully and return a clear, supported conclusion.' )
		if answer_only:
			lines.append( '- Return the final answer without extra prefatory narration.' )
		if use_self_check:
			lines.append( '- Verify the conclusion against the supplied information before answering.' )
		if deterministic_reasoning:
			lines.append( '- Prefer stable, conservative reasoning over creative variation.' )
	elif task_preset == 'Coding':
		lines.append( f'- Code Language: {coding_language}' )
		lines.append( f'- Coding Task: {coding_task}' )
		lines.append( '- Produce executable, editor-ready source code when code is requested.' )
		if coding_include_comments:
			lines.append( '- Include useful documentation and inline comments when appropriate.' )
		else:
			lines.append( '- Minimize comments unless required for clarity.' )
		if coding_editor_format:
			lines.append( '- Format generated code as editor-ready source rather than pseudo-code.' )
		if coding_fenced_output:
			lines.append( '- Return generated code inside fenced Markdown code blocks.' )
		else:
			lines.append( '- Return generated code without fenced Markdown blocks.' )
	elif task_preset in ( 'Writing', 'Editing' ):
		lines.append( f'- Writing Operation: {writing_task}' )
		lines.append( f'- Tone: {writing_tone}' )
		lines.append( f'- Audience: {writing_audience}' )
		lines.append( f'- Writing Length: {writing_length}' )
		lines.append( '- Preserve the supplied facts and requested intent.' )
	elif task_preset == 'Translation':
		lines.append( f'- Source Language: {translation_source_language}' )
		lines.append( f'- Target Language: {translation_target_language}' )
		lines.append( f'- Translation Mode: {translation_mode}' )
		if translation_preserve_formatting:
			lines.append( '- Preserve source formatting and structural cues where practical.' )
		lines.append( '- Preserve original meaning, tone, terminology, and structure where practical.' )
	elif task_preset == 'Summarization':
		lines.append( '- Summarize the supplied content clearly and faithfully.' )
		lines.append( '- Preserve material facts, names, dates, and conclusions.' )
	elif task_preset == 'Extraction':
		lines.append( '- Extract only requested facts supported by the supplied content.' )
		lines.append( '- Do not invent values that are not present.' )
	elif task_preset == 'Classification':
		lines.append( f'- Classification Type: {classification_type}' )
		lines.append( '- Classify only from the evidence supplied in the request.' )
		if classification_confidence:
			lines.append( '- Include a concise confidence assessment with the classification.' )
		if classification_allow_unknown:
			lines.append( '- Use Unknown when the supplied evidence does not support a reliable class.' )
		if classification_explain:
			lines.append( '- Briefly explain the evidence supporting the selected classification.' )
	elif task_preset == 'Comparison':
		lines.append( '- Compare the supplied items using consistent criteria and identify material differences.' )
	elif task_preset == 'Structured Output':
		lines.append( '- Follow the requested output structure exactly and omit unrelated prose.' )
	else:
		lines.append( '- Respond as a general-purpose assistant.' )

	if response_format == 'JSON':
		lines.append( '- Return one valid JSON object and no surrounding commentary.' )
	elif response_format == 'XML':
		lines.append( '- Return well-formed XML.' )
	elif response_format == 'YAML':
		lines.append( '- Return valid YAML.' )
	elif response_format == 'CSV':
		lines.append( '- Return valid CSV with a header row.' )
	elif response_format == 'Markdown Table':
		lines.append( '- Return the primary result as a Markdown table.' )
	elif response_format == 'Bullet List':
		lines.append( '- Return the primary result as a concise bullet list.' )
	elif response_format == 'Numbered List':
		lines.append( '- Return the primary result as a numbered list.' )

	return '\n'.join( lines ).strip( )


def build_effective_prompt_preview( user_input: str ) -> str:
	"""Builds a readable preview of the system, task, and user prompt content.

	Purpose:
		Shows the effective logical content sent through the Gemma 3 chat-template path without
		exposing model-specific control tokens.

	Args:
		user_input (str): User request text.

	Returns:
		str: Human-readable prompt preview.
	"""
	system_instructions = get_effective_system_instructions( )
	task_block = build_task_instruction_block( )
	preview_parts: List[ str ] = [ ]
	if system_instructions:
		preview_parts.extend( [ '[System Instructions]', system_instructions ] )
	if task_block:
		preview_parts.extend( [ '[Task Instructions]', task_block ] )
	preview_parts.extend( [ '[User Input]', user_input or '' ] )
	return '\n\n'.join( preview_parts ).strip( )


def get_runtime_llm( ) -> Llama:
	"""Loads or retrieves the cached llama.cpp runtime.

	Purpose:
		Loads the configured GGUF model using current context-window and CPU-thread settings and fails
		with a controlled error when the model cannot be initialized.

	Returns:
		Llama: Loaded llama.cpp model runtime.
	"""
	ctx_value = int( st.session_state.get( 'context_window', cfg.DEFAULT_CTX ) or cfg.DEFAULT_CTX )
	thread_value = int( st.session_state.get( 'cpu_threads', cfg.CORES ) or cfg.CORES )
	repeat_window_value = int( st.session_state.get( 'repeat_window', 0 ) or 0 )
	if ctx_value <= 0:
		ctx_value = int( cfg.DEFAULT_CTX )
	if thread_value <= 0:
		thread_value = int( cfg.CORES )
	if repeat_window_value <= 0:
		repeat_window_value = 64
	batch_size_value = int( st.session_state.get( 'batch_size', 512 ) or 512 )
	micro_batch_size_value = int( st.session_state.get( 'micro_batch_size', 128 ) or 128 )
	runtime_llm = load_llm( ctx_value, thread_value, repeat_window_value, batch_size_value,
		micro_batch_size_value )
	if runtime_llm is None:
		raise RuntimeError( 'The configured Gemma 3 GGUF model could not be loaded.' )
	return runtime_llm


def build_chat_messages( user_input: str ) -> List[ Dict[ str, str ] ]:
	"""Builds the chat-message sequence consumed by llama.cpp.

	Purpose:
		Combines system instructions, task controls, semantic context, document context, and prior chat
		history without manually emitting model-specific special tokens or duplicating the current user
		turn.

	Args:
		user_input (str): Current user request text.

	Returns:
		List[Dict[str, str]]: Ordered chat messages for Gemma 3 chat-template serialization.
	"""
	throw_if( 'user_input', user_input )
	system_instructions = get_effective_system_instructions( )
	task_block = build_task_instruction_block( )
	use_semantic = bool( st.session_state.get( 'use_semantic', False ) )
	use_chat_history = bool( st.session_state.get( 'use_chat_history', True ) )
	use_document_context = bool( st.session_state.get( 'use_document_context', False ) )
	basic_docs = st.session_state.get( 'basic_docs', [ ] )
	messages = st.session_state.get( 'messages', [ ] )
	semantic_top_k_value = int( st.session_state.get( 'semantic_top_k', 8 ) or 8 )

	system_parts: List[ str ] = [ ]
	if system_instructions:
		system_parts.append( system_instructions )
	if task_block:
		system_parts.append( task_block )

	if use_semantic and embedder is not None:
		try:
			with sqlite3.connect( cfg.DB_PATH ) as conn:
				rows = conn.execute( 'SELECT chunk, vector FROM embeddings' ).fetchall( )
			if rows:
				query_vector = np.asarray( embedder.encode( [ user_input ] )[ 0 ], dtype=np.float32 )
				scored: List[ Tuple[ str, float ] ] = [ ]
				for chunk, vector in rows:
					if not vector:
						continue
					stored_vector = np.frombuffer( vector, dtype=np.float32 )
					if stored_vector.size != query_vector.size:
						continue
					scored.append( (str( chunk or '' ), cosine_similarity( query_vector, stored_vector )) )
				for chunk, _score in sorted( scored, key=lambda item: item[ 1 ], reverse=True )[ :semantic_top_k_value ]:
					if chunk:
						system_parts.append( f'Semantic Context:\n{chunk}' )
		except Exception as e:
			exception = Error( e )
			exception.module = 'app'
			exception.cause = 'build_chat_messages'
			exception.method = 'build_chat_messages( user_input: str ) -> List[Dict[str, str]]'
			Logger( ).write( exception )

	if use_document_context and isinstance( basic_docs, list ):
		for document_text in basic_docs[ :6 ]:
			if document_text:
				system_parts.append( f'Document Context:\n{document_text}' )

	chat_messages: List[ Dict[ str, str ] ] = [ ]
	if system_parts:
		chat_messages.append( { 'role': 'system', 'content': '\n\n'.join( system_parts ).strip( ) } )

	if use_chat_history and isinstance( messages, list ):
		history_messages: List[ Dict[ str, str ] ] = [ ]
		for message_index, message in enumerate( messages ):
			role = ''
			content = ''
			if isinstance( message, (tuple, list) ) and len( message ) == 2:
				role = str( message[ 0 ] or '' ).strip( )
				content = str( message[ 1 ] or '' )
			elif isinstance( message, dict ):
				role = str( message.get( 'role', '' ) or '' ).strip( )
				content = str( message.get( 'content', '' ) or '' )
			if role not in ( 'user', 'assistant' ) or not content:
				continue
			is_current_turn = (
				role == 'user' and content == user_input and message_index == len( messages ) - 1 )
			if not is_current_turn:
				history_messages.append( { 'role': role, 'content': content } )

		expected_role = 'user'
		normalized_history: List[ Dict[ str, str ] ] = [ ]
		for history_message in history_messages:
			if history_message[ 'role' ] != expected_role:
				continue
			normalized_history.append( history_message )
			expected_role = 'assistant' if expected_role == 'user' else 'user'
		if normalized_history and normalized_history[ -1 ][ 'role' ] == 'user':
			normalized_history.pop( )
		chat_messages.extend( normalized_history )

	chat_messages.append( { 'role': 'user', 'content': user_input } )
	return chat_messages


def build_prompt( user_input: str ) -> str:
	"""Builds a readable compatibility prompt from the current chat-message sequence.

	Purpose:
		Preserves the legacy helper for diagnostics while model execution uses llama.cpp chat-template
		serialization through create_chat_completion.

	Args:
		user_input (str): Current user request text.

	Returns:
		str: Readable role-separated prompt representation.
	"""
	messages = build_chat_messages( user_input )
	parts: List[ str ] = [ ]
	for message in messages:
		parts.append( f'[{message[ "role" ].title( )}]\n{message[ "content" ]}' )
	return '\n\n'.join( parts ).strip( )


def run_direct_llm_turn( system_instruction: str, user_input: str, temperature: float,
		top_p: float, repeat_penalty: float, max_tokens: int, stream: bool,
		output: Any = None, response_format: str = 'Markdown' ) -> str:
	"""Executes an isolated Gemma text turn without shared Text Generation context.

	Purpose:
		Provides Prompt Generator and Document Q&A with a clean execution path that does not inherit
		Text Generation task presets, chat history, document context, or semantic context.

	Args:
		system_instruction (str): Workflow-specific system instruction.
		user_input (str): Workflow-specific user request.
		temperature (float): Sampling temperature.
		top_p (float): Nucleus-sampling probability.
		repeat_penalty (float): Repeat penalty.
		max_tokens (int): Maximum generated tokens.
		stream (bool): True to stream output.
		output (Any): Optional Streamlit output placeholder.
		response_format (str): Response-format contract.

	Returns:
		str: Generated response text.
	"""
	if not user_input or not str( user_input ).strip( ):
		return ''
	try:
		runtime_llm = get_runtime_llm( )
		messages: List[ Dict[ str, str ] ] = [ ]
		if system_instruction and str( system_instruction ).strip( ):
			messages.append( { 'role': 'system', 'content': str( system_instruction ).strip( ) } )
		messages.append( { 'role': 'user', 'content': str( user_input ).strip( ) } )
		generation_args: Dict[ str, Any ] = {
			'messages': messages,
			'max_tokens': int( max_tokens ) if int( max_tokens ) > 0 else 1024,
			'temperature': float( temperature ),
			'top_p': float( top_p ),
			'repeat_penalty': float( repeat_penalty ),
			'stream': stream
		}
		top_k_value = int( st.session_state.get( 'top_k', 0 ) )
		seed_value = int( st.session_state.get( 'random_seed', 0 ) )
		frequency_penalty_value = float( st.session_state.get( 'frequency_penalty', 0.0 ) )
		presence_penalty_value = float( st.session_state.get( 'presense_penalty', 0.0 ) )
		if top_k_value > 0:
			generation_args[ 'top_k' ] = top_k_value
		if seed_value > 0:
			generation_args[ 'seed' ] = seed_value
		if frequency_penalty_value > 0:
			generation_args[ 'frequency_penalty' ] = frequency_penalty_value
		if presence_penalty_value > 0:
			generation_args[ 'presence_penalty' ] = presence_penalty_value
		if response_format == 'JSON':
			generation_args[ 'response_format' ] = { 'type': 'json_object' }
		if not stream:
			response = runtime_llm.create_chat_completion( **generation_args )
			choices = response.get( 'choices', [ ] ) if isinstance( response, dict ) else [ ]
			if not choices:
				return ''
			return str( choices[ 0 ].get( 'message', { } ).get( 'content', '' ) or '' ).strip( )
		buffer = ''
		if output is None:
			output = st.empty( )
		for chunk in runtime_llm.create_chat_completion( **generation_args ):
			if not isinstance( chunk, dict ):
				continue
			choices = chunk.get( 'choices', [ ] )
			if not choices:
				continue
			piece = str( choices[ 0 ].get( 'delta', { } ).get( 'content', '' ) or '' )
			if piece:
				buffer += piece
				output.markdown( buffer + '▌' )
		output.markdown( buffer )
		return buffer.strip( )
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'run_direct_llm_turn'
		exception.method = 'run_direct_llm_turn( ... ) -> str'
		Logger( ).write( exception )
		message = f'Generation failed: {e}'
		if output is not None:
			output.error( message )
		else:
			st.error( message )
		return ''


def run_llm_turn( user_input: str, temperature: float, top_p: float, repeat_penalty: float,
		max_tokens: int, stream: bool, output: Any = None ) -> str:
	"""Executes one local-model generation turn through Gemma-compatible chat completion.

	Purpose:
		Uses the GGUF chat template exposed through llama.cpp, connects all applicable inference
		controls, handles streaming and non-streaming responses, and converts runtime failures into a
		controlled Streamlit-visible error instead of an uncaught exception.

	Args:
		user_input (str): User request text.
		temperature (float): Sampling temperature.
		top_p (float): Nucleus-sampling probability.
		repeat_penalty (float): Repeat penalty.
		max_tokens (int): Maximum generated tokens.
		stream (bool): True to stream output tokens.
		output (Any): Optional Streamlit output placeholder.

	Returns:
		str: Generated assistant response, or an empty string when generation fails.
	"""
	if user_input is None or not str( user_input ).strip( ):
		return ''

	try:
		runtime_llm = get_runtime_llm( )
		chat_messages = build_chat_messages( str( user_input ).strip( ) )
		max_token_value = int( max_tokens ) if int( max_tokens ) > 0 else 1024
		temperature_value = float( temperature )
		top_p_value = float( top_p )
		repeat_penalty_value = float( repeat_penalty )
		top_k_value = int( st.session_state.get( 'top_k', 0 ) )
		seed_value = int( st.session_state.get( 'random_seed', 0 ) )
		frequency_penalty_value = float( st.session_state.get( 'frequency_penalty', 0.0 ) )
		presence_penalty_value = float( st.session_state.get( 'presense_penalty', 0.0 ) )
		response_format_value = str( st.session_state.get( 'response_format', 'Markdown' ) )

		generation_args: Dict[ str, Any ] = {
			'messages': chat_messages,
			'max_tokens': max_token_value,
			'temperature': temperature_value,
			'top_p': top_p_value,
			'repeat_penalty': repeat_penalty_value,
			'stream': stream
		}
		if top_k_value > 0:
			generation_args[ 'top_k' ] = top_k_value
		if seed_value > 0:
			generation_args[ 'seed' ] = seed_value
		if frequency_penalty_value > 0:
			generation_args[ 'frequency_penalty' ] = frequency_penalty_value
		if presence_penalty_value > 0:
			generation_args[ 'presence_penalty' ] = presence_penalty_value
		if response_format_value == 'JSON':
			generation_args[ 'response_format' ] = { 'type': 'json_object' }

		if not stream:
			response = runtime_llm.create_chat_completion( **generation_args )
			choices = response.get( 'choices', [ ] ) if isinstance( response, dict ) else [ ]
			if not choices:
				return ''
			message = choices[ 0 ].get( 'message', { } )
			return str( message.get( 'content', '' ) or '' ).strip( )

		buffer = ''
		if output is None:
			output = st.empty( )
		for chunk in runtime_llm.create_chat_completion( **generation_args ):
			if not isinstance( chunk, dict ):
				continue
			choices = chunk.get( 'choices', [ ] )
			if not choices:
				continue
			delta = choices[ 0 ].get( 'delta', { } )
			text = str( delta.get( 'content', '' ) or '' )
			if text:
				buffer += text
				output.markdown( buffer + '▌' )
		output.markdown( buffer )
		return buffer.strip( )
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'run_llm_turn'
		exception.method = 'run_llm_turn( user_input, temperature, top_p, repeat_penalty, max_tokens, stream, output ) -> str'
		Logger( ).write( exception )
		message = f'Generation failed: {e}'
		if output is not None:
			output.error( message )
		else:
			st.error( message )
		return ''


def vision_runtime_available( ) -> bool:
	"""Determines whether the configured Image-to-Text runtime is available.

	Purpose:
		Confirms model/projector files exist and the installed llama-cpp-python build exposes the MTMD
		chat handler required by the configured Gemma 3 multimodal path.

	Returns:
		bool: True when the required local assets and MTMD handler are available.
	"""
	if not local_model_available( ) or MMPROJ_PATH_OBJ is None or not MMPROJ_PATH_OBJ.exists( ):
		return False
	try:
		from llama_cpp.llama_chat_format import MTMDChatHandler
		return MTMDChatHandler is not None
	except Exception:
		return False


@st.cache_resource
def load_vision_llm( ctx: int, threads: int, repeat_window: int, batch_size: int,
		micro_batch_size: int, mmproj_path: str ) -> Any | None:
	"""Loads the Gemma 3 multimodal llama.cpp runtime.

	Purpose:
		Loads Gemma 3 with llama-cpp-python's MTMD chat handler and the matching multimodal projector
		so local image bytes can be supplied as OpenAI-compatible image_url content parts.

	Args:
		ctx (int): Context-window size.
		threads (int): CPU thread count.
		repeat_window (int): Repeat-history token window.
		batch_size (int): llama.cpp logical batch size.
		micro_batch_size (int): llama.cpp physical micro-batch size.
		mmproj_path (str): Local multimodal projector path.

	Returns:
		Any | None: Loaded multimodal runtime when available; otherwise None.
	"""
	try:
		throw_if( 'mmproj_path', mmproj_path )
		from llama_cpp import Llama
		from llama_cpp.llama_chat_format import MTMDChatHandler
		projector_path = Path( mmproj_path )
		if not local_model_available( ) or not projector_path.exists( ):
			return None
		ctx_value = int( ctx ) if int( ctx ) > 0 else int( cfg.DEFAULT_CTX )
		thread_value = int( threads ) if int( threads ) > 0 else int( cfg.CORES )
		repeat_window_value = int( repeat_window ) if int( repeat_window ) > 0 else 64
		batch_size_value = int( batch_size ) if int( batch_size ) > 0 else 512
		micro_batch_size_value = int( micro_batch_size ) if int( micro_batch_size ) > 0 else 128
		micro_batch_size_value = min( micro_batch_size_value, batch_size_value )
		chat_handler = MTMDChatHandler( clip_model_path=str( projector_path ), verbose=False )
		return Llama( model_path=str( cfg.MODEL_PATH ), chat_handler=chat_handler,
			n_ctx=ctx_value, n_threads=thread_value, n_batch=batch_size_value,
			n_ubatch=micro_batch_size_value, last_n_tokens_size=repeat_window_value, verbose=False )
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'load_vision_llm'
		exception.method = ('load_vision_llm( ctx, threads, repeat_window, batch_size, '
			'micro_batch_size, mmproj_path ) -> Any | None')
		Logger( ).write( exception )
		return None


def get_vision_runtime_llm( ) -> Any:
	"""Loads or retrieves the cached Image-to-Text runtime.

	Purpose:
		Builds the multimodal runtime from the same context/runtime controls used by text generation and
		fails through a controlled exception when the projector is unavailable.

	Returns:
		Any: Loaded llama.cpp multimodal runtime.
	"""
	if not vision_runtime_available( ) or MMPROJ_PATH_OBJ is None:
		raise RuntimeError( 'A compatible Gemma 3 mmproj GGUF file is required for Image-to-Text.' )
	ctx_value = int( st.session_state.get( 'context_window', cfg.DEFAULT_CTX ) or cfg.DEFAULT_CTX )
	thread_value = int( st.session_state.get( 'cpu_threads', cfg.CORES ) or cfg.CORES )
	repeat_window_value = int( st.session_state.get( 'repeat_window', 64 ) or 64 )
	batch_size_value = int( st.session_state.get( 'batch_size', 512 ) or 512 )
	micro_batch_size_value = int( st.session_state.get( 'micro_batch_size', 128 ) or 128 )
	runtime_llm = load_vision_llm( ctx_value, thread_value, repeat_window_value, batch_size_value,
		micro_batch_size_value, str( MMPROJ_PATH_OBJ ) )
	if runtime_llm is None:
		raise RuntimeError( 'The Gemma 3 multimodal runtime could not be initialized.' )
	return runtime_llm


def image_bytes_to_data_uri( image_bytes: bytes, mime_type: str ) -> str:
	"""Converts local image bytes to a data URI.

	Purpose:
		Encodes uploaded local image bytes into the image_url representation accepted by
		llama-cpp-python's MTMD chat handler without making outbound network requests.

	Args:
		image_bytes (bytes): Image content.
		mime_type (str): Image MIME type.

	Returns:
		str: Base64 data URI.
	"""
	throw_if( 'image_bytes', image_bytes )
	mime_value = str( mime_type or 'image/png' ).strip( )
	encoded = base64.b64encode( image_bytes ).decode( 'utf-8' )
	return f'data:{mime_value};base64,{encoded}'


def build_vision_instruction( user_input: str ) -> str:
	"""Builds the Image-to-Text task instruction.

	Purpose:
		Converts bounded vision controls into a direct image-understanding request while preserving an
		optional user-authored question as the highest-specificity task instruction.

	Args:
		user_input (str): Optional user question or image-analysis request.

	Returns:
		str: Vision task instruction.
	"""
	task = str( st.session_state.get( 'vision_task', 'Extract Visible Text' ) )
	detail = str( st.session_state.get( 'vision_detail', 'Detailed' ) )
	response_format = str( st.session_state.get( 'vision_response_format', 'Plain Text' ) )
	response_language = str( st.session_state.get( 'vision_response_language', 'English' ) )
	preserve_layout = bool( st.session_state.get( 'vision_preserve_layout', True ) )
	include_visible_text = bool( st.session_state.get( 'vision_include_visible_text', True ) )
	task_map = {
		'Extract Visible Text': 'Extract all visible text accurately from the supplied image.',
		'Describe Image': 'Describe the supplied image accurately and identify important visual details.',
		'Answer Questions': 'Answer the user question using only information visible in the supplied image.',
		'Analyze Screenshot': 'Analyze the screenshot, including visible interface elements, text, and state.',
		'Analyze Chart': 'Analyze the chart, identify axes, series, values, trends, and material comparisons.',
		'Analyze Diagram': 'Analyze the diagram, identify components, labels, relationships, and flow.',
		'Extract Structured Data': 'Extract visible structured fields and values from the supplied image.',
		'Compare Images': 'Compare the supplied images and identify material similarities and differences.'
	}
	lines: List[ str ] = [ task_map.get( task, task_map[ 'Extract Visible Text' ] ),
		f'Detail Level: {detail}.', f'Response Format: {response_format}.',
		f'Response Language: {response_language}.' ]
	if preserve_layout:
		lines.append( 'Preserve meaningful visual ordering, sections, rows, columns, and line breaks.' )
	if include_visible_text and task != 'Extract Visible Text':
		lines.append( 'Include materially relevant visible text in the analysis.' )
	if user_input and str( user_input ).strip( ):
		lines.append( f'User Request: {str( user_input ).strip( )}' )
	return '\n'.join( lines ).strip( )


def run_vision_turn( image_payloads: List[ Dict[ str, Any ] ], user_input: str = '',
		stream: bool = False, output: Any = None, show_errors: bool = True,
		instruction_override: str = '', response_format_override: str = '' ) -> str:
	"""Executes one Gemma 3 Image-to-Text turn.

	Purpose:
		Sends local image bytes and bounded vision instructions through the MTMD multimodal chat path,
		using the same inference controls as text generation and converting failures into controlled UI
		feedback.

	Args:
		image_payloads (List[Dict[str, Any]]): Image dictionaries containing bytes and MIME metadata.
		user_input (str): Optional user-authored image question.
		stream (bool): True to stream response text.
		output (Any): Optional Streamlit output placeholder.
		show_errors (bool): True to display controlled runtime failures in the UI.
		instruction_override (str): Explicit vision instruction used by internal OCR workflows.
		response_format_override (str): Explicit response format used by internal OCR workflows.

	Returns:
		str: Generated Image-to-Text response, or an empty string when unavailable.
	"""
	if not image_payloads:
		return ''
	task = str( st.session_state.get( 'vision_task', 'Extract Visible Text' ) )
	if task == 'Compare Images' and len( image_payloads ) < 2:
		if show_errors:
			st.error( 'Compare Images requires at least two uploaded images.' )
		return ''
	if task == 'Answer Questions' and not str( user_input or '' ).strip( ) and not instruction_override:
		if show_errors:
			st.error( 'Answer Questions requires a non-empty Image Request.' )
		return ''
	try:
		runtime_llm = get_vision_runtime_llm( )
		instruction_text = (str( instruction_override ).strip( ) if instruction_override else
			build_vision_instruction( user_input ))
		content_parts: List[ Dict[ str, Any ] ] = [
			{ 'type': 'text', 'text': instruction_text }
		]
		for payload in image_payloads:
			image_bytes = payload.get( 'bytes', b'' )
			mime_type = str( payload.get( 'mime_type', 'image/png' ) or 'image/png' )
			if not image_bytes:
				continue
			content_parts.append( {
				'type': 'image_url',
				'image_url': { 'url': image_bytes_to_data_uri( image_bytes, mime_type ) }
			} )
		if len( content_parts ) == 1:
			return ''
		messages: List[ Dict[ str, Any ] ] = [ ]
		system_instructions = get_effective_system_instructions( )
		if system_instructions:
			messages.append( { 'role': 'system', 'content': system_instructions } )
		messages.append( { 'role': 'user', 'content': content_parts } )
		generation_args: Dict[ str, Any ] = {
			'messages': messages,
			'max_tokens': int( st.session_state.get( 'max_tokens', 1024 ) or 1024 ),
			'temperature': float( st.session_state.get( 'temperature', 0.0 ) ),
			'top_p': float( st.session_state.get( 'top_percent', 0.95 ) or 0.95 ),
			'repeat_penalty': float( st.session_state.get( 'repeat_penalty', 1.1 ) or 1.1 ),
			'stream': bool( stream )
		}
		top_k_value = int( st.session_state.get( 'top_k', 0 ) )
		seed_value = int( st.session_state.get( 'random_seed', 0 ) )
		if top_k_value > 0:
			generation_args[ 'top_k' ] = top_k_value
		if seed_value > 0:
			generation_args[ 'seed' ] = seed_value
		generation_args[ 'presence_penalty' ] = float( st.session_state.get( 'presense_penalty', 0.0 ) )
		generation_args[ 'frequency_penalty' ] = float( st.session_state.get( 'frequency_penalty', 0.0 ) )
		vision_format = str( response_format_override or
			st.session_state.get( 'vision_response_format', 'Plain Text' ) )
		if vision_format == 'JSON':
			generation_args[ 'response_format' ] = { 'type': 'json_object' }
		if not stream:
			response = runtime_llm.create_chat_completion( **generation_args )
			choices = response.get( 'choices', [ ] ) if isinstance( response, dict ) else [ ]
			if not choices:
				return ''
			return str( choices[ 0 ].get( 'message', { } ).get( 'content', '' ) or '' ).strip( )
		buffer = ''
		if output is None:
			output = st.empty( )
		for chunk in runtime_llm.create_chat_completion( **generation_args ):
			if not isinstance( chunk, dict ):
				continue
			choices = chunk.get( 'choices', [ ] )
			if not choices:
				continue
			delta = choices[ 0 ].get( 'delta', { } )
			piece = str( delta.get( 'content', '' ) or '' )
			if piece:
				buffer += piece
				output.markdown( buffer + '▌' )
		output.markdown( buffer )
		return buffer.strip( )
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'run_vision_turn'
		exception.method = ('run_vision_turn( image_payloads, user_input, stream, output, show_errors, '
			'instruction_override, response_format_override ) -> str')
		Logger( ).write( exception )
		if show_errors:
			message = f'Image-to-Text failed: {e}'
			if output is not None:
				output.error( message )
			else:
				st.error( message )
		return ''


def build_starter_prompt_template( category: str, task_type: str, response_format: str,
		language: str ) -> str:
	"""Builds a starter system prompt from prompt metadata before optional local-model drafting.

	Purpose:
		Builds a starter system prompt aligned to the persisted category taxonomy and bounded task,
		format, and language controls.

	Args:
		category (str): Prompt category.
		task_type (str): Task type.
		response_format (str): Response format.
		language (str): Preferred response language.

	Returns:
		str: Starter system prompt.
	"""
	category_value = str( category or '' ).strip( )
	task_value = str( task_type or 'Chat' ).strip( )
	format_value = str( response_format or 'Markdown' ).strip( )
	language_value = str( language or 'English' ).strip( )
	lines: List[ str ] = [
		f'You are Bro, a local AI assistant operating in the category "{category_value}".',
		f'Primary task type: {task_value}.', f'Response format: {format_value}.',
		f'Preferred language: {language_value}.'
	]
	category_guidance = {
		'Research / Academic': 'Provide careful, structured analysis grounded in supplied information.',
		'Software Engineering': 'Produce correct, editor-ready software guidance and code when requested.',
		'Writing / Administrative': 'Draft or revise content while preserving supplied facts, purpose, and audience.',
		'Data Analytics & Governance': 'Apply precise analytical and data-governance reasoning.',
		'Business / Finance / Marketing': 'Apply appropriate business, financial, or marketing analysis.',
		'Compliance / Legal / Budget': 'Apply precise compliance, legal, or budget analysis to supplied information.',
		'Prompt Engineering': 'Produce clear, reusable prompt instructions appropriate to the stated task.',
		'Instruction/ Training / Planning': 'Produce clear instructional, training, or planning guidance.',
		'Image Analysis': 'Produce instructions for image understanding and visible-text extraction.'
	}
	lines.append( category_guidance.get( category_value, 'Respond helpfully, accurately, and concisely.' ) )
	lines.append( 'If required information is missing, state that clearly.' )
	return '\n'.join( lines ).strip( )


def generate_prompt_template_draft( goal: str, constraints: str, style: str,
		category: str, task_type: str, response_format: str, language: str ) -> str:
	"""Uses the local model to draft a system prompt from bounded metadata controls.

	Purpose:
		Uses the shared Gemma chat-completion path to create an editable system-prompt draft from the
		selected category, task, response format, language, goal, constraints, and style.

	Args:
		goal (str): Prompt-generation goal.
		constraints (str): Prompt constraints.
		style (str): Generator style.
		category (str): Prompt category.
		task_type (str): Task type.
		response_format (str): Response format.
		language (str): Preferred language.

	Returns:
		str: Generated prompt-template draft.
	"""
	prompt = f'''Create a strong system prompt for the Bro local AI application.

Category: {category}
Task Type: {task_type}
Response Format: {response_format}
Language: {language}
Goal: {goal}
Constraints: {constraints}
Style: {style}

Write only the system prompt text. Do not add explanation.'''.strip( )
	return run_direct_llm_turn(
		system_instruction='Generate only the requested reusable system prompt. Do not use unrelated conversation context.',
		user_input=prompt,
		temperature=float( st.session_state.get( 'temperature', 0.2 ) ),
		top_p=float( st.session_state.get( 'top_percent', 0.95 ) ),
		repeat_penalty=float( st.session_state.get( 'repeat_penalty', 1.05 ) ),
		max_tokens=512,
		stream=False,
		output=None,
		response_format='Plain Text' )


def apply_prompt_to_text_generation( prompt_text: str ) -> None:
	"""Copies selected prompt text into shared system instructions for Text Generation mode.

	Purpose:
		Updates the shared system-instruction state consumed by Text Generation.

	Args:
		prompt_text (str): Prompt text to apply.

	Returns:
		None: This function performs its work through Streamlit session state.
	"""
	st.session_state[ 'system_instructions' ] = str( prompt_text or '' )


def apply_prompt_to_document_qna( prompt_text: str ) -> None:
	"""Copies selected prompt text into shared system instructions for Document Q&A.

	Purpose:
		Updates shared system instructions and enables the existing grounded Document Q&A defaults.

	Args:
		prompt_text (str): Prompt text to apply.

	Returns:
		None: This function performs its work through Streamlit session state.
	"""
	st.session_state[ 'system_instructions' ] = str( prompt_text or '' )
	st.session_state[ 'require_grounding' ] = True
	st.session_state[ 'answer_from_excerpts_only' ] = True


def apply_prompt_metadata_to_shared_state( category: str, task_type: str,
		response_format: str, language: str ) -> None:
	"""Applies selected prompt metadata to shared generation controls.

	Purpose:
		Synchronizes selected Prompt Engineering metadata with Text Generation without overloading
		database fields.

	Args:
		category (str): Prompt category.
		task_type (str): Task type.
		response_format (str): Response format.
		language (str): Preferred language.

	Returns:
		None: This function performs its work through Streamlit session state.
	"""
	st.session_state[ 'task_preset' ] = str( task_type or 'Chat' )
	st.session_state[ 'response_format' ] = str( response_format or 'Markdown' )
	st.session_state[ 'response_language' ] = str( language or 'English' )


def clone_prompt_record( source_prompt: Dict[ str, Any ] | None ) -> None:
	"""Copies a selected prompt record into the edit surface as a new prompt draft.

	Purpose:
		Clones mutable prompt fields while intentionally clearing the immutable database ID.

	Args:
		source_prompt (Dict[str, Any] | None): Source prompt record.

	Returns:
		None: This function performs its work through Streamlit session state.
	"""
	if not isinstance( source_prompt, dict ):
		return
	st.session_state.pe_selected_id = None
	st.session_state.pe_caption = f'{str( source_prompt.get( "Caption", "" ) )} Copy'.strip( )
	st.session_state.pe_name = str( source_prompt.get( 'Name', '' ) or '' )
	st.session_state.pe_category = str( source_prompt.get( 'Category', '' ) or '' )
	st.session_state.pe_text = str( source_prompt.get( 'Text', '' ) or '' )



def render_system_instructions_controls( workflow: str, include_preset: bool,
		include_preview: bool ) -> None:
	"""Renders mode-aware category and template System Instructions controls.

	Purpose:
		Preserves the existing System Instructions edit, conversion, clear, preset, and preview
		functionality while exposing only populated database categories that are appropriate for the
		active workflow. Category and template selector state is isolated by workflow.

	Args:
		workflow (str): Text Generation, Document Q&A, or Image to Text workflow name.
		include_preset (bool): True to render the existing Apply Preset button.
		include_preview (bool): True to render the existing Preview Prompt button and preview surface.

	Returns:
		None: This function renders Streamlit controls and updates session state.
	"""
	throw_if( 'workflow', workflow )
	workflow_key = str( workflow ).lower( ).replace( ' ', '_' ).replace( '&', 'and' )
	category_key = f'{workflow_key}_instruction_category'
	prompt_key = f'{workflow_key}_instruction_prompt_id'
	categories = fetch_mode_prompt_categories( cfg.DB_PATH, workflow )

	if category_key not in st.session_state:
		st.session_state[ category_key ] = categories[ 0 ] if categories else ''
	if prompt_key not in st.session_state:
		st.session_state[ prompt_key ] = None
	if st.session_state.get( category_key ) not in categories:
		st.session_state[ category_key ] = categories[ 0 ] if categories else ''
		st.session_state[ prompt_key ] = None

	def _on_category_change( ) -> None:
		st.session_state[ prompt_key ] = None
		st.session_state[ 'active_prompt_caption' ] = ''

	def _on_template_change( ) -> None:
		prompt_id = st.session_state.get( prompt_key )
		if not prompt_id:
			return
		prompt_row = fetch_prompt_by_id( int( prompt_id ) )
		if not prompt_row:
			return
		category = str( prompt_row.get( 'Category', '' ) or '' )
		if not is_prompt_category_allowed_for_workflow( category, workflow ):
			return
		st.session_state[ 'system_instructions' ] = str( prompt_row.get( 'Text', '' ) or '' )
		st.session_state[ 'active_prompt_caption' ] = str( prompt_row.get( 'Caption', '' ) or '' )
		st.session_state[ category_key ] = category

	def _on_clear( ) -> None:
		st.session_state[ 'system_instructions' ] = ''
		st.session_state[ prompt_key ] = None
		st.session_state[ 'active_prompt_caption' ] = ''

	def _on_convert_system_instructions( ) -> None:
		instruction_text = st.session_state.get( 'system_instructions', '' )
		if not isinstance( instruction_text, str ) or not instruction_text.strip( ):
			return
		source = instruction_text.strip( )
		if cfg.XML_BLOCK_PATTERN.search( source ):
			st.session_state[ 'system_instructions' ] = convert_xml( source )
		else:
			st.session_state[ 'system_instructions' ] = convert_markdown( source )

	def _on_apply_preset_template( ) -> None:
		task_preset = str( st.session_state.get( 'task_preset', 'Chat' ) or 'Chat' ).strip( )
		preset_map = {
			'Chat': 'You are Bro, a helpful local assistant. Be accurate, practical, and concise.',
			'Analysis': 'Analyze the supplied information carefully and provide a clear, supported conclusion.',
			'Reasoning': 'Solve the task carefully and return a clear, supported answer.',
			'Coding': 'Produce correct, editor-ready code and explain only as needed.',
			'Writing': 'Draft clear, accurate writing for the selected audience and tone.',
			'Editing': 'Revise the supplied writing while preserving facts and intended meaning.',
			'Translation': 'Translate faithfully while preserving meaning, terminology, and tone.',
			'Summarization': 'Summarize faithfully and preserve material facts.',
			'Extraction': 'Extract only supported facts and do not invent missing values.',
			'Classification': 'Classify from supplied evidence using the requested classification scheme.',
			'Comparison': 'Compare the supplied items consistently and identify material differences.',
			'Structured Output': 'Follow the requested output structure exactly.'
		}
		st.session_state[ 'system_instructions' ] = preset_map.get( task_preset, preset_map[ 'Chat' ] )

	in_left, in_right = st.columns( [ 0.8, 0.2 ] )
	with in_left:
		st.text_area( label='Enter Text', height=120, width='stretch',
			help=cfg.SYSTEM_INSTRUCTIONS, key='system_instructions' )
	with in_right:
		if categories:
			st.selectbox( label='Category', options=categories, key=category_key,
				on_change=_on_category_change )
			prompt_options = fetch_prompt_options(
				cfg.DB_PATH, str( st.session_state.get( category_key, '' ) or '' ) )
			prompt_ids = [ option[ 0 ] for option in prompt_options ]
			prompt_labels = { option[ 0 ]: option[ 1 ] for option in prompt_options }
			if st.session_state.get( prompt_key ) not in prompt_ids:
				st.session_state[ prompt_key ] = None
			st.selectbox( label='Use Template', options=prompt_ids, index=None,
				format_func=lambda prompt_id: prompt_labels.get( prompt_id, str( prompt_id ) ),
				key=prompt_key, on_change=_on_template_change, placeholder='Select Template' )
		else:
			st.selectbox( label='Category', options=[ 'No Compatible Templates Found' ],
				disabled=True, key=f'{workflow_key}_empty_category' )
			st.selectbox( label='Use Template', options=[ 'No Templates Found' ],
				disabled=True, key=f'{workflow_key}_empty_template' )

	button_count = 2 + int( include_preset ) + int( include_preview )
	button_columns = st.columns( [ 1.0 / button_count ] * button_count )
	button_index = 0
	with button_columns[ button_index ]:
		st.button( label='Clear Instructions', width='stretch', on_click=_on_clear, icon='🧹',
			key=f'{workflow_key}_clear_instructions' )
	button_index += 1
	with button_columns[ button_index ]:
		st.button( label='XML ↔️ Markdown', width='stretch',
			on_click=_on_convert_system_instructions, key=f'{workflow_key}_convert_instructions' )
	button_index += 1
	if include_preset:
		with button_columns[ button_index ]:
			st.button( label='Apply Preset', width='stretch', on_click=_on_apply_preset_template,
				key=f'{workflow_key}_apply_preset' )
		button_index += 1
	if include_preview:
		with button_columns[ button_index ]:
			if st.button( label='Preview Prompt', width='stretch',
					key=f'{workflow_key}_preview_prompt' ):
				st.session_state[ 'preview_effective_prompt' ] = not bool(
					st.session_state.get( 'preview_effective_prompt', False ) )
		if bool( st.session_state.get( 'preview_effective_prompt', False ) ):
			st.text_area( label='Effective Prompt Preview',
				value=build_effective_prompt_preview(
					str( st.session_state.get( 'last_preview_input', '' ) or '' ) ),
				height=220, disabled=True, key=f'{workflow_key}_effective_prompt_preview' )

# ----------- DATABASE UTILITIES -------------------------

def initialize_database( ) -> None:
	"""Creates required SQLite tables for chat history, embeddings, prompts, document metadata, chunks, and image metadata.

	Purpose:
		Creates required SQLite tables for chat history, embeddings, prompts, document
		metadata, chunks, and image metadata.
	"""
	Path( 'stores/sqlite' ).mkdir( parents=True, exist_ok=True )
	
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute( """
                      CREATE TABLE IF NOT EXISTS chat_history
                      (
                          id
                          INTEGER
                          PRIMARY
                          KEY
                          AUTOINCREMENT,
                          role
                          TEXT,
                          content
                          TEXT
                      )
		              """ )
		
		conn.execute( """
                      CREATE TABLE IF NOT EXISTS embeddings
                      (
                          id
                          INTEGER
                          PRIMARY
                          KEY
                          AUTOINCREMENT,
                          chunk
                          TEXT,
                          vector
                          BLOB
                      )
		              """ )
		
		conn.execute( """
                      CREATE TABLE IF NOT EXISTS Prompts
                      (
                          ID INTEGER NOT NULL UNIQUE,
                          Caption TEXT(80),
                          Name TEXT(80),
                          Category TEXT(80),
                          Text TEXT(2048),
                          PRIMARY KEY(ID AUTOINCREMENT)
                      )
		              """ )
		
		conn.execute( """
                      CREATE TABLE IF NOT EXISTS documents
                      (
                          DocumentId
                          INTEGER
                          PRIMARY
                          KEY
                          AUTOINCREMENT,
                          Name
                          TEXT
                          NOT
                          NULL,
                          Type
                          TEXT,
                          SizeBytes
                          INTEGER,
                          Source
                          TEXT,
                          Fingerprint
                          TEXT,
                          TextLength
                          INTEGER,
                          ChunkCount
                          INTEGER,
                          CreatedOn
                          TEXT
                      )
		              """ )
		
		conn.execute( """
                      CREATE TABLE IF NOT EXISTS document_chunks
                      (
                          ChunkId
                          INTEGER
                          PRIMARY
                          KEY
                          AUTOINCREMENT,
                          DocumentName
                          TEXT
                          NOT
                          NULL,
                          ChunkIndex
                          INTEGER,
                          ChunkText
                          TEXT,
                          ChunkLength
                          INTEGER,
                          Fingerprint
                          TEXT,
                          CreatedOn
                          TEXT
                      )
		              """ )
		
		conn.execute( """
                      CREATE TABLE IF NOT EXISTS document_embeddings
                      (
                          EmbeddingId
                          INTEGER
                          PRIMARY
                          KEY
                          AUTOINCREMENT,
                          DocumentName
                          TEXT
                          NOT
                          NULL,
                          ChunkIndex
                          INTEGER,
                          VectorDim
                          INTEGER,
                          Fingerprint
                          TEXT,
                          CreatedOn
                          TEXT
                      )
		              """ )
		
		conn.execute( """
                      CREATE TABLE IF NOT EXISTS images
                      (
                          ImageId
                          INTEGER
                          PRIMARY
                          KEY
                          AUTOINCREMENT,
                          Name
                          TEXT
                          NOT
                          NULL,
                          MimeType
                          TEXT,
                          SizeBytes
                          INTEGER,
                          Fingerprint
                          TEXT,
                          Source
                          TEXT,
                          CreatedOn
                          TEXT
                      )
		              """ )
		
		prompt_columns = [ row[ 1 ] for row in
		                   conn.execute( 'PRAGMA table_info("Prompts");' ).fetchall( ) ]
		required_prompt_columns = { 'ID', 'Caption', 'Name', 'Category', 'Text' }
		if not required_prompt_columns.issubset( set( prompt_columns ) ):
			raise RuntimeError(
				'The Prompts table schema is incompatible. Expected ID, Caption, Name, Category, and Text.' )

		conn.commit( )

def create_connection( ) -> sqlite3.Connection:
	"""Creates a SQLite connection to the configured application database path.

	Purpose:
		Creates a SQLite connection to the configured application database path.

	Returns:
		sqlite3.Connection: SQLite connection object.
	"""
	return sqlite3.connect( cfg.DB_PATH )

def list_tables( ) -> List[ str ]:
	"""Returns table names from the configured SQLite database for Data Management browsing and administration workflows.

	Purpose:
		Returns table names from the configured SQLite database for Data Management browsing
		and administration workflows.

	Returns:
		List[str]: Result produced by the operation.
	"""
	with create_connection( ) as conn:
		_query = "SELECT name FROM sqlite_master WHERE type='table' ORDER BY name;"
		rows = conn.execute( _query ).fetchall( )
		return [ r[ 0 ] for r in rows ]

def create_schema( table: str ) -> List[ Tuple ]:
	"""Returns SQLite schema metadata for a selected table.

	Purpose:
		Returns SQLite schema metadata for a selected table.

	Args:
		table: SQLite table name.

	Returns:
		List[Tuple]: Result produced by the operation.
	"""
	with create_connection( ) as conn:
		return conn.execute( f'PRAGMA table_info("{table}");' ).fetchall( )

def read_table( table: str, limit: int = None, offset: int = 0 ) -> pd.DataFrame:
	"""Reads rows from a selected SQLite table into a pandas DataFrame with optional paging controls.

	Purpose:
		Reads rows from a selected SQLite table into a pandas DataFrame with optional paging
		controls.

	Args:
		table: SQLite table name.
		limit: limit value used by this workflow.
		offset: offset value used by this workflow.

	Returns:
		pd.DataFrame: DataFrame produced by the operation.
	"""
	query = f'SELECT rowid, * FROM "{table}"'
	if limit:
		query += f" LIMIT {limit} OFFSET {offset}"
	with create_connection( ) as conn:
		return pd.read_sql_query( query, conn )

def drop_table( table: str ) -> None:
	"""Drops a selected SQLite table when requested by the Data Management administration workflow.

	Purpose:
		Drops a selected SQLite table when requested by the Data Management administration
		workflow.

	Args:
		table: SQLite table name.
	"""
	if not table:
		return
	
	with create_connection( ) as conn:
		conn.execute( f'DROP TABLE IF EXISTS "{table}";' )
		conn.commit( )

def rename_table( old_name: str, new_name: str ) -> None:
	"""Renames a SQLite table using native ALTER TABLE support or a schema-preserving rebuild fallback.

	Purpose:
		Renames a SQLite table using native ALTER TABLE support or a schema-preserving
		rebuild fallback.

	Args:
		old_name: Existing table or column name.
		new_name: Replacement table or column name.
	"""
	if not old_name or not new_name:
		return
	
	with create_connection( ) as conn:
		try:
			conn.execute( f'ALTER TABLE "{old_name}" RENAME TO "{new_name}";' )
			conn.commit( )
			return
		except Exception as e:
			exception = Error( e )
			exception.module = 'app'
			exception.cause = 'rename_table'
			exception.method = 'rename_table( old_name: str, new_name: str ) -> None'
			Logger( ).write( exception )
			pass
		
		row = conn.execute(
			"""
            SELECT sql
            FROM sqlite_master
            WHERE type ='table' AND name =?
			""",
			(old_name,)
		).fetchone( )
		
		if not row or not row[ 0 ]:
			raise ValueError( "Table definition not found." )
		
		create_sql = row[ 0 ]
		indexes = conn.execute(
			"""
            SELECT sql
            FROM sqlite_master
            WHERE type ='index' AND tbl_name=? AND sql IS NOT NULL
			""",
			(old_name,)
		).fetchall( )
		
		open_paren = create_sql.find( "(" )
		if open_paren == -1:
			raise ValueError( "Malformed CREATE TABLE statement." )
		
		temp_name = f"{new_name}__rebuild_temp"
		conn.execute( "BEGIN" )
		conn.execute( f'CREATE TABLE "{temp_name}" {create_sql[ open_paren: ]}' )
		cols = [ r[ 1 ] for r in conn.execute( f'PRAGMA table_info("{old_name}");' ).fetchall( ) ]
		col_list = ", ".join( [ f'"{c}"' for c in cols ] )
		
		conn.execute(
			f'INSERT INTO "{temp_name}" ({col_list}) SELECT {col_list} FROM "{old_name}";'
		)
		
		conn.execute( f'DROP TABLE "{old_name}";' )
		conn.execute( f'ALTER TABLE "{temp_name}" RENAME TO "{new_name}";' )
		
		for idx in indexes:
			idx_sql = idx[ 0 ]
			if idx_sql:
				idx_sql = idx_sql.replace( f'ON "{old_name}"', f'ON "{new_name}"' )
				conn.execute( idx_sql )
		
		conn.commit( )

def rename_column( table_name: str, old_name: str, new_name: str ) -> None:
	"""Renames a SQLite table column using native ALTER TABLE support or a schema-preserving rebuild fallback.

	Purpose:
		Renames a SQLite table column using native ALTER TABLE support or a schema-
		preserving rebuild fallback.

	Args:
		table_name: SQLite table name.
		old_name: Existing table or column name.
		new_name: Replacement table or column name.
	"""
	if not table_name or not old_name or not new_name:
		return
	
	with create_connection( ) as conn:
		try:
			conn.execute(
				f'ALTER TABLE "{table_name}" RENAME COLUMN "{old_name}" TO "{new_name}";'
			)
			conn.commit( )
			return
		except Exception as e:
			exception = Error( e )
			exception.module = 'app'
			exception.cause = 'rename_column'
			exception.method = 'rename_column( table_name: str, old_name: str, new_name: str ) -> None'
			Logger( ).write( exception )
			pass
		
		row = conn.execute( """
                            SELECT sql
                            FROM sqlite_master
                            WHERE type ='table' AND name =?
		                    """, (table_name,) ).fetchone( )
		
		if not row or not row[ 0 ]:
			raise ValueError( "Table definition not found." )
		
		create_sql = row[ 0 ]
		indexes = conn.execute( """
                                SELECT sql
                                FROM sqlite_master
                                WHERE type ='index' AND tbl_name=? AND sql IS NOT NULL
		                        """, (table_name,) ).fetchall( )
		
		schema = conn.execute( f'PRAGMA table_info("{table_name}");' ).fetchall( )
		cols = [ r[ 1 ] for r in schema ]
		if old_name not in cols:
			raise ValueError( "Column not found." )
		
		mapped_cols = [ (new_name if c == old_name else c) for c in cols ]
		temp_table = f"{table_name}__rebuild_temp"
		col_defs: List[ str ] = [ ]
		pk_cols = [ r for r in schema if int( r[ 5 ] or 0 ) > 0 ]
		single_pk = len( pk_cols ) == 1
		
		for row in schema:
			col_name = row[ 1 ]
			col_type = row[ 2 ] or ''
			not_null = int( row[ 3 ] or 0 )
			default_value = row[ 4 ]
			pk = int( row[ 5 ] or 0 )
			
			out_name = new_name if col_name == old_name else col_name
			col_def = f'"{out_name}" {col_type}'.strip( )
			
			if not_null:
				col_def += ' NOT NULL'
			
			if default_value is not None:
				col_def += f' DEFAULT {default_value}'
			
			if single_pk and pk == 1:
				col_def += ' PRIMARY KEY'
			
			col_defs.append( col_def )
		
		new_create_sql = f'CREATE TABLE "{temp_table}" ({", ".join( col_defs )});'
		
		old_select = ", ".join( [ f'"{c}"' for c in cols ] )
		new_insert = ", ".join( [ f'"{c}"' for c in mapped_cols ] )
		
		conn.execute( "BEGIN" )
		conn.execute( new_create_sql )
		conn.execute(
			f'INSERT INTO "{temp_table}" ({new_insert}) SELECT {old_select} FROM "{table_name}";'
		)
		
		conn.execute( f'DROP TABLE "{table_name}";' )
		conn.execute( f'ALTER TABLE "{temp_table}" RENAME TO "{table_name}";' )
		
		for idx in indexes:
			idx_sql = idx[ 0 ]
			if idx_sql:
				idx_sql = idx_sql.replace( f'"{old_name}"', f'"{new_name}"' )
				conn.execute( idx_sql )
		
		conn.commit( )

def create_index( table: str, column: str ) -> None:
	"""Creates a safe SQLite index for a validated table and column.

	Purpose:
		Creates a safe SQLite index for a validated table and column.

	Args:
		table: SQLite table name.
		column: SQLite column name.
	"""
	if not table or not column:
		return
	
	# ----------  Validate table exists
	tables = list_tables( )
	if table not in tables:
		raise ValueError( "Invalid table name." )
	
	# ----------  Validate column exists
	schema = create_schema( table )
	valid_columns = [ col[ 1 ] for col in schema ]
	
	if column not in valid_columns:
		raise ValueError( "Invalid column name." )
	
	# ----------  Sanitize index name (identifier only)
	safe_index_name = re.sub( r"[^0-9a-zA-Z_]+", "_", f"idx_{table}_{column}" )
	
	# ----------  Create index safely (quote identifiers)
	sql = f'CREATE INDEX IF NOT EXISTS "{safe_index_name}" ON "{table}"("{column}");'
	
	with create_connection( ) as conn:
		conn.execute( sql )
		conn.commit( )

def apply_filters( df: pd.DataFrame ) -> pd.DataFrame:
	"""Renders advanced Streamlit filter controls and applies the selected filter to a DataFrame.

	Purpose:
		Renders advanced Streamlit filter controls and applies the selected filter to a
		DataFrame.

	Args:
		df: DataFrame to process.

	Returns:
		pd.DataFrame: DataFrame produced by the operation.
	"""
	st.subheader( 'Advanced Filters' )
	col1, col2, col3 = st.columns( 3 )
	column = col1.selectbox( 'Column', df.columns )
	operator = col2.selectbox( 'Operator', [ '=', '!=', '>', '<', '>=', '<=', 'contains' ] )
	value = col3.text_input( 'Value' )
	if value:
		if operator == '=':
			df = df[ df[ column ] == value ]
		elif operator == '!=':
			df = df[ df[ column ] != value ]
		elif operator == '>':
			df = df[ df[ column ].astype( float ) > float( value ) ]
		elif operator == '<':
			df = df[ df[ column ].astype( float ) < float( value ) ]
		elif operator == '>=':
			df = df[ df[ column ].astype( float ) >= float( value ) ]
		elif operator == '<=':
			df = df[ df[ column ].astype( float ) <= float( value ) ]
		elif operator == 'contains':
			df = df[ df[ column ].astype( str ).str.contains( value ) ]
	
	return df

def create_aggregation( df: pd.DataFrame ):
	"""Renders Streamlit aggregation controls and displays an aggregate metric for a selected numeric column.

	Purpose:
		Renders Streamlit aggregation controls and displays an aggregate metric for a
		selected numeric column.

	Args:
		df: DataFrame to process.
	"""
	st.subheader( 'Aggregation Engine' )
	numeric_cols = df.select_dtypes( include=[ 'number' ] ).columns.tolist( )
	if not numeric_cols:
		st.info( 'No numeric columns available.' )
		return
	
	col = st.selectbox( 'Column', numeric_cols )
	agg = st.selectbox( 'Aggregation', [ 'COUNT', 'SUM', 'AVG', 'MIN', 'MAX', 'MEDIAN' ] )
	if agg == 'COUNT':
		result = df[ col ].count( )
	elif agg == 'SUM':
		result = df[ col ].sum( )
	elif agg == 'AVG':
		result = df[ col ].mean( )
	elif agg == 'MIN':
		result = df[ col ].min( )
	elif agg == 'MAX':
		result = df[ col ].max( )
	elif agg == 'MEDIAN':
		result = df[ col ].median( )
	
	st.metric( 'Result', result )

def create_visualization( df: pd.DataFrame ):
	"""Renders Streamlit chart controls and displays Plotly charts for selected DataFrame columns.

	Purpose:
		Renders Streamlit chart controls and displays Plotly charts for selected DataFrame
		columns.

	Args:
		df: DataFrame to process.
	"""
	st.subheader( 'Visualization Engine' )
	numeric_cols = df.select_dtypes( include=[ 'number' ] ).columns.tolist( )
	categorical_cols = df.select_dtypes( include=[ 'object' ] ).columns.tolist( )
	chart = st.selectbox( 'Chart Type',
		[ 'Histogram', 'Bar', 'Line', 'Scatter', 'Box', 'Pie', 'Correlation' ] )
	
	if chart == 'Histogram' and numeric_cols:
		col = st.selectbox( 'Column', numeric_cols )
		fig = px.histogram( df, x=col )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Bar':
		x = st.selectbox( 'X', df.columns )
		y = st.selectbox( 'Y', numeric_cols )
		fig = px.bar( df, x=x, y=y )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Line':
		x = st.selectbox( 'X', df.columns )
		y = st.selectbox( 'Y', numeric_cols )
		fig = px.line( df, x=x, y=y )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Scatter':
		x = st.selectbox( 'X', numeric_cols )
		y = st.selectbox( 'Y', numeric_cols )
		fig = px.scatter( df, x=x, y=y )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Box':
		col = st.selectbox( 'Column', numeric_cols )
		fig = px.box( df, y=col )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Pie':
		col = st.selectbox( 'Category Column', categorical_cols )
		fig = px.pie( df, names=col )
		st.plotly_chart( fig, use_container_width=True )
	
	elif chart == 'Correlation' and len( numeric_cols ) > 1:
		corr = df[ numeric_cols ].corr( )
		fig = px.imshow( corr, text_auto=True )
		st.plotly_chart( fig, use_container_width=True )

def convert_dataframe( table_name: str, df: pd.DataFrame ):
	"""Creates a SQLite table definition from a DataFrame column layout and inferred SQLite column types.

	Purpose:
		Creates a SQLite table definition from a DataFrame column layout and inferred SQLite
		column types.

	Args:
		table_name: SQLite table name.
		df: DataFrame to process.
	"""
	columns = [ ]
	for col in df.columns:
		sql_type = get_sqlite_type( df[ col ].dtype )
		safe_col = col.replace( ' ', '_' )
		columns.append( f'{safe_col} {sql_type}' )
	
	create_stmt = f'CREATE TABLE IF NOT EXISTS {table_name} ({", ".join( columns )});'
	
	with create_connection( ) as conn:
		conn.execute( create_stmt )
		conn.commit( )

def insert_data( table_name: str, df: pd.DataFrame ):
	"""Inserts DataFrame rows into a SQLite table after normalizing column names for SQLite compatibility.

	Purpose:
		Inserts DataFrame rows into a SQLite table after normalizing column names for SQLite
		compatibility.

	Args:
		table_name: SQLite table name.
		df: DataFrame to process.
	"""
	df = df.copy( )
	df.columns = [ c.replace( ' ', '_' ) for c in df.columns ]
	
	placeholders = ', '.join( [ '?' ] * len( df.columns ) )
	stmt = f'INSERT INTO {table_name} VALUES ({placeholders});'
	
	with create_connection( ) as conn:
		conn.executemany( stmt, df.values.tolist( ) )
		conn.commit( )

def get_sqlite_type( dtype: str ) -> str:
	"""Maps a pandas dtype to the SQLite column type used by table-creation workflows.

	Purpose:
		Maps a pandas dtype to the SQLite column type used by table-creation workflows.

	Args:
		dtype: str Pandas dtype to map.

	Returns:
		str: Text produced by the operation.
	"""
	dtype_str = str( dtype ).lower( )
	
	# ----------  Integer Types
	if "int" in dtype_str:
		return "INTEGER"
	
	# ----------  Float Types
	if "float" in dtype_str:
		return "REAL"
	
	# ----------  Boolean
	if "bool" in dtype_str:
		return "INTEGER"
	
	# ----------  Datetime
	if "datetime" in dtype_str:
		return "TEXT"
	
	# ----------  Categorical
	if "category" in dtype_str:
		return "TEXT"
	
	# ----------  Default fallback
	return "TEXT"

def create_custom_table( table_name: str, columns: list ) -> None:
	"""Creates a custom SQLite table from validated column-definition metadata.

	Purpose:
		Creates a custom SQLite table from validated column-definition metadata.

	Args:
		table_name: SQLite table name.
		columns: columns value used by this workflow.
	"""
	if not table_name:
		raise ValueError( "Table name required." )
	
	# ----------  Validate identifier
	if not re.match( r"^[A-Za-z_][A-Za-z0-9_]*$", table_name ):
		raise ValueError( "Invalid table name." )
	
	col_defs = [ ]
	for col in columns:
		col_name = col[ "name" ]
		col_type = col[ "type" ].upper( )
		if not re.match( r"^[A-Za-z_][A-Za-z0-9_]*$", col_name ):
			raise ValueError( f"Invalid column name: {col_name}" )
		
		definition = f'"{col_name}" {col_type}'
		if col[ "primary_key" ]:
			definition += " PRIMARY KEY"
			if col[ "auto_increment" ] and col_type == "INTEGER":
				definition += " AUTOINCREMENT"
		
		if col[ "not_null" ]:
			definition += " NOT NULL"
		
		col_defs.append( definition )
	
	sql = f'CREATE TABLE IF NOT EXISTS "{table_name}" ({", ".join( col_defs )});'
	with create_connection( ) as conn:
		conn.execute( sql )
		conn.commit( )

def is_safe_query( query: str ) -> bool:
	"""Determines whether a SQL query is read-only and safe for the guarded SQL console.

	Purpose:
		Determines whether a SQL query is read-only and safe for the guarded SQL console.

	Args:
		query: Query text.

	Returns:
		bool: Boolean result produced by the operation.
	"""
	if not query or not isinstance( query, str ):
		return False
	
	q = query.strip( ).lower( )
	
	# ----------  Block multiple statements
	if ';' in q[ :-1 ]:
		return False
	
	# ----------  Remove SQL comments
	q = re.sub( r"--.*?$", "", q, flags=re.MULTILINE )
	q = re.sub( r"/\*.*?\*/", "", q, flags=re.DOTALL )
	q = q.strip( )
	
	# ----------  Allowed starting keywords
	allowed_starts = ('select', 'with', 'explain', 'pragma')
	if not q.startswith( allowed_starts ):
		return False
	
	# ----------  Block dangerous keywords anywhere
	blocked_keywords = ('insert ', 'update ', 'delete ', 'drop ', 'alter ',
	                    'create ', 'attach ', 'detach ', 'vacuum ', 'replace ', 'trigger ')
	
	for keyword in blocked_keywords:
		if keyword in q:
			return False
	
	return True

def create_identifier( name: str ) -> str:
	"""Sanitizes arbitrary text into a safe SQLite identifier for schema administration workflows.

	Purpose:
		Sanitizes arbitrary text into a safe SQLite identifier for schema administration
		workflows.

	Args:
		name: Prompt caption or environment variable name.

	Returns:
		str: Text produced by the operation.
	"""
	if not name or not isinstance( name, str ):
		raise ValueError( 'Invalid Identifier.' )
	
	safe = re.sub( r'[^0-9a-zA-Z_]', '_', name.strip( ) )
	if not re.match( r'^[A-Za-z_]', safe ):
		safe = f'_{safe}'
	
	if not safe:
		raise ValueError( 'Invalid identifier after sanitization.' )
	
	return safe

def get_indexes( table: str ):
	"""Returns SQLite index metadata for a selected table.

	Purpose:
		Returns SQLite index metadata for a selected table.

	Args:
		table: SQLite table name.
	"""
	with create_connection( ) as conn:
		rows = conn.execute( f'PRAGMA index_list("{table}");' ).fetchall( )
		return rows

def add_column( table: str, column: str, col_type: str ):
	"""Adds a sanitized column to a selected SQLite table with the requested SQLite type.

	Purpose:
		Adds a sanitized column to a selected SQLite table with the requested SQLite type.

	Args:
		table: SQLite table name.
		column: SQLite column name.
		col_type: col type value used by this workflow.
	"""
	column = create_identifier( column )
	col_type = col_type.upper( )
	
	with create_connection( ) as conn:
		conn.execute(
			f'ALTER TABLE "{table}" ADD COLUMN "{column}" {col_type};' )
		conn.commit( )

def create_profile_table( table: str ):
	"""Builds a profile DataFrame summarizing nulls, distinct values, and numeric ranges for a selected table.

	Purpose:
		Builds a profile DataFrame summarizing nulls, distinct values, and numeric ranges
		for a selected table.

	Args:
		table: SQLite table name.
	"""
	df = read_table( table )
	profile_rows = [ ]
	total_rows = len( df )
	for col in df.columns:
		series = df[ col ]
		null_count = series.isna( ).sum( )
		distinct_count = series.nunique( dropna=True )
		row = \
			{
				'column': col, 'dtype': str( series.dtype ),
				'null_%': round( (null_count / total_rows) * 100, 2 ) if total_rows else 0,
				'distinct_%': round( ( distinct_count / total_rows) * 100, 2 ) if total_rows else 0,
			}
		
		if pd.api.types.is_numeric_dtype( series ):
			row[ "min" ] = series.min( )
			row[ "max" ] = series.max( )
			row[ "mean" ] = series.mean( )
		else:
			row[ "min" ] = None
			row[ "max" ] = None
			row[ "mean" ] = None
		
		profile_rows.append( row )
	
	return pd.DataFrame( profile_rows )

def drop_column( table: str, column: str ):
	"""Drops a column by rebuilding the SQLite table while preserving remaining columns, data, and usable indexes.

	Purpose:
		Drops a column by rebuilding the SQLite table while preserving remaining columns,
		data, and usable indexes.

	Args:
		table: SQLite table name.
		column: SQLite column name.
	"""
	if not table or not column:
		raise ValueError( "Table and column required." )
	
	with create_connection( ) as conn:
		schema = conn.execute( f'PRAGMA table_info("{table}");' ).fetchall( )
		if not schema:
			raise ValueError( "Table definition not found." )
		
		col_names = [ r[ 1 ] for r in schema ]
		if column not in col_names:
			raise ValueError( "Column not found." )
		
		remaining = [ r for r in schema if r[ 1 ] != column ]
		if not remaining:
			raise ValueError( "Cannot drop the only remaining column." )
		
		temp_table = f"{table}_rebuild_temp"
		
		pk_cols = [ r for r in remaining if int( r[ 5 ] or 0 ) > 0 ]
		single_pk = len( pk_cols ) == 1
		
		new_defs: List[ str ] = [ ]
		for row in remaining:
			col_name = row[ 1 ]
			col_type = row[ 2 ] or ''
			not_null = int( row[ 3 ] or 0 )
			default_value = row[ 4 ]
			pk = int( row[ 5 ] or 0 )
			
			col_def = f'"{col_name}" {col_type}'.strip( )
			
			if not_null:
				col_def += ' NOT NULL'
			
			if default_value is not None:
				col_def += f' DEFAULT {default_value}'
			
			if single_pk and pk == 1:
				col_def += ' PRIMARY KEY'
			
			new_defs.append( col_def )
		
		new_create_sql = f'CREATE TABLE "{temp_table}" ({", ".join( new_defs )});'
		
		conn.execute( "BEGIN" )
		conn.execute( new_create_sql )
		
		remaining_cols = [ r[ 1 ] for r in remaining ]
		col_list = ", ".join( [ f'"{c}"' for c in remaining_cols ] )
		
		conn.execute( f'INSERT INTO "{temp_table}" ({col_list}) '
		              f'SELECT {col_list} FROM "{table}";' )
		
		indexes = conn.execute(
			"""
            SELECT sql
            FROM sqlite_master
            WHERE type ='index' AND tbl_name=? AND sql IS NOT NULL
			""",
			(table,)
		).fetchall( )
		
		conn.execute( f'DROP TABLE "{table}";' )
		conn.execute( f'ALTER TABLE "{temp_table}" RENAME TO "{table}";' )
		
		for idx in indexes:
			idx_sql = idx[ 0 ]
			if idx_sql and column not in idx_sql:
				conn.execute( idx_sql )
		
		conn.commit( )

def reset_selection( ) -> None:
	"""Clears the Prompt Engineering edit surface and selected prompt identifier in Streamlit session state.

	Purpose:
		Clears the authoritative prompt edit fields without referencing removed legacy schema columns.

	Returns:
		None: This function performs its work through Streamlit session state.
	"""
	st.session_state.pe_selected_id = None
	st.session_state.pe_caption = ''
	st.session_state.pe_name = ''
	st.session_state.pe_category = ''
	st.session_state.pe_text = ''


def load_prompt( pid: int ) -> None:
	"""Loads a prompt record into the Prompt Engineering edit surface by primary key.

	Purpose:
		Loads ID, Caption, Name, Category, and Text from the authoritative Prompts schema.

	Args:
		pid (int): Prompt primary key.

	Returns:
		None: This function performs its work through Streamlit session state.
	"""
	prompt_row = fetch_prompt_by_id( pid )
	if not prompt_row:
		return
	st.session_state.pe_selected_id = int( prompt_row[ 'ID' ] )
	st.session_state.pe_caption = str( prompt_row.get( 'Caption', '' ) or '' )
	st.session_state.pe_name = str( prompt_row.get( 'Name', '' ) or '' )
	st.session_state.pe_category = str(
		prompt_row.get( 'Category', '' ) or '' )
	st.session_state.pe_text = str( prompt_row.get( 'Text', '' ) or '' )


def get_ai_asset_tables( ) -> List[ str ]:
	"""Returns the SQLite table names used for AI asset governance metadata.

	Purpose:
		Returns the SQLite table names used for AI asset governance metadata.

	Returns:
		List[str]: Result produced by the operation.
	"""
	return [ 'documents', 'document_chunks', 'document_embeddings', 'images' ]

def get_timestamp_text( ) -> str:
	"""Returns a timestamp string used for metadata rows written to local SQLite tables.

	Purpose:
		Returns a timestamp string used for metadata rows written to local SQLite tables.

	Returns:
		str: Text produced by the operation.
	"""
	return time.strftime( '%Y-%m-%d %H:%M:%S' )

def register_session_documents( ) -> Dict[ str, int ]:
	"""Registers active uploaded documents into the governed documents table with size, type, fingerprint, and chunk metadata.

	Purpose:
		Registers active uploaded documents into the governed documents table with size,
		type, fingerprint, and chunk metadata.

	Returns:
		Dict[str, int]: Result produced by the operation.
	"""
	active_docs = st.session_state.get( 'active_docs', [ ] )
	doc_bytes = st.session_state.get( 'doc_bytes', { } )
	
	inserted = 0
	updated = 0
	
	with create_connection( ) as conn:
		for name in active_docs:
			file_bytes = doc_bytes.get( name, b'' )
			if not file_bytes:
				continue
			
			text = extract_text( file_bytes, name )
			chunks = chunk_text( text ) if text else [ ]
			fingerprint = hashlib.sha256( file_bytes ).hexdigest( )
			file_type = Path( name ).suffix.lower( ).replace( '.', '' )
			created_on = get_timestamp_text( )
			
			existing = conn.execute(
				'''
                SELECT DocumentId
                FROM documents
                WHERE Name = ?
                  AND Fingerprint = ?
				''',
				(name, fingerprint)
			).fetchone( )
			
			if existing:
				conn.execute(
					'''
                    UPDATE documents
                    SET Type       = ?,
                        SizeBytes  = ?,
                        Source     = ?,
                        TextLength = ?,
                        ChunkCount = ?,
                        CreatedOn  = ?
                    WHERE DocumentId = ?
					''',
					(
							file_type,
							len( file_bytes ),
							'uploadlocal',
							len( text ),
							len( chunks ),
							created_on,
							existing[ 0 ]
					)
				)
				updated += 1
			else:
				conn.execute(
					'''
                    INSERT INTO documents
                    (Name,
                     Type,
                     SizeBytes,
                     Source,
                     Fingerprint,
                     TextLength,
                     ChunkCount,
                     CreatedOn)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?)
					''',
					(
							name,
							file_type,
							len( file_bytes ),
							'uploadlocal',
							fingerprint,
							len( text ),
							len( chunks ),
							created_on
					)
				)
				inserted += 1
		
		conn.commit( )
	
	return { 'inserted': inserted, 'updated': updated }

def register_session_chunks( ) -> Dict[ str, int ]:
	"""Registers active document chunks into the governed document_chunks table for local asset traceability.

	Purpose:
		Registers active document chunks into the governed document_chunks table for local
		asset traceability.

	Returns:
		Dict[str, int]: Result produced by the operation.
	"""
	active_docs = st.session_state.get( 'active_docs', [ ] )
	doc_bytes = st.session_state.get( 'doc_bytes', { } )
	inserted = 0
	
	with create_connection( ) as conn:
		for name in active_docs:
			file_bytes = doc_bytes.get( name, b'' )
			if not file_bytes:
				continue
			
			text = extract_text( file_bytes, name )
			chunks = chunk_text( text ) if text else [ ]
			file_fingerprint = hashlib.sha256( file_bytes ).hexdigest( )
			created_on = get_timestamp_text( )
			
			conn.execute(
				'DELETE FROM document_chunks WHERE DocumentName = ? AND Fingerprint = ?',
				(name, file_fingerprint)
			)
			
			for idx, chunk_value in enumerate( chunks ):
				conn.execute(
					'''
                    INSERT INTO document_chunks
                    (DocumentName,
                     ChunkIndex,
                     ChunkText,
                     ChunkLength,
                     Fingerprint,
                     CreatedOn)
                    VALUES (?, ?, ?, ?, ?, ?)
					''',
					(
							name,
							idx,
							chunk_value,
							len( chunk_value ),
							file_fingerprint,
							created_on
					)
				)
				inserted += 1
		
		conn.commit( )
	
	return { 'inserted': inserted }

def register_session_embeddings( ) -> Dict[ str, int ]:
	"""Registers document embedding metadata into the governed document_embeddings table when an embedder is available.

	Purpose:
		Registers document embedding metadata into the governed document_embeddings table
		when an embedder is available.

	Returns:
		Dict[str, int]: Result produced by the operation.
	"""
	active_docs = st.session_state.get( 'active_docs', [ ] )
	doc_bytes = st.session_state.get( 'doc_bytes', { } )
	inserted = 0
	
	if embedder is None:
		return { 'inserted': 0 }
	
	vector_dim = getattr( embedder, 'get_sentence_embedding_dimension', lambda: 384 )( )
	vector_dim = int( vector_dim ) if vector_dim else 384
	with create_connection( ) as conn:
		for name in active_docs:
			file_bytes = doc_bytes.get( name, b'' )
			if not file_bytes:
				continue
			
			text = extract_text( file_bytes, name )
			chunks = chunk_text( text ) if text else [ ]
			file_fingerprint = hashlib.sha256( file_bytes ).hexdigest( )
			created_on = get_timestamp_text( )
			
			conn.execute(
				'DELETE FROM document_embeddings WHERE DocumentName = ? AND Fingerprint = ?',
				(name, file_fingerprint) )
			
			for idx, _chunk_value in enumerate( chunks ):
				conn.execute( '''
                              INSERT INTO document_embeddings
                              (DocumentName,
                               ChunkIndex,
                               VectorDim,
                               Fingerprint,
                               CreatedOn)
                              VALUES (?, ?, ?, ?, ?)
				              ''', (name, idx, vector_dim, file_fingerprint, created_on) )
				inserted += 1
		
		conn.commit( )
	
	return { 'inserted': inserted }

def register_upload_images( uploaded_files: List[ Any ] ) -> Dict[ str, int ]:
	"""Registers uploaded image metadata into the governed images table with MIME type, size, fingerprint, and source metadata.

	Purpose:
		Registers uploaded image metadata into the governed images table with MIME type,
		size, fingerprint, and source metadata.

	Args:
		uploaded_files: Uploaded Streamlit file objects.

	Returns:
		Dict[str, int]: Result produced by the operation.
	"""
	inserted = 0
	updated = 0
	
	with create_connection( ) as conn:
		for f in uploaded_files:
			try:
				name = str( getattr( f, 'name', '' ) or '' ).strip( )
				file_bytes = f.getvalue( )
				mime_type = str( getattr( f, 'type', '' ) or '' ).strip( )
			except Exception as e:
				exception = Error( e )
				exception.module = 'app'
				exception.cause = 'register_upload_images'
				exception.method = 'register_upload_images( uploaded_files: List[Any] ) -> Dict[str, int]'
				Logger( ).write( exception )
				continue
			
			if not name or not file_bytes:
				continue
			
			fingerprint = hashlib.sha256( file_bytes ).hexdigest( )
			created_on = get_timestamp_text( )
			
			existing = conn.execute(
				'''
                SELECT ImageId
                FROM images
                WHERE Name = ?
                  AND Fingerprint = ?
				''',
				(name, fingerprint)
			).fetchone( )
			
			if existing:
				conn.execute(
					'''
                    UPDATE images
                    SET MimeType  = ?,
                        SizeBytes = ?,
                        Source    = ?,
                        CreatedOn = ?
                    WHERE ImageId = ?
					''',
					(
							mime_type,
							len( file_bytes ),
							'uploadlocal',
							created_on,
							existing[ 0 ]
					)
				)
				updated += 1
			else:
				conn.execute(
					'''
                    INSERT INTO images
                    (Name,
                     MimeType,
                     SizeBytes,
                     Fingerprint,
                     Source,
                     CreatedOn)
                    VALUES (?, ?, ?, ?, ?, ?)
					''',
					(
							name,
							mime_type,
							len( file_bytes ),
							fingerprint,
							'uploadlocal',
							created_on
					)
				)
				inserted += 1
		
		conn.commit( )
	
	return { 'inserted': inserted, 'updated': updated }

# -------------- LLM  UTILITIES -------------------

@st.cache_resource
def load_llm( ctx: int, threads: int, repeat_window: int, batch_size: int,
		micro_batch_size: int ) -> Any | None:
	"""Loads the configured local GGUF model through llama.cpp.

	Purpose:
		Loads the configured Gemma 3 GGUF model through cached Streamlit resource management using the
		selected context window, CPU thread count, and repeat-history window.

	Args:
		ctx (int): Context-window size.
		threads (int): CPU thread count.
		repeat_window (int): Number of recent tokens retained for repetition penalties.
		batch_size (int): llama.cpp logical batch size.
		micro_batch_size (int): llama.cpp physical micro-batch size.

	Returns:
		Any | None: Loaded llama.cpp runtime when available; otherwise None.
	"""
	try:
		if not local_model_available( ):
			return None
		from llama_cpp import Llama
		ctx_value = int( ctx ) if int( ctx ) > 0 else int( cfg.DEFAULT_CTX )
		thread_value = int( threads ) if int( threads ) > 0 else int( cfg.CORES )
		repeat_window_value = int( repeat_window ) if int( repeat_window ) > 0 else 64
		batch_size_value = int( batch_size ) if int( batch_size ) > 0 else 512
		micro_batch_size_value = int( micro_batch_size ) if int( micro_batch_size ) > 0 else 128
		micro_batch_size_value = min( micro_batch_size_value, batch_size_value )
		return Llama( model_path=str( cfg.MODEL_PATH ), n_ctx=ctx_value, n_threads=thread_value,
			n_batch=batch_size_value, n_ubatch=micro_batch_size_value,
			last_n_tokens_size=repeat_window_value, verbose=False )
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'load_llm'
		exception.method = ('load_llm( ctx: int, threads: int, repeat_window: int, batch_size: int, '
			'micro_batch_size: int ) -> Any | None')
		Logger( ).write( exception )
		return None

@st.cache_resource
def load_embedder( ) -> Any | None:
	"""Loads the sentence-transformer embedding model through cached Streamlit resource management when the dependency is available.

	Purpose:
		Loads the sentence-transformer embedding model through cached Streamlit resource
		management when the dependency is available.

	Returns:
		Any | None: Runtime object when available; otherwise None.
	"""
	try:
		from sentence_transformers import SentenceTransformer
		
		return SentenceTransformer( 'all-MiniLM-L6-v2' )
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'load_embedder'
		exception.method = 'load_embedder(  ) -> Any | None'
		Logger( ).write( exception )
		return None

# ------------- DOCQNA UTILITIES ----------------------

def create_docqna_instruction( action_name: str ) -> str:
	"""Returns action-specific Document Q&A guidance for the selected document workflow.

	Purpose:
		Returns action-specific Document Q&A guidance for the selected document workflow.

	Args:
		action_name: Selected Document Q&A action.

	Returns:
		str: Text produced by the operation.
	"""
	action = str( action_name or 'Answer Question' ).strip( )
	action_map = {
			'Answer Question':
				'Answer the user question directly using the retrieved excerpts.',
			'Summarize Active Document':
				'Provide a clear, structured summary of the active document.',
			'Extract Key Points':
				'Extract the most important points as a concise bullet list.',
			'Generate Outline':
				'Generate a structured outline of the document.',
			'Extract Entities':
				'Extract named entities, important organizations, dates, and references.',
			'Extract Tables':
				'Describe tabular information or structured fields present in the excerpts.',
			'Compare Active Documents':
				'Compare the active documents, noting agreements, differences, and gaps.',
			'Classify Document':
				'Classify the active document using the requested or most appropriate supported categories.',
			'Find Evidence':
				'Identify the excerpts that directly support or contradict the requested proposition.',
			'Generate Executive Summary':
				'Generate a concise executive summary focused on material facts, conclusions, and decisions.',
			'Extract Dates':
				'Extract important dates and explain their associated events or obligations.',
			'Extract Organizations':
				'Extract named organizations and describe their roles when supported by the excerpts.',
			'Extract Requirements':
				'Extract explicit requirements, constraints, and acceptance criteria from the excerpts.',
			'Extract Action Items':
				'Extract explicit or clearly supported action items and responsible parties when available.',
			'Identify Contradictions':
				'Identify material contradictions or inconsistencies across the active document excerpts.',
			'Identify Missing Information':
				'Identify information required by the request that is absent from the retrieved evidence.'
	}
	
	return action_map.get( action, action_map[ 'Answer Question' ] )

def build_instruction_block( ) -> str:
	"""Builds the unified Document Q&A instruction block from grounding, response-format, and action settings.

	Purpose:
		Builds the unified Document Q&A instruction block from grounding, response-format,
		and action settings.

	Returns:
		str: Text produced by the operation.
	"""
	require_grounding = bool( st.session_state.get( 'require_grounding', True ) )
	answer_from_excerpts_only = bool( st.session_state.get( 'answer_from_excerpts_only', True ) )
	grounding_failure_behavior = str( st.session_state.get( 'grounding_failure_behavior',
		'State Insufficient Information' ) )
	response_format = str(
		st.session_state.get( 'response_format', 'Markdown' ) or 'Markdown' ).strip( )
	doc_action = str(
		st.session_state.get( 'docqna_action', 'Answer Question' ) or 'Answer Question' )
	lines: List[ str ] = [ ]
	lines.append( 'Document Q&A Instructions:' )
	lines.append( f'- Action: {doc_action}' )
	lines.append( f'- Response Format: {response_format}' )
	lines.append( f'- Action Guidance: {create_docqna_instruction( doc_action )}' )
	if require_grounding:
		lines.append( '- Ground every answer in the retrieved document excerpts.' )
	
	if answer_from_excerpts_only:
		lines.append( '- Answer only from retrieved excerpts and supplied document context.' )
	if grounding_failure_behavior == 'Return Retrieved Excerpts':
		lines.append( '- When the answer is unsupported, return the most relevant excerpts instead.' )
	elif grounding_failure_behavior == 'Best Supported Answer':
		lines.append( '- When evidence is incomplete, provide the best supported answer and identify the gap.' )
	else:
		lines.append( '- When evidence is insufficient, state clearly that there is not enough information.' )
	
	if response_format == 'JSON':
		lines.append( '- Return valid JSON only.' )
	
	return '\n'.join( lines ).strip( )

def extract_text_bytes( file_bytes: bytes, file_name: str = '' ) -> str:
	"""Extracts text from supported document bytes.

	Purpose:
		Extracts TXT content directly, parses DOCX content with python-docx, and processes PDF pages
		individually so native text and Gemma vision OCR can be combined within mixed digital/scanned
		documents. OCR cache identity includes the extraction settings and active model/projector.

	Args:
		file_bytes (bytes): Document bytes to parse.
		file_name (str): Source file name used to select the parser.

	Returns:
		str: Extracted document text.
	"""
	if not file_bytes:
		return ''

	file_name_value = str( file_name or '' ).lower( )
	if file_name_value.endswith( '.docx' ):
		try:
			from docx import Document
			document = Document( BytesIO( file_bytes ) )
			parts: List[ str ] = [
				paragraph.text for paragraph in document.paragraphs
				if paragraph.text and paragraph.text.strip( )
			]
			for table in document.tables:
				for row in table.rows:
					values = [ cell.text.strip( ) for cell in row.cells ]
					if any( values ):
						parts.append( ' | '.join( values ) )
			return '\n'.join( parts ).strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'app'
			exception.cause = 'extract_text_bytes'
			exception.method = 'extract_text_bytes( file_bytes: bytes, file_name: str ) -> str'
			Logger( ).write( exception )
			return ''

	is_pdf = file_name_value.endswith( '.pdf' ) or file_name_value == ''
	if not is_pdf:
		try:
			return file_bytes.decode( errors='ignore' ).strip( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'app'
			exception.cause = 'extract_text_bytes'
			exception.method = 'extract_text_bytes( file_bytes: bytes, file_name: str ) -> str'
			Logger( ).write( exception )
			return ''

	if fitz is None:
		return ''

	include_page_markers = bool( st.session_state.get( 'include_page_markers', False ) )
	prefer_native_pdf_text = bool( st.session_state.get( 'prefer_native_pdf_text', True ) )
	ocr_enabled = bool( st.session_state.get( 'ocr_enabled', False ) )
	page_limit_value = str( st.session_state.get( 'ocr_page_limit', '5 Pages' ) )
	cache_seed = '|'.join( [
		hashlib.sha256( file_bytes ).hexdigest( ),
		str( ocr_enabled ),
		str( prefer_native_pdf_text ),
		str( include_page_markers ),
		page_limit_value,
		str( cfg.MODEL_PATH ),
		str( MMPROJ_PATH_OBJ or '' )
	] )
	cache_key = hashlib.sha256( cache_seed.encode( 'utf-8', errors='ignore' ) ).hexdigest( )
	ocr_cache = st.session_state.get( 'docqna_ocr_cache', { } )
	if isinstance( ocr_cache, dict ) and cache_key in ocr_cache:
		return str( ocr_cache[ cache_key ] or '' )

	try:
		doc = fitz.open( stream=file_bytes, filetype='pdf' )
		page_limit_map = {
			'1 Page': 1,
			'2 Pages': 2,
			'5 Pages': 5,
			'10 Pages': 10,
			'All Pages': len( doc )
		}
		ocr_page_limit = min( len( doc ), page_limit_map.get( page_limit_value, 5 ) )
		parts: List[ str ] = [ ]

		for page_index, page in enumerate( doc, start=1 ):
			page_text = page.get_text( 'text' ) or '' if prefer_native_pdf_text else ''
			usable_text = page_text.strip( )
			if (not usable_text and ocr_enabled and page_index <= ocr_page_limit
					and vision_runtime_available( )):
				pixmap = page.get_pixmap( matrix=fitz.Matrix( 2.0, 2.0 ), alpha=False )
				png_bytes = pixmap.tobytes( 'png' )
				usable_text = run_vision_turn(
					[ {
						'name': f'{file_name or "document.pdf"} page {page_index}',
						'bytes': png_bytes,
						'mime_type': 'image/png'
					} ],
					stream=False,
					output=None,
					show_errors=False,
					instruction_override='Extract all visible text from this document page. Return the text only.',
					response_format_override='Plain Text'
				)
			if include_page_markers:
				parts.append( f'[Page {page_index}]' )
			if usable_text:
				parts.append( usable_text )

		result = '\n'.join( parts ).strip( )
		if isinstance( ocr_cache, dict ):
			ocr_cache[ cache_key ] = result
			st.session_state[ 'docqna_ocr_cache' ] = ocr_cache
		return result
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'extract_text_bytes'
		exception.method = 'extract_text_bytes( file_bytes: bytes, file_name: str ) -> str'
		Logger( ).write( exception )
		return ''


def route_document_query( prompt: str ) -> str:
	"""Routes a document question or action through retrieved context and the local generation pipeline.

	Purpose:
		Routes a document question or action through retrieved context and the local
		generation pipeline.

	Args:
		prompt: Prompt or document request text.

	Returns:
		str: Text produced by the operation.
	"""
	user_input = build_docqna_input( user_query=prompt,
		k=int( st.session_state.get( 'retrieval_k', 6 ) ) )
	
	if not user_input:
		user_input = (prompt or '').strip( )
	
	return run_direct_llm_turn(
		system_instruction=get_effective_system_instructions( ),
		user_input=user_input,
		temperature=float( st.session_state.get( 'temperature', 0.0 ) ),
		top_p=float( st.session_state.get( 'top_percent', 0.95 ) ),
		repeat_penalty=float( st.session_state.get( 'repeat_penalty', 1.1 ) ),
		max_tokens=int( st.session_state.get( 'max_tokens', 1024 ) ) or 1024,
		stream=False,
		output=None,
		response_format=str( st.session_state.get( 'response_format', 'Markdown' ) ) )

def summarize_document( ) -> str:
	"""Requests a structured summary of the active document set through the Document Q&A routing layer.

	Purpose:
		Requests a structured summary of the active document set through the Document Q&A
		routing layer.

	Returns:
		str: Text produced by the operation.
	"""
	summary_prompt = """
		Provide a clear, structured summary of the active document set.
		Include:
		- Purpose
		- Key themes
		- Major conclusions
		- Important data points
		- Open questions or uncertainties
	"""
	
	return route_document_query( summary_prompt.strip( ) )

def compute_fingerprint( active_docs: List[ str ], doc_bytes: Dict[ str, bytes ] ) -> str:
	"""Computes a stable fingerprint for active document names and byte contents to support index cache invalidation.

	Purpose:
		Computes a stable fingerprint for active document names and byte contents to support
		index cache invalidation.

	Args:
		active_docs: Active document names.
		doc_bytes: Mapping of document names to byte contents.

	Returns:
		str: Text produced by the operation.
	"""
	h = hashlib.sha256( )
	for name in sorted( active_docs ):
		b = doc_bytes.get( name, b'' )
		h.update( name.encode( 'utf-8', errors='ignore' ) )
		h.update( len( b ).to_bytes( 8, 'little', signed=False ) )
		h.update( hashlib.sha256( b ).digest( ) )
	return h.hexdigest( )

def extract_text( file_bytes: bytes, file_name: str = '' ) -> str:
	"""Extracts document text using the configured byte-level parsing function.

	Purpose:
		Extracts document text through the byte-level parser without recursively invoking the public
		helper.

	Args:
		file_bytes (bytes): Document bytes to parse.
		file_name (str): Source file name used to infer parsing behavior.

	Returns:
		str: Extracted document text.
	"""
	return extract_text_bytes( file_bytes=file_bytes, file_name=file_name )

def load_sqlite_vec( conn: sqlite3.Connection ) -> bool:
	"""Attempts to load the sqlite-vec extension into an active SQLite connection.

	Purpose:
		Attempts to load the sqlite-vec extension into an active SQLite connection.

	Args:
		conn: conn value used by this workflow.

	Returns:
		bool: Boolean result produced by the operation.
	"""
	try:
		import sqlite_vec
		
		sqlite_vec.load( conn )
		return True
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'load_sqlite_vec'
		exception.method = 'load_sqlite_vec( conn: sqlite3.Connection ) -> bool'
		Logger( ).write( exception )
		return False

def ensure_schema( dim: int ) -> bool:
	"""Creates the sqlite-vec virtual table used by Document Q&A retrieval when vector support is available.

	Purpose:
		Creates the sqlite-vec virtual table used by Document Q&A retrieval when vector
		support is available.

	Args:
		dim: Embedding dimension.

	Returns:
		bool: Boolean result produced by the operation.
	"""
	conn = create_connection( )
	try:
		ok = load_sqlite_vec( conn )
		if not ok:
			return False
		
		cur = conn.cursor( )
		cur.execute(
			f'''
			CREATE VIRTUAL TABLE IF NOT EXISTS docqna_vec
			USING vec0(
				embedding float[{int( dim )}],
				doc_name TEXT,
				chunk TEXT
			);
			'''
		)
		conn.commit( )
		return True
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'ensure_schema'
		exception.method = 'ensure_schema( dim: int ) -> bool'
		Logger( ).write( exception )
		return False
	finally:
		conn.close( )

def build_docqna_inventory( ) -> List[ Dict[ str, Any ] ]:
	"""Builds document inventory rows for active uploads, including byte size, extracted text length, and chunk count.

	Purpose:
		Builds document inventory rows for active uploads, including byte size, extracted
		text length, and chunk count.

	Returns:
		List[Dict[str, Any]]: Result produced by the operation.
	"""
	rows: List[ Dict[ str, Any ] ] = [ ]
	active_docs = st.session_state.get( 'active_docs', [ ] )
	doc_bytes = st.session_state.get( 'doc_bytes', { } )
	for name in active_docs:
		b = doc_bytes.get( name, b'' )
		text = extract_text( b, name ) if b else ''
		chunks = chunk_text( text ) if text else [ ]
		rows.append( {
				'Name': name,
				'SizeBytes': len( b ) if b else 0,
				'TextLength': len( text ) if text else 0,
				'ChunkCount': len( chunks ),
				'Loaded': bool( b )
		} )
	
	return rows

def get_docqna_names( ) -> str:
	"""Builds a human-readable list of active document names for Document Q&A prompt context.

	Purpose:
		Builds a human-readable list of active document names for Document Q&A prompt
		context.

	Returns:
		str: Text produced by the operation.
	"""
	active_docs = st.session_state.get( 'active_docs', [ ] )
	if not isinstance( active_docs, list ) or len( active_docs ) == 0:
		return 'No active documents'
	return ', '.join( [ str( name ) for name in active_docs ] )

def rebuild_index( embedder: Any | None ) -> None:
	"""Builds or refreshes the Document Q&A vector index when active documents or chunk settings change.

	Purpose:
		Builds or refreshes the Document Q&A vector index when active documents or chunk
		settings change.

	Args:
		embedder: Sentence embedding model used to encode document chunks.
	"""
	if embedder is None:
		st.session_state[ 'docqna_vec_ready' ] = False
		st.session_state[ 'docqna_fallback_rows' ] = [ ]
		st.session_state[ 'docqna_chunk_count' ] = 0
		return

	active_docs: List[ str ] = st.session_state.get( 'active_docs', [ ] )
	doc_bytes: Dict[ str, bytes ] = st.session_state.get( 'doc_bytes', { } )
	retrieval_chunk_size = int( st.session_state.get( 'retrieval_chunk_size', 1200 ) )
	retrieval_chunk_overlap = int( st.session_state.get( 'retrieval_chunk_overlap', 200 ) )
	
	fp_seed = (
		f'{retrieval_chunk_size}|{retrieval_chunk_overlap}|'
		f'{bool( st.session_state.get( "ocr_enabled", False ) )}|'
		f'{bool( st.session_state.get( "prefer_native_pdf_text", True ) )}|'
		f'{bool( st.session_state.get( "include_page_markers", False ) )}|'
		f'{str( st.session_state.get( "ocr_page_limit", "5 Pages" ) )}|'
		f'{str( cfg.MODEL_PATH )}|{str( MMPROJ_PATH_OBJ or "" )}|'
	)
	fp_seed += compute_fingerprint( active_docs, doc_bytes )
	fp = hashlib.sha256( fp_seed.encode( 'utf-8', errors='ignore' ) ).hexdigest( )
	
	if fp and fp == st.session_state.get( 'docqna_fingerprint', '' ):
		st.session_state[ 'docqna_inventory_rows' ] = build_docqna_inventory( )
		return
	
	st.session_state[ 'docqna_fingerprint' ] = fp
	st.session_state[ 'docqna_chunk_count' ] = 0
	st.session_state[ 'docqna_fallback_rows' ] = [ ]
	st.session_state[ 'docqna_inventory_rows' ] = build_docqna_inventory( )
	
	dim_value = getattr( embedder, 'get_sentence_embedding_dimension', lambda: 384 )( )
	dim = int( dim_value ) if dim_value else 384
	
	prefer_sqlite_vec = bool( st.session_state.get( 'prefer_sqlite_vec', True ) )
	vec_ready = False
	if prefer_sqlite_vec:
		vec_ready = ensure_schema( dim )
	
	st.session_state[ 'docqna_vec_ready' ] = bool( vec_ready )
	
	conn = create_connection( )
	try:
		cur = conn.cursor( )
		
		if vec_ready:
			try:
				cur.execute( 'DELETE FROM docqna_vec;' )
				conn.commit( )
			except Exception as e:
				exception = Error( e )
				exception.module = 'app'
				exception.cause = 'rebuild_index'
				exception.method = 'rebuild_index( embedder: Any | None ) -> None'
				Logger( ).write( exception )
				st.session_state[ 'docqna_vec_ready' ] = False
				vec_ready = False
		
		total_chunks = 0
		fallback_rows: List[ Tuple[ str, str, bytes ] ] = [ ]
		
		for name in active_docs:
			b = doc_bytes.get( name )
			if not b:
				continue
			
			text = extract_text( b, name )
			if not text:
				continue
			
			chunks = chunk_text(
				text,
				size=retrieval_chunk_size,
				overlap=retrieval_chunk_overlap
			)
			if not chunks:
				continue
			
			vecs = embedder.encode( chunks, show_progress_bar=False )
			vecs = np.asarray( vecs, dtype=np.float32 )
			
			if vec_ready:
				for chunk_text_value, v in zip( chunks, vecs ):
					cur.execute(
						'INSERT INTO docqna_vec ( embedding, doc_name, chunk ) VALUES ( ?, ?, ? );',
						(v.tobytes( ), name, chunk_text_value)
					)
			else:
				for chunk_text_value, v in zip( chunks, vecs ):
					fallback_rows.append( (name, chunk_text_value, v.tobytes( )) )
			
			total_chunks += int( len( chunks ) )
		
		conn.commit( )
		st.session_state[ 'docqna_chunk_count' ] = total_chunks
		
		if not vec_ready:
			st.session_state[ 'docqna_fallback_rows' ] = fallback_rows
		else:
			st.session_state[ 'docqna_fallback_rows' ] = [ ]
	
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'rebuild_index'
		exception.method = 'rebuild_index( embedder: Any | None ) -> None'
		Logger( ).write( exception )
		st.session_state[ 'docqna_vec_ready' ] = False
		st.session_state[ 'docqna_fallback_rows' ] = [ ]
		st.session_state[ 'docqna_chunk_count' ] = 0
	finally:
		conn.close( )

def retrieve_chunks( query: str, k: int = None ) -> List[ Tuple[ str, str, float ] ]:
	"""Retrieves top-ranked document chunks for a query using sqlite-vec when available and cosine fallback otherwise.

	Purpose:
		Retrieves top-ranked document chunks for a query using sqlite-vec when available and
		cosine fallback otherwise.

	Args:
		query: Query text.
		k: Number of chunks to retrieve.

	Returns:
		List[Tuple[str, str, float]]: Result produced by the operation.
	"""
	if not query or not query.strip( ):
		return [ ]
	if embedder is None:
		st.session_state[ 'docqna_vec_ready' ] = False
		st.session_state[ 'docqna_fallback_rows' ] = [ ]
		return [ ]
	
	try:
		if bool( st.session_state.get( 'docqna_rebuild_each_query', False ) ):
			st.session_state[ 'docqna_fingerprint' ] = ''
		rebuild_index( embedder )
		k_value = int( k ) if k is not None else int( st.session_state.get( 'retrieval_k', 6 ) )
		if k_value <= 0:
			k_value = 6
		qv = embedder.encode( [ query ], show_progress_bar=False )
		qv = np.asarray( qv, dtype=np.float32 )[ 0 ]
	except Exception as e:
		exception = Error( e )
		exception.module = 'app'
		exception.cause = 'retrieve_chunks'
		exception.method = 'retrieve_chunks( query: str, k: int ) -> List[Tuple[str, str, float]]'
		Logger( ).write( exception )
		return [ ]
	backend = str( st.session_state.get( 'retrieval_backend', 'Automatic' ) )
	use_vector_backend = backend in ( 'Automatic', 'sqlite-vec' )
	if use_vector_backend and bool( st.session_state.get( 'docqna_vec_ready', False ) ):
		conn = create_connection( )
		try:
			load_sqlite_vec( conn )
			cur = conn.cursor( )
			cur.execute(
				'''
                SELECT doc_name, chunk, distance
                FROM docqna_vec
                WHERE embedding MATCH ?
                ORDER BY distance ASC LIMIT ?;
				''',
				(qv.tobytes( ), int( k_value )) )
			rows = cur.fetchall( )
			return [ (r[ 0 ], r[ 1 ], float( r[ 2 ] )) for r in rows ]
		except Exception as e:
			exception = Error( e )
			exception.module = 'app'
			exception.cause = 'retrieve_chunks'
			exception.method = 'retrieve_chunks( query: str, k: int ) -> List[Tuple[str, str, float]]'
			Logger( ).write( exception )
			st.session_state[ 'docqna_vec_ready' ] = False
		finally:
			conn.close( )
	
	if backend == 'sqlite-vec':
		return [ ]
	if not bool( st.session_state.get( 'allow_similarity_fallback', True ) ):
		return [ ]
	
	fallback_rows: List[ Tuple[ str, str, bytes ] ] = st.session_state.get( 'docqna_fallback_rows',
		[ ] )
	results: List[ Tuple[ str, str, float ] ] = [ ]
	for doc_name, chunk_text_value, vec_blob in fallback_rows:
		if not vec_blob:
			continue
		
		v = np.frombuffer( vec_blob, dtype=np.float32 )
		if v.size == 0 or v.size != qv.size:
			continue
		try:
			score = cosine_similarity( qv, v )
		except Exception as e:
			exception = Error( e )
			exception.module = 'app'
			exception.cause = 'retrieve_chunks'
			exception.method = 'retrieve_chunks( query: str, k: int ) -> List[Tuple[str, str, float]]'
			Logger( ).write( exception )
			continue
		results.append( (doc_name, chunk_text_value, float( score )) )
	
	results.sort( key=lambda r: r[ 2 ], reverse=True )
	return results[ : int( k_value ) ]

def build_docqna_input( user_query: str, k: int = None ) -> str:
	"""Builds a grounded Document Q&A prompt from instructions, active document names, retrieved excerpts, and user request.

	Purpose:
		Builds a grounded Document Q&A prompt from instructions, active document names,
		retrieved excerpts, and user request.

	Args:
		user_query: user query value used by this workflow.
		k: Number of chunks to retrieve.

	Returns:
		str: Text produced by the operation.
	"""
	doc_instruction_block = build_instruction_block( )
	hits = retrieve_chunks( user_query, k=k )
	st.session_state[ 'docqna_last_retrieval' ] = hits
	context_blocks: List[ str ] = [ ]
	for doc_name, chunk, score in hits:
		context_blocks.append( f'[Document: {doc_name}]\n{chunk}'.strip( ) )
	
	semantic_blocks: List[ str ] = [ ]
	semantic_context_buffer = st.session_state.get( 'semantic_context_buffer', [ ] )
	if bool( st.session_state.get( 'docqna_include_semantic_context', True ) ) and isinstance(
			semantic_context_buffer, list ):
		for value in semantic_context_buffer:
			if isinstance( value, str ) and value.strip( ):
				semantic_blocks.append( f'[Semantic Context]\n{value.strip( )}' )
	if str( st.session_state.get( 'docqna_context_order', 'Retrieved First' ) ) == 'Semantic First':
		context_blocks = semantic_blocks + context_blocks
	else:
		context_blocks.extend( semantic_blocks )
	
	context = '\n\n'.join( context_blocks ).strip( )
	active_doc_names = get_docqna_names( )
	prompt_parts: List[ str ] = [ ]
	if doc_instruction_block:
		prompt_parts.append( doc_instruction_block )
	
	prompt_parts.append( f'Active Documents:\n{active_doc_names}' )
	
	if context:
		prompt_parts.append(
			'Use the following retrieved document excerpts as the evidence base for your answer.\n\n'
			f'{context}' )
	else:
		prompt_parts.append(
			'No retrieved document excerpts were available for this question.' )
	
	prompt_parts.append( f'User Request:\n{user_query}\n\nAnswer:' )
	return '\n\n'.join( prompt_parts ).strip( )

# ------------ SEMANTIC SEARCH UTLITIES -------------------

def decode_embedding_rows( ) -> List[ Tuple[ str, np.ndarray ] ]:
	"""Reads semantic-search embedding rows and decodes vectors.

	Purpose:
		Reads semantic chunk/vector rows from the existing embeddings schema without requiring a
		database migration.

	Returns:
		List[Tuple[str, np.ndarray]]: Chunk text and numeric vector rows.
	"""
	rows_out: List[ Tuple[ str, np.ndarray ] ] = [ ]
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		rows = conn.execute( 'SELECT chunk, vector FROM embeddings' ).fetchall( )
	for chunk_text_value, vector_blob in rows:
		if not vector_blob:
			continue
		vector = np.frombuffer( vector_blob, dtype=np.float32 )
		if vector.size == 0:
			continue
		rows_out.append( (str( chunk_text_value or '' ), vector) )
	return rows_out


def clear_semantic_index( ) -> None:
	"""Clears the semantic-search embeddings table and related diagnostics.

	Purpose:
		Removes indexed chunks while preserving the embeddings schema and resets all Streamlit semantic
		search result state.

	Returns:
		None: This function performs its work through database and session-state side effects.
	"""
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		conn.execute( 'DELETE FROM embeddings' )
		conn.commit( )
	st.session_state[ 'semantic_result_rows' ] = [ ]
	st.session_state[ 'semantic_selected_rows' ] = [ ]
	st.session_state[ 'semantic_index_chunk_count' ] = 0
	st.session_state[ 'semantic_index_dim' ] = 0
	st.session_state[ 'semantic_index_doc_count' ] = 0
	st.session_state[ 'semantic_uploaded_names' ] = [ ]
	st.session_state[ 'semantic_last_query' ] = ''


def build_semantic_index( uploaded_files: List[ Any ] ) -> Dict[ str, Any ]:
	"""Builds or appends a document-aware semantic chunk index.

	Purpose:
		Extracts text, creates bounded overlapping chunks, embeds them, and persists vectors using the
		existing embeddings table schema.

	Args:
		uploaded_files (List[Any]): Uploaded Streamlit files.

	Returns:
		Dict[str, Any]: Index build status and diagnostics.
	"""
	if embedder is None:
		return { 'success': False, 'message': 'Embedding model unavailable.', 'doc_count': 0,
			'chunk_count': 0, 'vector_dim': 0 }
	chunk_size = int( st.session_state.get( 'semantic_chunk_size', 1200 ) )
	chunk_overlap = int( st.session_state.get( 'semantic_chunk_overlap', 200 ) )
	clear_existing = bool( st.session_state.get( 'semantic_clear_existing', True ) )
	append_existing = bool( st.session_state.get( 'semantic_append_existing', False ) )
	if clear_existing and not append_existing:
		clear_semantic_index( )
	chunk_rows: List[ Tuple[ str, str ] ] = [ ]
	doc_names: List[ str ] = [ ]
	for uploaded_file in uploaded_files:
		try:
			file_name = str( getattr( uploaded_file, 'name', '' ) or '' ).strip( )
			file_bytes = uploaded_file.getvalue( )
		except Exception as e:
			exception = Error( e )
			exception.module = 'app'
			exception.cause = 'build_semantic_index'
			exception.method = 'build_semantic_index( uploaded_files: List[Any] ) -> Dict[str, Any]'
			Logger( ).write( exception )
			continue
		if not file_name or not file_bytes:
			continue
		text_value = extract_text( file_bytes=file_bytes, file_name=file_name )
		if not text_value:
			try:
				text_value = file_bytes.decode( errors='ignore' )
			except Exception:
				text_value = ''
		if not text_value:
			continue
		chunks = chunk_text( text=text_value, size=chunk_size, overlap=chunk_overlap )
		for chunk_value in chunks:
			chunk_rows.append( (file_name, chunk_value) )
		if chunks:
			doc_names.append( file_name )
	if not chunk_rows:
		return { 'success': False, 'message': 'No extractable text was found in the uploaded files.',
			'doc_count': 0, 'chunk_count': 0, 'vector_dim': 0 }
	chunk_values = [ row[ 1 ] for row in chunk_rows ]
	vecs = np.asarray( embedder.encode( chunk_values, show_progress_bar=False ), dtype=np.float32 )
	with sqlite3.connect( cfg.DB_PATH ) as conn:
		for (_document_name, chunk_value), vec in zip( chunk_rows, vecs ):
			conn.execute( 'INSERT INTO embeddings (chunk, vector) VALUES (?, ?)',
				(chunk_value, vec.tobytes( )) )
		conn.commit( )
	vector_dim = int( vecs.shape[ 1 ] ) if len( vecs.shape ) == 2 else 0
	st.session_state[ 'semantic_uploaded_names' ] = doc_names
	st.session_state[ 'semantic_index_doc_count' ] = len( set( doc_names ) )
	st.session_state[ 'semantic_index_chunk_count' ] = len( chunk_rows )
	st.session_state[ 'semantic_index_dim' ] = vector_dim
	return { 'success': True, 'message': 'Semantic index built successfully.',
		'doc_count': len( set( doc_names ) ), 'chunk_count': len( chunk_rows ), 'vector_dim': vector_dim }


def query_semantic_index( query_text: str ) -> List[ Dict[ str, Any ] ]:
	"""Queries the semantic index and returns ranked chunk rows.

	Purpose:
		Ranks stored semantic chunks using cosine similarity without altering the existing embeddings
		database schema.

	Args:
		query_text (str): Semantic query text.

	Returns:
		List[Dict[str, Any]]: Ranked semantic result rows.
	"""
	if not query_text or not query_text.strip( ) or embedder is None:
		return [ ]
	top_k = int( st.session_state.get( 'semantic_top_k', 8 ) )
	min_similarity = float( st.session_state.get( 'semantic_min_similarity', 0.0 ) )
	rows = decode_embedding_rows( )
	if not rows:
		return [ ]
	query_vector = np.asarray(
		embedder.encode( [ query_text.strip( ) ], show_progress_bar=False )[ 0 ],
		dtype=np.float32
	)
	scored_rows: List[ Dict[ str, Any ] ] = [ ]
	for chunk_text_value, vector in rows:
		if vector.size != query_vector.size:
			continue
		score = cosine_similarity( query_vector, vector )
		if score < min_similarity:
			continue
		scored_rows.append( {
			'Selected': False,
			'Score': float( score ),
			'Chunk': chunk_text_value,
			'Length': len( chunk_text_value )
		} )
	scored_rows.sort( key=lambda row: row[ 'Score' ], reverse=True )
	results = scored_rows[ :top_k ]
	for rank, row in enumerate( results, start=1 ):
		row[ 'Rank' ] = rank
	st.session_state[ 'semantic_last_query' ] = query_text.strip( )
	st.session_state[ 'semantic_result_rows' ] = results
	return results


def create_semantic_context( ) -> str:
	"""Builds a reusable semantic-context text block from selected semantic-search rows.

	Purpose:
		Builds a reusable semantic-context text block from selected semantic-search rows.

	Returns:
		str: Text produced by the operation.
	"""
	selected_rows = st.session_state.get( 'semantic_selected_rows', [ ] )
	if not isinstance( selected_rows, list ) or len( selected_rows ) == 0:
		return ''
	
	context_parts: List[ str ] = [ ]
	for idx, row in enumerate( selected_rows, start=1 ):
		chunk_text_value = str( row.get( 'Chunk', '' ) or '' ).strip( )
		score_value = row.get( 'Score', '' )
		if not chunk_text_value:
			continue
		
		context_parts.append( f'[Semantic Chunk {idx} | Score: {score_value}]\n{chunk_text_value}' )
	
	return '\n\n'.join( context_parts ).strip( )

def extract_selected_rows( edited_rows: List[ Dict[ str, Any ] ] ) -> List[ Dict[ str, Any ] ]:
	"""Extracts selected rows from a Streamlit data editor payload for semantic-context routing.

	Purpose:
		Extracts selected rows from a Streamlit data editor payload for semantic-context
		routing.

	Args:
		edited_rows: edited rows value used by this workflow.

	Returns:
		List[Dict[str, Any]]: Result produced by the operation.
	"""
	selected: List[ Dict[ str, Any ] ] = [ ]
	if not isinstance( edited_rows, list ):
		return selected
	
	for row in edited_rows:
		if isinstance( row, dict ) and bool( row.get( 'Selected', False ) ):
			selected.append( row )
	
	return selected

def send_text_chunks( ) -> None:
	"""Adds selected semantic chunks to the shared document context buffer used by Text Generation.

	Purpose:
		Adds selected semantic chunks to the shared document context buffer used by Text
		Generation.
	"""
	context_text = create_semantic_context( )
	if not context_text:
		return
	
	existing_docs = st.session_state.get( 'basic_docs', [ ] )
	if not isinstance( existing_docs, list ):
		existing_docs = [ ]
	
	existing_docs.append( context_text )
	st.session_state[ 'basic_docs' ] = existing_docs
	st.session_state[ 'use_semantic' ] = True

def send_docqna_chunks( ) -> None:
	"""Adds selected semantic chunks to the context buffer used by Document Q&A prompt construction.

	Purpose:
		Adds selected semantic chunks to the context buffer used by Document Q&A prompt
		construction.
	"""
	context_text = create_semantic_context( )
	if not context_text:
		return
	
	buffer_rows = st.session_state.get( 'semantic_context_buffer', [ ] )
	if not isinstance( buffer_rows, list ):
		buffer_rows = [ ]
	
	buffer_rows.append( context_text )
	st.session_state[ 'semantic_context_buffer' ] = buffer_rows

# ==============================================================================
# Init
# ==============================================================================
initialize_database( )
embedder = load_embedder( )

if not isinstance( st.session_state.get( 'messages' ), list ):
	st.session_state[ 'messages' ] = [ ]

if len( st.session_state[ 'messages' ] ) == 0:
	st.session_state[ 'messages' ] = load_history( )

if 'system_instructions' not in st.session_state:
	st.session_state[ 'system_instructions' ] = ''

st.set_page_config( page_title='Bro', layout='wide', page_icon=cfg.FAVICON )
st.caption( cfg.APP_SUBTITLE )

# ==============================================================================
# SIDEBAR
# ==============================================================================
with st.sidebar:
	style_subheaders( )
	st.logo( cfg.LOGO, size='large' )
	st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
	with st.expander( label='Mode', expanded=True ):
		mode_options = list( cfg.MODES )
		if 'Image to Text' not in mode_options:
			mode_options.append( 'Image to Text' )
		mode = st.radio( label='Select', options=mode_options, index=0 )

# ==============================================================================
# TEXT GENERATION MODE
# ==============================================================================
if mode == 'Text Generation':
	messages = st.session_state.get( 'messages', [ ] )
	max_tokens = st.session_state.get( 'max_tokens', 0 )
	top_percent = st.session_state.get( 'top_percent', 0.0 )
	top_k = st.session_state.get( 'top_k', 0 )
	temperature = st.session_state.get( 'temperature', 0.0 )
	is_grounded = st.session_state.get( 'is_grounded', False )
	frequency_penalty = st.session_state.get( 'frequency_penalty', 0.0 )
	presense_penalty = st.session_state.get( 'presense_penalty', 0.0 )
	repeat_penalty = st.session_state.get( 'repeat_penalty', 0.0 )
	repeat_window = st.session_state.get( 'repeat_window', 0.0 )
	cpu_threads = st.session_state.get( 'cpu_threads', cfg.CORES )
	context_window = st.session_state.get( 'context_window', cfg.DEFAULT_CTX )
	left, center, right = st.columns( [ 0.025, 0.95, 0.025 ] )
	with center:
		st.subheader( '💬 Text Generation', help=cfg.TEXT_GENERATION )
		st.divider( )
		
		# ------------------------------------------------------------------
		# Expander — Mind Controls
		# ------------------------------------------------------------------
		with st.expander( label='Mind Controls', icon='🧠', expanded=False ):
			if st.session_state.get( 'task_preset' ) not in get_prompt_task_types( ):
				st.session_state[ 'task_preset' ] = 'Chat'
			if st.session_state.get( 'response_format' ) not in get_response_formats( ):
				st.session_state[ 'response_format' ] = 'Markdown'
			if st.session_state.get( 'response_language' ) not in get_spoken_languages( False ):
				st.session_state[ 'response_language' ] = 'English'
			if st.session_state.get( 'translation_source_language' ) not in get_spoken_languages( True ):
				st.session_state[ 'translation_source_language' ] = 'Auto Detect'
			if st.session_state.get( 'translation_target_language' ) not in get_spoken_languages( False ):
				st.session_state[ 'translation_target_language' ] = 'English'

			with st.expander( label='Task Preset', icon='🧭', expanded=False ):
				task_c1, task_c2, task_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True,
					gap='medium' )
				with task_c1:
					st.selectbox( label='Task Type', options=get_prompt_task_types( ), key='task_preset' )
				with task_c2:
					st.selectbox( label='Task Detail', options=TASK_DETAIL_OPTIONS, key='task_detail' )
				with task_c3:
					st.selectbox( label='Task Focus', options=TASK_FOCUS_OPTIONS, key='task_focus' )
				if st.button( label='Reset', key='task_preset_reset', width='stretch', icon='🔄' ):
					for key in [ 'task_preset', 'task_detail', 'task_focus' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Reasoning Controls', icon='🧩', expanded=False ):
				reason_c1, reason_c2, reason_c3, reason_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with reason_c1:
					st.selectbox( label='Reasoning Depth', options=[ 'Low', 'Medium', 'High' ],
						key='reasoning_depth' )
				with reason_c2:
					st.toggle( label='Answer Only', value=bool( st.session_state.get( 'answer_only', False ) ),
						key='answer_only' )
				with reason_c3:
					st.toggle( label='Use Self-Check', value=bool( st.session_state.get( 'use_self_check', False ) ),
						key='use_self_check' )
				with reason_c4:
					st.toggle( label='Prefer Deterministic Reasoning',
						value=bool( st.session_state.get( 'deterministic_reasoning', False ) ),
						key='deterministic_reasoning' )
				if st.button( label='Reset', key='reasoning_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'reasoning_depth', 'answer_only', 'use_self_check', 'deterministic_reasoning' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Coding Controls', icon='🧾', expanded=False ):
				code_c1, code_c2, code_c3, code_c4, code_c5 = st.columns(
					[ 0.2, 0.2, 0.2, 0.2, 0.2 ], border=True, gap='medium' )
				with code_c1:
					st.selectbox( label='Code Language', options=CODING_LANGUAGE_OPTIONS, key='coding_language' )
				with code_c2:
					st.selectbox( label='Coding Task', options=CODING_TASK_OPTIONS, key='coding_task' )
				with code_c3:
					st.toggle( label='Include Comments',
						value=bool( st.session_state.get( 'coding_include_comments', True ) ),
						key='coding_include_comments' )
				with code_c4:
					st.toggle( label='Use Editor Format',
						value=bool( st.session_state.get( 'coding_editor_format', True ) ),
						key='coding_editor_format' )
				with code_c5:
					st.toggle( label='Emit Fenced Code',
						value=bool( st.session_state.get( 'coding_fenced_output', True ) ),
						key='coding_fenced_output' )
				if st.button( label='Reset', key='coding_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'coding_language', 'coding_task', 'coding_include_comments',
							'coding_editor_format', 'coding_fenced_output' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Writing Controls', icon='✍️', expanded=False ):
				write_c1, write_c2, write_c3, write_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with write_c1:
					st.selectbox( label='Writing Task', options=WRITING_TASK_OPTIONS, key='writing_task' )
				with write_c2:
					st.selectbox( label='Tone', options=WRITING_TONE_OPTIONS, key='writing_tone' )
				with write_c3:
					st.selectbox( label='Audience', options=WRITING_AUDIENCE_OPTIONS, key='writing_audience' )
				with write_c4:
					st.selectbox( label='Length', options=RESPONSE_LENGTH_OPTIONS, key='writing_length' )
				if st.button( label='Reset', key='writing_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'writing_task', 'writing_tone', 'writing_audience', 'writing_length' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Translation Controls', icon='🌐', expanded=False ):
				trans_c1, trans_c2, trans_c3, trans_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with trans_c1:
					st.selectbox( label='Source Language', options=get_spoken_languages( True ),
						key='translation_source_language' )
				with trans_c2:
					st.selectbox( label='Target Language', options=get_spoken_languages( False ),
						key='translation_target_language' )
				with trans_c3:
					st.selectbox( label='Translation Mode', options=TRANSLATION_MODE_OPTIONS,
						key='translation_mode' )
				with trans_c4:
					st.toggle( label='Preserve Formatting',
						value=bool( st.session_state.get( 'translation_preserve_formatting', True ) ),
						key='translation_preserve_formatting' )
				if st.button( label='Reset', key='translation_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'translation_source_language', 'translation_target_language',
							'translation_mode', 'translation_preserve_formatting' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Classification Controls', icon='🏷️', expanded=False ):
				class_c1, class_c2, class_c3, class_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with class_c1:
					st.selectbox( label='Classification Type', options=CLASSIFICATION_TYPE_OPTIONS,
						key='classification_type' )
				with class_c2:
					st.toggle( label='Return Confidence',
						value=bool( st.session_state.get( 'classification_confidence', False ) ),
						key='classification_confidence' )
				with class_c3:
					st.toggle( label='Allow Unknown',
						value=bool( st.session_state.get( 'classification_allow_unknown', True ) ),
						key='classification_allow_unknown' )
				with class_c4:
					st.toggle( label='Explain Classification',
						value=bool( st.session_state.get( 'classification_explain', False ) ),
						key='classification_explain' )
				if st.button( label='Reset', key='classification_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'classification_type', 'classification_confidence',
							'classification_allow_unknown', 'classification_explain' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Response Controls', icon='↔️', expanded=False ):
				resp_c1, resp_c2, resp_c3, resp_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with resp_c1:
					st.selectbox( label='Response Format', options=get_response_formats( ), key='response_format' )
				with resp_c2:
					st.selectbox( label='Response Language', options=get_spoken_languages( False ),
						key='response_language' )
				with resp_c3:
					st.selectbox( label='Response Length', options=RESPONSE_LENGTH_OPTIONS,
						key='response_length' )
				with resp_c4:
					st.toggle( label='Include Headings',
						value=bool( st.session_state.get( 'response_include_headings', True ) ),
						key='response_include_headings' )
				if st.button( label='Reset', key='response_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'response_format', 'response_language', 'response_length',
							'response_include_headings' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Inference Settings', icon='🎚️', expanded=False ):
				inf_c1, inf_c2, inf_c3, inf_c4, inf_c5 = st.columns(
					[ 0.2, 0.2, 0.2, 0.2, 0.2 ], border=True, gap='medium' )
				with inf_c1:
					st.slider( label='Temperature', min_value=0.0, max_value=1.0,
						key='temperature', help=cfg.TEMPERATURE )
				with inf_c2:
					st.slider( label='Top-P', min_value=0.0, max_value=1.0, step=0.01,
						key='top_percent', help=cfg.TOP_P )
				with inf_c3:
					st.slider( label='Top-K', min_value=0, max_value=50, step=1,
						key='top_k', help=cfg.TOP_K )
				with inf_c4:
					st.slider( label='Repeat Penalty', min_value=0.0, max_value=2.0, step=0.05,
						key='repeat_penalty', help=cfg.REPEAT_PENALTY )
				with inf_c5:
					st.slider( label='Repeat Window', min_value=0, max_value=1024, step=16,
						key='repeat_window', help=cfg.REPEAT_WINDOW )
				inf_c6, inf_c7, inf_c8, inf_c9 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with inf_c6:
					st.slider( label='Presence Penalty', min_value=0.0, max_value=2.0, step=0.05,
						key='presense_penalty', help=cfg.PRESENCE_PENALTY )
				with inf_c7:
					st.slider( label='Frequency Penalty', min_value=0.0, max_value=2.0, step=0.05,
						key='frequency_penalty', help=cfg.FREQUENCY_PENALTY )
				with inf_c8:
					st.slider( label='Random Seed', min_value=0, max_value=4096, step=1,
						key='random_seed', help=cfg.SEED )
				with inf_c9:
					st.slider( label='Max Tokens', min_value=0, max_value=8192, step=128,
						key='max_tokens', help=cfg.MAX_TOKENS )
				if st.button( label='Reset', key='probability_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'temperature', 'top_percent', 'top_k', 'repeat_penalty', 'repeat_window',
							'presense_penalty', 'frequency_penalty', 'random_seed', 'max_tokens' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Context Controls', icon='🎛️', expanded=False ):
				ctx_c1, ctx_c2, ctx_c3, ctx_c4, ctx_c5 = st.columns(
					[ 0.2, 0.2, 0.2, 0.2, 0.2 ], border=True, gap='medium' )
				with ctx_c1:
					st.slider( label='Context Window', min_value=0, max_value=131072, step=512,
						key='context_window', help=cfg.CONTEXT_WINDOW )
				with ctx_c2:
					st.toggle( label='Use Conversation History',
						value=bool( st.session_state.get( 'use_chat_history', True ) ), key='use_chat_history' )
				with ctx_c3:
					st.toggle( label='Use Document Context',
						value=bool( st.session_state.get( 'use_document_context', False ) ),
						key='use_document_context' )
				with ctx_c4:
					st.toggle( label='Use Semantic Context',
						value=bool( st.session_state.get( 'use_semantic', False ) ), key='use_semantic' )
				with ctx_c5:
					st.toggle( label='Use Grounding',
						value=bool( st.session_state.get( 'is_grounded', False ) ), key='is_grounded' )
				if st.button( label='Reset', key='context_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'context_window', 'use_chat_history', 'use_document_context', 'use_semantic',
							'is_grounded' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Runtime Settings', icon='⚙️', expanded=False ):
				run_c1, run_c2, run_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True, gap='medium' )
				with run_c1:
					st.slider( label='CPU Threads', min_value=1, max_value=max( 1, cfg.CORES ), step=1,
						key='cpu_threads', help=cfg.CPU_CORES )
				with run_c2:
					st.slider( label='Batch Size', min_value=64, max_value=2048, step=64, key='batch_size' )
				with run_c3:
					st.slider( label='Micro Batch Size', min_value=32, max_value=1024, step=32,
						key='micro_batch_size' )
				if st.button( label='Reset', key='runtime_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'cpu_threads', 'batch_size', 'micro_batch_size' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

		# ------------------------------------------------------------------
		# Expander — System Instructions
		# ------------------------------------------------------------------
		with st.expander( label='System Instructions', icon='🖥️', expanded=False,
				width='stretch' ):
			render_system_instructions_controls( workflow='Text Generation', include_preset=True, include_preview=True )

		st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
		
		for r, c in st.session_state.messages:
			with st.chat_message( r ):
				st.markdown( c )
		
		user_input = st.chat_input( 'Ask Bro…' )
		if user_input:
			st.session_state[ 'last_preview_input' ] = str( user_input )
			save_message( 'user', user_input )
			st.session_state.messages.append( ('user', user_input) )
			with st.chat_message( 'user' ):
				st.markdown( user_input )
			
			with st.chat_message( 'assistant' ):
				out = st.empty( )
				buf = run_llm_turn( user_input=user_input,
					temperature=float( st.session_state.get( 'temperature', 0.0 ) ),
					top_p=float( st.session_state.get( 'top_percent', 0.95 ) ),
					repeat_penalty=float( st.session_state.get( 'repeat_penalty', 1.1 ) ),
					max_tokens=int( st.session_state.get( 'max_tokens', 1024 ) ) or 1024,
					stream=True, output=out )
			
			save_message( 'assistant', buf )
			st.session_state.messages.append( ('assistant', buf) )
		
		if st.button( '🧹 Clear Chat' ):
			clear_history( )
			st.session_state.messages = [ ]
			st.rerun( )


# ==============================================================================
# IMAGE TO TEXT MODE
# ==============================================================================
elif mode == 'Image to Text':
	left, center, right = st.columns( [ 0.025, 0.95, 0.025 ] )
	with center:
		st.subheader( '🖼️ Image to Text' )
		st.divider( )
		if not vision_runtime_available( ):
			st.warning( 'Image-to-Text requires an explicitly configured compatible Gemma 3 mmproj GGUF file. Configure '
				'cfg.MMPROJ_PATH / cfg.MM_PROJ_PATH, BRO_MMPROJ_PATH, GEMMA_MMPROJ_PATH, or place '
				'cfg.MMPROJ_PATH, cfg.MM_PROJ_PATH, BRO_MMPROJ_PATH, or GEMMA_MMPROJ_PATH.' )

		with st.expander( label='Mind Controls', icon='🧠', expanded=False ):
			with st.expander( label='Vision Controls', icon='👁️', expanded=False ):
				vis_c1, vis_c2, vis_c3, vis_c4, vis_c5, vis_c6 = st.columns(
					[ 1 / 6, 1 / 6, 1 / 6, 1 / 6, 1 / 6, 1 / 6 ], border=True, gap='medium' )
				with vis_c1:
					st.selectbox( label='Vision Task', options=VISION_TASK_OPTIONS, key='vision_task' )
				with vis_c2:
					st.selectbox( label='Image Detail', options=VISION_DETAIL_OPTIONS, key='vision_detail' )
				with vis_c3:
					st.selectbox( label='Response Format', options=get_response_formats( ),
						key='vision_response_format' )
				with vis_c4:
					st.selectbox( label='Response Language', options=get_spoken_languages( False ),
						key='vision_response_language' )
				with vis_c5:
					st.toggle( label='Preserve Layout',
						value=bool( st.session_state.get( 'vision_preserve_layout', True ) ),
						key='vision_preserve_layout' )
				with vis_c6:
					st.toggle( label='Include Visible Text',
						value=bool( st.session_state.get( 'vision_include_visible_text', True ) ),
						key='vision_include_visible_text' )
				if st.button( label='Reset', key='vision_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'vision_task', 'vision_detail', 'vision_response_format',
							'vision_response_language', 'vision_preserve_layout', 'vision_include_visible_text' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Inference Settings', icon='🎚️', expanded=False ):
				inf_c1, inf_c2, inf_c3, inf_c4, inf_c5 = st.columns(
					[ 0.2, 0.2, 0.2, 0.2, 0.2 ], border=True, gap='medium' )
				with inf_c1:
					st.slider( label='Temperature', min_value=0.0, max_value=1.0, key='temperature',
						help=cfg.TEMPERATURE )
				with inf_c2:
					st.slider( label='Top-P', min_value=0.0, max_value=1.0, step=0.01,
						key='top_percent', help=cfg.TOP_P )
				with inf_c3:
					st.slider( label='Top-K', min_value=0, max_value=50, step=1, key='top_k', help=cfg.TOP_K )
				with inf_c4:
					st.slider( label='Repeat Penalty', min_value=0.0, max_value=2.0, step=0.05,
						key='repeat_penalty', help=cfg.REPEAT_PENALTY )
				with inf_c5:
					st.slider( label='Repeat Window', min_value=0, max_value=1024, step=16,
						key='repeat_window', help=cfg.REPEAT_WINDOW )
				inf_c6, inf_c7, inf_c8, inf_c9 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with inf_c6:
					st.slider( label='Presence Penalty', min_value=0.0, max_value=2.0, step=0.05,
						key='presense_penalty', help=cfg.PRESENCE_PENALTY )
				with inf_c7:
					st.slider( label='Frequency Penalty', min_value=0.0, max_value=2.0, step=0.05,
						key='frequency_penalty', help=cfg.FREQUENCY_PENALTY )
				with inf_c8:
					st.slider( label='Random Seed', min_value=0, max_value=4096, step=1,
						key='random_seed', help=cfg.SEED )
				with inf_c9:
					st.slider( label='Max Tokens', min_value=0, max_value=8192, step=128,
						key='max_tokens', help=cfg.MAX_TOKENS )
				if st.button( label='Reset', key='vision_inference_reset', width='stretch', icon='🔄' ):
					for key in [ 'temperature', 'top_percent', 'top_k', 'repeat_penalty', 'repeat_window',
							'presense_penalty', 'frequency_penalty', 'random_seed', 'max_tokens' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Runtime Settings', icon='⚙️', expanded=False ):
				run_c1, run_c2, run_c3, run_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with run_c1:
					st.slider( label='Context Window', min_value=0, max_value=131072, step=512,
						key='context_window', help=cfg.CONTEXT_WINDOW )
				with run_c2:
					st.slider( label='CPU Threads', min_value=1, max_value=max( 1, cfg.CORES ), step=1,
						key='cpu_threads', help=cfg.CPU_CORES )
				with run_c3:
					st.slider( label='Batch Size', min_value=64, max_value=2048, step=64, key='batch_size' )
				with run_c4:
					st.slider( label='Micro Batch Size', min_value=32, max_value=1024, step=32,
						key='micro_batch_size' )
				if st.button( label='Reset', key='vision_runtime_reset', width='stretch', icon='🔄' ):
					for key in [ 'context_window', 'cpu_threads', 'batch_size', 'micro_batch_size' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

		with st.expander( label='System Instructions', icon='🖥️', expanded=False, width='stretch' ):
			render_system_instructions_controls( workflow='Image to Text', include_preset=False, include_preview=False )

		uploaded_images = st.file_uploader( label='Upload Image(s)',
			type=[ 'png', 'jpg', 'jpeg', 'webp' ], accept_multiple_files=True, key='vision_uploads' )
		image_payloads: List[ Dict[ str, Any ] ] = [ ]
		if uploaded_images:
			preview_columns = st.columns( min( 4, len( uploaded_images ) ) )
			for image_index, uploaded_image in enumerate( uploaded_images ):
				image_bytes = uploaded_image.getvalue( )
				mime_type = str( getattr( uploaded_image, 'type', '' ) or 'image/png' )
				image_payloads.append( { 'name': uploaded_image.name, 'bytes': image_bytes,
					'mime_type': mime_type } )
				with preview_columns[ image_index % len( preview_columns ) ]:
					st.image( image_bytes, caption=uploaded_image.name, use_container_width=True )
		vision_question = st.text_area( label='Image Request', height=100,
			placeholder='Optional question or instruction for the selected image task.',
			key='vision_user_request' )
		if st.button( label='Run Image Analysis', icon='🖼️', width='stretch',
				key='vision_run_analysis', disabled=not bool( image_payloads ) ):
			with st.chat_message( 'assistant' ):
				vision_output = st.empty( )
				response = run_vision_turn( image_payloads=image_payloads, user_input=vision_question,
					stream=True, output=vision_output, show_errors=True )
				st.session_state[ 'vision_last_response' ] = response

# ==============================================================================
# RETRIEVAL AUGMENTATION
# ==============================================================================
elif mode == 'Document Q&A':
	messages = st.session_state.get( 'messages', [ ] )
	uploaded = st.session_state.get( 'uploaded', [ ] )
	active_docs = st.session_state.get( 'active_docs', [ ] )
	doc_bytes = st.session_state.get( 'doc_bytes', { } )
	max_tokens = st.session_state.get( 'max_tokens', 0 )
	top_percent = st.session_state.get( 'top_percent', 0.0 )
	top_k = st.session_state.get( 'top_k', 0 )
	temperature = st.session_state.get( 'temperature', 0.0 )
	frequency_penalty = st.session_state.get( 'frequency_penalty', 0.0 )
	presense_penalty = st.session_state.get( 'presense_penalty', 0.0 )
	repeat_penalty = st.session_state.get( 'repeat_penalty', 0.0 )
	repeat_window = st.session_state.get( 'repeat_window', 0.0 )
	cpu_threads = st.session_state.get( 'cpu_threads', cfg.CORES )
	context_window = st.session_state.get( 'context_window', cfg.DEFAULT_CTX )
	
	left, center, right = st.columns( [ 0.025, 0.95, 0.025 ] )
	with center:
		st.subheader( '📚 Retrieval Augementation', help=cfg.RETRIEVAL_AUGMENTATION )
		st.divider( )
		
		# ------------------------------------------------------------------
		# Expander — Mind Controls
		# ------------------------------------------------------------------
		with st.expander( label='Mind Controls', icon='🧠', expanded=False ):
			with st.expander( label='Retrieval Controls', icon='🧲', expanded=False ):
				ret_c1, ret_c2, ret_c3, ret_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with ret_c1:
					st.slider( label='Chunks to Retrieve', min_value=1, max_value=20, step=1,
						key='retrieval_k' )
				with ret_c2:
					st.slider( label='Chunk Size', min_value=256, max_value=4000, step=64,
						key='retrieval_chunk_size' )
				with ret_c3:
					st.slider( label='Chunk Overlap', min_value=0, max_value=1000, step=25,
						key='retrieval_chunk_overlap' )
				with ret_c4:
					st.toggle( label='Show Retrieved Chunks',
						value=bool( st.session_state.get( 'show_retrieved_chunks', True ) ),
						key='show_retrieved_chunks' )
				if st.button( label='Reset', key='doc_retrieval_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'retrieval_k', 'retrieval_chunk_size', 'retrieval_chunk_overlap',
							'show_retrieved_chunks' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Grounding Controls', icon='🛡️', expanded=False ):
				ground_c1, ground_c2, ground_c3 = st.columns(
					[ 0.34, 0.33, 0.33 ], border=True, gap='medium' )
				with ground_c1:
					st.toggle( label='Require Grounding',
						value=bool( st.session_state.get( 'require_grounding', True ) ), key='require_grounding' )
				with ground_c2:
					st.toggle( label='Answer From Excerpts Only',
						value=bool( st.session_state.get( 'answer_from_excerpts_only', True ) ),
						key='answer_from_excerpts_only' )
				with ground_c3:
					st.selectbox( label='If Evidence Is Insufficient', options=GROUNDING_FAILURE_OPTIONS,
						key='grounding_failure_behavior' )
				if st.button( label='Reset', key='doc_grounding_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'require_grounding', 'answer_from_excerpts_only',
							'grounding_failure_behavior' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Retrieval Backend', icon='🧮', expanded=False ):
				back_c1, back_c2, back_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True, gap='medium' )
				with back_c1:
					st.selectbox( label='Backend', options=RETRIEVAL_BACKEND_OPTIONS, key='retrieval_backend' )
				with back_c2:
					st.toggle( label='Fallback Cosine Search',
						value=bool( st.session_state.get( 'allow_similarity_fallback', True ) ),
						key='allow_similarity_fallback' )
				with back_c3:
					st.toggle( label='Rebuild Index Each Query',
						value=bool( st.session_state.get( 'docqna_rebuild_each_query', False ) ),
						key='docqna_rebuild_each_query' )
				st.session_state[ 'prefer_sqlite_vec' ] = st.session_state.get( 'retrieval_backend' ) != 'Cosine Similarity'
				if st.button( label='Reset', key='doc_backend_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'retrieval_backend', 'allow_similarity_fallback', 'docqna_rebuild_each_query' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Document Actions', icon='🗂️', expanded=False ):
				action_c1, action_c2, action_c3 = st.columns( [ 0.5, 0.25, 0.25 ], border=True, gap='medium' )
				with action_c1:
					st.selectbox( label='Action', options=[ 'Answer Question', 'Summarize Active Document',
						'Extract Key Points', 'Generate Outline', 'Extract Entities', 'Extract Tables',
						'Compare Active Documents', 'Classify Document', 'Find Evidence',
						'Generate Executive Summary', 'Extract Dates', 'Extract Organizations',
						'Extract Requirements', 'Extract Action Items', 'Identify Contradictions',
						'Identify Missing Information' ], key='docqna_action' )
				with action_c2:
					st.selectbox( label='Action Detail', options=DOC_ACTION_DETAIL_OPTIONS,
						key='docqna_action_detail' )
				with action_c3:
					st.markdown( '<br>', unsafe_allow_html=True )
					run_action = st.button( 'Run Action', key='doc_run_action', width='stretch' )
				if run_action:
					action_name = str( st.session_state.get( 'docqna_action', 'Answer Question' ) )
					action_detail = str( st.session_state.get( 'docqna_action_detail', 'Standard' ) )
					action_prompts = {
						'Summarize Active Document': 'Summarize the active document set clearly and faithfully.',
						'Extract Key Points': 'Extract the key points from the active document set.',
						'Generate Outline': 'Generate an outline of the active document set.',
						'Extract Entities': 'Extract named entities, dates, organizations, and references from the active document set.',
						'Extract Tables': 'Describe the tabular or structured information visible in the active document set.',
						'Compare Active Documents': 'Compare the active documents and explain major agreements, differences, and gaps.',
						'Classify Document': 'Classify the active document set using supported evidence.',
						'Find Evidence': 'Find excerpts that directly support or contradict the requested proposition.',
						'Generate Executive Summary': 'Generate an executive summary of the active document set.',
						'Extract Dates': 'Extract important dates and their associated events or obligations.',
						'Extract Organizations': 'Extract named organizations and their supported roles.',
						'Extract Requirements': 'Extract explicit requirements, constraints, and acceptance criteria.',
						'Extract Action Items': 'Extract action items and responsible parties when supported.',
						'Identify Contradictions': 'Identify material contradictions or inconsistencies across the active documents.',
						'Identify Missing Information': 'Identify information required by the request that is absent from the active documents.' }
					if action_name != 'Answer Question':
						action_prompt = action_prompts.get( action_name, 'Summarize the active document set.' )
						action_prompt += f'\nDetail Level: {action_detail}.'
						with st.chat_message( 'assistant' ):
							out = st.empty( )
							response = run_direct_llm_turn( system_instruction=get_effective_system_instructions( ), user_input=build_docqna_input( user_query=action_prompt,
								k=int( st.session_state.get( 'retrieval_k', 6 ) ) ),
								temperature=float( st.session_state.get( 'temperature', 0.0 ) ),
								top_p=float( st.session_state.get( 'top_percent', 0.95 ) ),
								repeat_penalty=float( st.session_state.get( 'repeat_penalty', 1.1 ) ),
								max_tokens=int( st.session_state.get( 'max_tokens', 1024 ) ) or 1024,
								stream=True, output=out, response_format=str( st.session_state.get( 'response_format', 'Markdown' ) ) )
						save_message( 'assistant', response )
						st.session_state.messages.append( ('assistant', response) )
				if st.button( label='Reset', key='doc_action_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'docqna_action', 'docqna_action_detail' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Document Parsing', icon='📄', expanded=False ):
				parse_c1, parse_c2, parse_c3, parse_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with parse_c1:
					st.toggle( label='Enable OCR', value=bool( st.session_state.get( 'ocr_enabled', False ) ),
						key='ocr_enabled' )
				with parse_c2:
					st.toggle( label='Prefer Native PDF Text',
						value=bool( st.session_state.get( 'prefer_native_pdf_text', True ) ),
						key='prefer_native_pdf_text' )
				with parse_c3:
					st.toggle( label='Include Page Markers',
						value=bool( st.session_state.get( 'include_page_markers', False ) ),
						key='include_page_markers' )
				with parse_c4:
					st.selectbox( label='OCR Page Limit', options=OCR_PAGE_LIMIT_OPTIONS, key='ocr_page_limit' )
				if st.button( label='Reset', key='doc_parsing_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'ocr_enabled', 'prefer_native_pdf_text', 'include_page_markers', 'ocr_page_limit' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.session_state[ 'docqna_ocr_cache' ] = { }
					st.rerun( )


			with st.expander( label='Diagnostics', icon='🔎', expanded=False ):
				diag_c1, diag_c2, diag_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True, gap='medium' )
				with diag_c1:
					st.toggle( label='Show Diagnostics',
						value=bool( st.session_state.get( 'show_docqna_diagnostics', False ) ),
						key='show_docqna_diagnostics' )
				with diag_c2:
					st.toggle( label='Show OCR Status',
						value=bool( st.session_state.get( 'show_ocr_status', True ) ), key='show_ocr_status' )
				with diag_c3:
					st.toggle( label='Show Runtime Metadata',
						value=bool( st.session_state.get( 'show_runtime_metadata', False ) ),
						key='show_runtime_metadata' )
				if st.button( label='Reset', key='doc_diagnostics_reset', width='stretch', icon='🔄' ):
					for key in [ 'show_docqna_diagnostics', 'show_ocr_status', 'show_runtime_metadata' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Response Controls', icon='↔️', expanded=False ):
				resp_c1, resp_c2, resp_c3, resp_c4 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with resp_c1:
					st.selectbox( label='Response Format', options=get_response_formats( ), key='response_format' )
				with resp_c2:
					st.selectbox( label='Response Language', options=get_spoken_languages( False ),
						key='response_language' )
				with resp_c3:
					st.selectbox( label='Response Length', options=RESPONSE_LENGTH_OPTIONS, key='response_length' )
				with resp_c4:
					st.toggle( label='Include Headings',
						value=bool( st.session_state.get( 'response_include_headings', True ) ),
						key='response_include_headings' )
				if st.button( label='Reset', key='doc_response_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'response_format', 'response_language', 'response_length',
							'response_include_headings' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Inference Settings', icon='🎚️', expanded=False ):
				inf_c1, inf_c2, inf_c3, inf_c4, inf_c5 = st.columns(
					[ 0.2, 0.2, 0.2, 0.2, 0.2 ], border=True, gap='medium' )
				with inf_c1:
					st.slider( label='Temperature', min_value=0.0, max_value=1.0, key='temperature',
						help=cfg.TEMPERATURE )
				with inf_c2:
					st.slider( label='Top-P', min_value=0.0, max_value=1.0, step=0.01,
						key='top_percent', help=cfg.TOP_P )
				with inf_c3:
					st.slider( label='Top-K', min_value=0, max_value=50, step=1, key='top_k', help=cfg.TOP_K )
				with inf_c4:
					st.slider( label='Repeat Penalty', min_value=0.0, max_value=2.0, step=0.05,
						key='repeat_penalty', help=cfg.REPEAT_PENALTY )
				with inf_c5:
					st.slider( label='Repeat Window', min_value=0, max_value=1024, step=16,
						key='repeat_window', help=cfg.REPEAT_WINDOW )
				inf_c6, inf_c7, inf_c8, inf_c9 = st.columns(
					[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
				with inf_c6:
					st.slider( label='Presence Penalty', min_value=0.0, max_value=2.0, step=0.05,
						key='presense_penalty', help=cfg.PRESENCE_PENALTY )
				with inf_c7:
					st.slider( label='Frequency Penalty', min_value=0.0, max_value=2.0, step=0.05,
						key='frequency_penalty', help=cfg.FREQUENCY_PENALTY )
				with inf_c8:
					st.slider( label='Random Seed', min_value=0, max_value=4096, step=1,
						key='random_seed', help=cfg.SEED )
				with inf_c9:
					st.slider( label='Max Tokens', min_value=0, max_value=8192, step=128,
						key='max_tokens', help=cfg.MAX_TOKENS )
				if st.button( label='Reset', key='doc_probability_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'temperature', 'top_percent', 'top_k', 'repeat_penalty', 'repeat_window',
							'presense_penalty', 'frequency_penalty', 'random_seed', 'max_tokens' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Context Controls', icon='🎛️', expanded=False ):
				ctx_c1, ctx_c2, ctx_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True, gap='medium' )
				with ctx_c1:
					st.slider( label='Context Window', min_value=0, max_value=131072, step=512,
						key='context_window', help=cfg.CONTEXT_WINDOW )
				with ctx_c2:
					st.toggle( label='Include Semantic Context',
						value=bool( st.session_state.get( 'docqna_include_semantic_context', True ) ),
						key='docqna_include_semantic_context' )
				with ctx_c3:
					st.selectbox( label='Context Order', options=DOC_CONTEXT_ORDER_OPTIONS,
						key='docqna_context_order' )
				if st.button( label='Reset', key='doc_context_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'context_window', 'docqna_include_semantic_context', 'docqna_context_order' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

			with st.expander( label='Runtime Settings', icon='⚙️', expanded=False ):
				run_c1, run_c2, run_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True, gap='medium' )
				with run_c1:
					st.slider( label='CPU Threads', min_value=1, max_value=max( 1, cfg.CORES ), step=1,
						key='cpu_threads', help=cfg.CPU_CORES )
				with run_c2:
					st.slider( label='Batch Size', min_value=64, max_value=2048, step=64, key='batch_size' )
				with run_c3:
					st.slider( label='Micro Batch Size', min_value=32, max_value=1024, step=32,
						key='micro_batch_size' )
				if st.button( label='Reset', key='doc_runtime_controls_reset', width='stretch', icon='🔄' ):
					for key in [ 'cpu_threads', 'batch_size', 'micro_batch_size' ]:
						if key in st.session_state:
							del st.session_state[ key ]
					st.rerun( )

		# ------------------------------------------------------------------
		# Expander — System Instructions
		# ------------------------------------------------------------------
		with st.expander( label='System Instructions', icon='🖥️', expanded=False, width='stretch' ):
			render_system_instructions_controls( workflow='Document Q&A', include_preset=False, include_preview=False )

		# ------------------------------------------------------------------
		# Document Selection UI
		# ------------------------------------------------------------------
		with st.expander( label='Document Loader', icon='📥', expanded=False,
				width='stretch' ):
			doc_left, doc_right = st.columns( [ 0.5, 0.5 ], gap='medium', border=True )
			with doc_left:
				st.radio( label='Document Source', options=[ 'uploadlocal' ],
					index=0, horizontal=True, key='doc_source' )
				
				uploaded = st.file_uploader( label='Upload document(s) (PDF, TXT, DOCX)',
					type=[ 'pdf', 'txt', 'docx' ], accept_multiple_files=True,
					label_visibility='visible' )
				
				if uploaded is not None and isinstance( uploaded, list ) and len( uploaded ) > 0:
					st.session_state.uploaded = uploaded
					names: List[ str ] = [ f.name for f in uploaded if getattr( f, 'name', None ) ]
					st.session_state.active_docs = names
					if 'doc_bytes' not in st.session_state or not isinstance(
							st.session_state.doc_bytes, dict ):
						st.session_state.doc_bytes = { }
					
					for f in uploaded:
						try:
							if getattr( f, 'name', None ):
								st.session_state.doc_bytes[ f.name ] = f.getvalue( )
						except Exception as e:
							exception = Error( e )
							exception.module = 'app'
							exception.cause = 'Streamlit UI'
							exception.method = 'streamlit_ui_block() -> None'
							Logger( ).write( exception )
							continue
					
					st.session_state[ 'docqna_inventory_rows' ] = build_docqna_inventory( )
				else:
					st.info( 'Load a document.' )
				
				if st.session_state.get( 'active_docs' ):
					st.multiselect( label='Active Documents',
						options=[ f.name for f in st.session_state.get( 'uploaded', [ ] ) ],
						default=st.session_state.get( 'active_docs', [ ] ),
						key='active_docs' )
				
				unload = st.button( label='Unload Document(s)', width='stretch' )
				if unload:
					st.session_state.uploaded = [ ]
					st.session_state.active_docs = [ ]
					st.session_state.doc_bytes = { }
					st.session_state[ 'docqna_inventory_rows' ] = [ ]
					st.session_state[ 'docqna_fingerprint' ] = ''
					st.session_state[ 'docqna_chunk_count' ] = 0
					st.session_state[ 'docqna_fallback_rows' ] = [ ]
					st.session_state[ 'docqna_last_retrieval' ] = [ ]
					st.rerun( )
				
				if bool( st.session_state.get( 'show_docqna_diagnostics', False ) ):
					st.caption(
						f'Chunk Size: {int( st.session_state.get( 'retrieval_chunk_size', 1200 ) )} '
						f'| Chunk Overlap: {int( st.session_state.get( 'retrieval_chunk_overlap', 200 ) )} '
						f'| Index Ready: {bool( st.session_state.get( 'docqna_vec_ready', False ) )} '
						f'| Chunk Count: {int( st.session_state.get( 'docqna_chunk_count', 0 ) )}' )
				if bool( st.session_state.get( 'show_ocr_status', True ) ):
					st.caption( f'OCR Enabled: {bool( st.session_state.get( "ocr_enabled", False ) )} '
						f'| Vision Runtime: {vision_runtime_available( )} '
						f'| OCR Page Limit: {str( st.session_state.get( "ocr_page_limit", "5 Pages" ) )}' )
				if bool( st.session_state.get( 'show_runtime_metadata', False ) ):
					st.caption( f'Context Window: {int( st.session_state.get( "context_window", cfg.DEFAULT_CTX ) or cfg.DEFAULT_CTX )} '
						f'| CPU Threads: {int( st.session_state.get( "cpu_threads", cfg.CORES ) or cfg.CORES )} '
						f'| Retrieval Backend: {str( st.session_state.get( "retrieval_backend", "Automatic" ) )}' )
			
			with doc_right:
				if st.session_state.get( 'active_docs' ):
					preview_name = st.session_state.active_docs[ 0 ]
					file_bytes = st.session_state.doc_bytes.get( preview_name )
					if file_bytes and str( preview_name ).lower( ).endswith( '.pdf' ):
						st.pdf( file_bytes, height=420 )
					elif file_bytes:
						preview_text = extract_text( file_bytes, preview_name )
						st.text_area( label=f'Preview: {preview_name}', value=preview_text[ :4000 ],
							height=420, disabled=True )
					else:
						st.info( 'Document loaded but preview unavailable.' )
				else:
					st.info( 'No document loaded.' )
			
			if st.session_state.get( 'docqna_inventory_rows' ):
				st.markdown( '### Active Document Inventory' )
				st.dataframe( pd.DataFrame( st.session_state.get( 'docqna_inventory_rows', [ ] ) ),
					use_container_width=True )
		
		# ------------------------------------------------------------------
		# Chat History Render
		# ------------------------------------------------------------------
		if 'messages' not in st.session_state or not isinstance( st.session_state.messages, list ):
			st.session_state.messages = [ ]
		
		for msg in st.session_state.messages:
			role = ''
			content = ''
			
			if isinstance( msg, dict ):
				role = str( msg.get( 'role', '' ) or '' ).strip( )
				content = msg.get( 'content', '' )
			else:
				if isinstance( msg, tuple ) or isinstance( msg, list ):
					if len( msg ) == 2:
						role = str( msg[ 0 ] or '' ).strip( )
						content = msg[ 1 ]
					else:
						role = ''
						content = ''
				else:
					role = ''
					content = ''
			
			if role not in ('user', 'assistant', 'system'):
				continue
			
			if content is None:
				content = ''
			elif not isinstance( content, str ):
				content = str( content )
			
			with st.chat_message( role ):
				st.markdown( content )
		
		st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
		
		# ------------------------------------------------------------------
		# Chat Input
		# ------------------------------------------------------------------
		user_input = st.chat_input( 'Ask a question about the document' )
		if user_input and isinstance( user_input, str ) and user_input.strip( ):
			user_input = user_input.strip( )
			
			if 'messages' not in st.session_state or not isinstance(
					st.session_state.messages, list ):
				st.session_state.messages = [ ]
			
			save_message( 'user', user_input )
			st.session_state.messages.append( ('user', user_input) )
			
			with st.chat_message( 'user' ):
				st.markdown( user_input )
			
			doc_user_input = build_docqna_input( user_query=user_input,
				k=int( st.session_state.get( 'retrieval_k', 6 ) ) )
			
			if not doc_user_input or not isinstance( doc_user_input,
					str ) or not doc_user_input.strip( ):
				doc_user_input = user_input
			
			with st.chat_message( 'assistant' ):
				out = st.empty( )
				response = run_direct_llm_turn( system_instruction=get_effective_system_instructions( ), user_input=doc_user_input,
					temperature=float( st.session_state.get( 'temperature', 0.0 ) ),
					top_p=float( st.session_state.get( 'top_percent', 0.95 ) ),
					repeat_penalty=float( st.session_state.get( 'repeat_penalty', 1.1 ) ),
					max_tokens=int( st.session_state.get( 'max_tokens', 1024 ) ) or 1024,
					stream=True, output=out, response_format=str( st.session_state.get( 'response_format', 'Markdown' ) ) )
			
			if response is None:
				response = ''
			elif not isinstance( response, str ):
				response = str( response )
			
			response = response.strip( )
			save_message( 'assistant', response )
			st.session_state.messages.append( ('assistant', response) )
			if bool( st.session_state.get( 'show_retrieved_chunks', True ) ):
				hits = st.session_state.get( 'docqna_last_retrieval', [ ] )
				if hits:
					with st.expander( 'Retrieved Chunks', expanded=False ):
						for idx, hit in enumerate( hits, start=1 ):
							doc_name = str( hit[ 0 ] )
							chunk_text_value = str( hit[ 1 ] )
							score_value = hit[ 2 ]
							
							st.markdown( f'**{idx}. {doc_name}**' )
							st.caption( f'Score / Distance: {score_value}' )
							st.text_area( label=f'Chunk {idx}', value=chunk_text_value,
								height=140, disabled=True, key=f'doc_hit_{idx}' )
		
		if st.button( '🧹 Clear Chat', key='doc_clear_chat' ):
			clear_history( )
			st.session_state.messages = [ ]
			st.rerun( )

# ==============================================================================
# SEMANTIC SEARCH
# ==============================================================================
elif mode == 'Semantic Search':
	left, center, right = st.columns( [ 0.025, 0.95, 0.025 ] )
	with center:
		st.subheader( '🔍 Semantic Search', help=cfg.SEMANTIC_SEARCH )
		st.divider( )
		
		with st.expander( label='Index Builder', icon='🧱', expanded=False ):
			idx_c1, idx_c2, idx_c3, idx_c4, idx_c5 = st.columns(
				[ 0.2, 0.2, 0.2, 0.2, 0.2 ], border=True, gap='medium' )
			with idx_c1:
				st.slider( label='Chunk Size', min_value=256, max_value=4000, step=64,
					key='semantic_chunk_size' )
			with idx_c2:
				st.slider( label='Chunk Overlap', min_value=0, max_value=1000, step=25,
					key='semantic_chunk_overlap' )
			with idx_c3:
				st.toggle( label='Clear Existing Index',
					value=bool( st.session_state.get( 'semantic_clear_existing', True ) ),
					key='semantic_clear_existing' )
			with idx_c4:
				st.toggle( label='Append to Existing Index',
					value=bool( st.session_state.get( 'semantic_append_existing', False ) ),
					key='semantic_append_existing' )
			with idx_c5:
				st.toggle( label='Show Embedding Diagnostics',
					value=bool( st.session_state.get( 'semantic_show_diagnostics', True ) ),
					key='semantic_show_diagnostics' )
			
			semantic_files = st.file_uploader( label='Upload for embedding',
				accept_multiple_files=True, type=[ 'pdf', 'txt', 'docx' ],
				key='semantic_file_uploader' )
			
			if st.button( 'Build Index', key='semantic_build_index', width='stretch' ):
				if semantic_files:
					result = build_semantic_index( semantic_files )
					if bool( result.get( 'success', False ) ):
						st.success( str( result.get( 'message', '' ) ) )
					else:
						st.error( str( result.get( 'message', 'Index build failed.' ) ) )
				else:
					st.info( 'Upload one or more files before building the index.' )
			
			if st.button( label='Reset', key='semantic_index_builder_reset', width='stretch', icon='🔄' ):
				for key in [ 'semantic_chunk_size', 'semantic_chunk_overlap', 'semantic_clear_existing',
						'semantic_append_existing', 'semantic_show_diagnostics' ]:
					if key in st.session_state:
						del st.session_state[ key ]
				st.rerun( )
			
			if bool( st.session_state.get( 'semantic_show_diagnostics', True ) ):
				diag_c1, diag_c2, diag_c3 = st.columns( [ 0.33, 0.33, 0.34 ], border=True )
				with diag_c1:
					st.metric( 'Indexed Documents',
						int( st.session_state.get( 'semantic_index_doc_count', 0 ) ) )
				
				with diag_c2:
					st.metric( 'Indexed Chunks',
						int( st.session_state.get( 'semantic_index_chunk_count', 0 ) ) )
				with diag_c3:
					st.metric( 'Vector Dimension',
						int( st.session_state.get( 'semantic_index_dim', 0 ) ) )
		
		with st.expander( label='Semantic Query', icon='🧠', expanded=False ):
			query_c1, query_c2, query_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True,
				gap='medium' )
			
			with query_c1:
				st.slider( label='Top K', min_value=1, max_value=25, step=1,
					key='semantic_top_k' )
			
			with query_c2:
				st.slider( label='Minimum Similarity', min_value=0.0, max_value=1.0, step=0.01,
					key='semantic_min_similarity' )
			
			with query_c3:
				st.toggle( label='Show Embedding Diagnostics',
					value=bool( st.session_state.get( 'semantic_show_diagnostics', True ) ),
					key='semantic_query_show_diagnostics' )
			
			semantic_query = st.text_area( label='Semantic Query', height=120,
				key='semantic_query_text' )
			
			if st.button( 'Run Semantic Search', key='semantic_run_query', width='stretch' ):
				rows = query_semantic_index( semantic_query )
				if len( rows ) == 0:
					st.info( 'No semantic matches found.' )
			
			if st.button( label='Reset', key='semantic_query_reset', width='stretch', icon='🔄' ):
				for key in [ 'semantic_top_k', 'semantic_min_similarity', 'semantic_query_show_diagnostics',
						'semantic_query_text' ]:
					if key in st.session_state:
						del st.session_state[ key ]
				st.session_state[ 'semantic_result_rows' ] = [ ]
				st.session_state[ 'semantic_selected_rows' ] = [ ]
				st.rerun( )
			
			result_rows = st.session_state.get( 'semantic_result_rows', [ ] )
			if isinstance( result_rows, list ) and len( result_rows ) > 0:
				edited_rows = st.data_editor( result_rows, hide_index=True,
					use_container_width=True,
					key='semantic_results_editor' )
				
				selected_rows = extract_selected_rows( edited_rows )
				st.session_state[ 'semantic_selected_rows' ] = selected_rows
				if len( selected_rows ) > 0:
					st.caption( f'Selected Chunks: {len( selected_rows )}' )
		
		with st.expander( label='Actions', icon='🔀', expanded=False ):
			act_c1, act_c2, act_c3 = st.columns( [ 0.34, 0.33, 0.33 ] )
			
			with act_c1:
				if st.button( 'Send Selected Chunks to Text Generation', width='stretch' ):
					send_text_chunks( )
					st.success( 'Selected chunks added to shared Text Generation context.' )
			
			with act_c2:
				if st.button( 'Send Selected Chunks to Document Q&A', width='stretch' ):
					send_docqna_chunks( )
					st.success( 'Selected chunks added to the shared Document Q&A context buffer.' )
			
			with act_c3:
				if st.button( 'Save Selected Chunks as Prompt Context', width='stretch' ):
					context_text = create_semantic_context( )
					if context_text:
						existing_docs = st.session_state.get( 'basic_docs', [ ] )
						if not isinstance( existing_docs, list ):
							existing_docs = [ ]
						existing_docs.append( context_text )
						st.session_state[ 'basic_docs' ] = existing_docs
						st.success( 'Selected chunks saved to shared prompt context.' )
					else:
						st.info( 'Select one or more chunks first.' )
			
			selected_rows = st.session_state.get( 'semantic_selected_rows', [ ] )
			if isinstance( selected_rows, list ) and len( selected_rows ) > 0:
				st.markdown( '### Selected Semantic Context Preview' )
				st.text_area( label='Selected Context', value=create_semantic_context( ),
					height=220, disabled=True )
		
		with st.expander( label='Index Maintenance', icon='🛠️', expanded=False ):
			maint_c1, maint_c2, maint_c3 = st.columns( [ 0.34, 0.33, 0.33 ] )
			
			with maint_c1:
				if st.button( 'Delete Index', width='stretch' ):
					clear_semantic_index( )
					st.success( 'Semantic index deleted.' )
			
			with maint_c2:
				if st.button( 'Recompute Diagnostics', width='stretch' ):
					rows = decode_embedding_rows( )
					st.session_state[ 'semantic_index_chunk_count' ] = len( rows )
					if len( rows ) > 0:
						st.session_state[ 'semantic_index_dim' ] = int( rows[ 0 ][ 2 ].shape[ 0 ] )
					else:
						st.session_state[ 'semantic_index_dim' ] = 0
					st.success( 'Diagnostics refreshed.' )
			
			with maint_c3:
				if st.button( 'Clear Query Results', width='stretch' ):
					st.session_state[ 'semantic_result_rows' ] = [ ]
					st.session_state[ 'semantic_selected_rows' ] = [ ]
					st.session_state[ 'semantic_last_query' ] = ''
					st.success( 'Query results cleared.' )
			
			if bool( st.session_state.get( 'semantic_show_diagnostics', True ) ):
				st.caption(
					f'Last Query: {str( st.session_state.get( "semantic_last_query", "" ) )} '
					f'| Uploaded Sources: {len( st.session_state.get( "semantic_uploaded_names", [ ] ) )}' )

# ==============================================================================
# PROMPT ENGINEERING MODE
# ==============================================================================
elif mode == 'Prompt Engineering':
	import math

	TABLE = 'Prompts'
	PAGE_SIZE = 10
	st.session_state.setdefault( 'pe_cascade_enabled', False )
	left, center, right = st.columns( [ 0.025, 0.95, 0.025 ] )
	with center:
		st.subheader( '📝 Prompt Engineering', help=cfg.PROMPT_ENGINEERING )
		st.divider( )
		
		st.checkbox( 'Cascade selection into shared System Instructions and task settings',
			key='pe_cascade_enabled' )
		
		# ------------------------------------------------------------------
		# Session state
		# ------------------------------------------------------------------
		st.session_state.setdefault( 'pe_page', 1 )
		st.session_state.setdefault( 'pe_search', '' )
		st.session_state.setdefault( 'pe_sort_col', 'ID' )
		st.session_state.setdefault( 'pe_sort_dir', 'ASC' )
		st.session_state.setdefault( 'pe_selected_id', None )
		st.session_state.setdefault( 'pe_caption', '' )
		st.session_state.setdefault( 'pe_name', '' )
		st.session_state.setdefault( 'pe_category', '' )
		st.session_state.setdefault( 'pe_text', '' )
		st.session_state.setdefault( 'pe_language', 'English' )
		st.session_state.setdefault( 'pe_language_edit', 'English' )
		st.session_state.setdefault( 'pe_task_type_edit', 'Chat' )
		
		# ------------------------------------------------------------------
		# DB helpers
		# ------------------------------------------------------------------
		def get_conn( ) -> sqlite3.Connection:
			"""Creates the Prompt Engineering SQLite connection.

			Purpose:
				Provides the existing local helper used by Prompt Engineering query and paging workflows.

			Returns:
				sqlite3.Connection: SQLite connection to the configured application database.
			"""
			return sqlite3.connect( cfg.DB_PATH )
		
		def reset_prompt_selection( ) -> None:
			"""Clears the Prompt Engineering edit surface.

			Purpose:
				Clears selected-record state while preserving Prompt Engineering filters and generator controls.

			Returns:
				None: This function performs its work through Streamlit session state.
			"""
			st.session_state.pe_selected_id = None
			st.session_state.pe_caption = ''
			st.session_state.pe_name = ''
			st.session_state.pe_category = ''
			st.session_state.pe_text = ''
		
		def load_prompt_for_edit( pid: int ) -> None:
			"""Loads a prompt record into the Prompt Engineering edit surface.

			Purpose:
				Loads the authoritative five-field prompt record by immutable primary key.

			Args:
				pid (int): Prompt primary key.

			Returns:
				None: This function performs its work through Streamlit session state.
			"""
			prompt_row = fetch_prompt_by_id( pid )
			if not prompt_row:
				return
			st.session_state.pe_selected_id = int( prompt_row[ 'ID' ] )
			st.session_state.pe_caption = str( prompt_row.get( 'Caption', '' ) or '' )
			st.session_state.pe_name = str( prompt_row.get( 'Name', '' ) or '' )
			st.session_state.pe_category = str(
				prompt_row.get( 'Category', '' ) or '' )
			st.session_state.pe_text = str( prompt_row.get( 'Text', '' ) or '' )
			st.session_state[ 'prompt_category' ] = st.session_state.pe_category
		
		# ------------------------------------------------------------------
		# Filters
		# ------------------------------------------------------------------
		c1, c2, c3, c4, c5 = st.columns( [ 3, 2, 2, 2, 3 ], border=True )
		available_categories = fetch_prompt_categories( cfg.DB_PATH )
		if st.session_state.get( 'pe_category' ) not in available_categories:
			st.session_state[ 'pe_category' ] = available_categories[ 0 ] if available_categories else ''
		if st.session_state.get( 'prompt_category_selection' ) not in [ 'All Categories' ] + available_categories:
			st.session_state[ 'prompt_category_selection' ] = 'All Categories'
		if st.session_state.get( 'pe_category' ) not in available_categories:
			st.session_state[ 'pe_category' ] = ''
		if st.session_state.get( 'prompt_response_format' ) not in get_response_formats( ):
			st.session_state[ 'prompt_response_format' ] = 'Markdown'
		if st.session_state.get( 'pe_language' ) not in get_spoken_languages( False ):
			st.session_state[ 'pe_language' ] = 'English'
		if st.session_state.get( 'pe_language_edit' ) not in get_spoken_languages( False ):
			st.session_state[ 'pe_language_edit' ] = 'English'
		
		with c1:
			st.text_input( 'Search (Caption / Name / Text)', key='pe_search' )
		
		with c2:
			st.selectbox( 'Category', [ 'All Categories' ] + available_categories,
				key='prompt_category_selection' )
		
		with c3:
			st.selectbox( 'Sort by', [ 'ID', 'Caption', 'Name', 'Category', 'Text' ],
				key='pe_sort_col' )
		
		with c4:
			st.selectbox( 'Direction', [ 'ASC', 'DESC' ], key='pe_sort_dir' )
		
		with c5:
			st.markdown(
				"<div style='font-size:0.95rem;font-weight:600;margin-bottom:0.25rem;'>Go to ID</div>",
				unsafe_allow_html=True )
			a1, a2, a3 = st.columns( [ 2, 1, 1 ] )
			with a1:
				jump_id = st.number_input( 'Go to ID', min_value=1, step=1,
					label_visibility='collapsed' )
			with a2:
				if st.button( 'Go' ):
					load_prompt_for_edit( int( jump_id ) )
			with a3:
				st.button( 'Clear', on_click=reset_prompt_selection )
		
		# ------------------------------------------------------------------
		# Load prompt table
		# ------------------------------------------------------------------
		where_clauses: List[ str ] = [ ]
		params: List[ Any ] = [ ]
		if st.session_state.pe_search:
			where_clauses.append( '(Caption LIKE ? OR Name LIKE ? OR Text LIKE ?)' )
			search_text = f"%{st.session_state.pe_search}%"
			params.extend( [ search_text, search_text, search_text ] )
		selected_category = str(
			st.session_state.get( 'prompt_category_selection', 'All Categories' ) or 'All Categories' )
		if selected_category != 'All Categories':
			where_clauses.append( 'Category = ?' )
			params.append( selected_category )
		where = 'WHERE ' + ' AND '.join( where_clauses ) if where_clauses else ''
		offset = (int( st.session_state.pe_page ) - 1) * PAGE_SIZE
		query = f'''SELECT ID, Caption, Name, Category, Text
			FROM {TABLE}
			{where}
			ORDER BY {st.session_state.pe_sort_col} {st.session_state.pe_sort_dir}
			LIMIT {PAGE_SIZE} OFFSET {offset}'''
		count_query = f'SELECT COUNT(*) FROM {TABLE} {where}'
		with get_conn( ) as conn:
			rows = conn.execute( query, params ).fetchall( )
			total_rows = int( conn.execute( count_query, params ).fetchone( )[ 0 ] )
		total_pages = max( 1, math.ceil( total_rows / PAGE_SIZE ) )
		if st.session_state.pe_page > total_pages:
			st.session_state.pe_page = total_pages
		
		# ------------------------------------------------------------------
		# Prompt table
		# ------------------------------------------------------------------
		table_rows: List[ Dict[ str, Any ] ] = [ ]
		for row in rows:
			table_rows.append( {
				'Selected': int( row[ 0 ] ) == st.session_state.pe_selected_id,
				'ID': int( row[ 0 ] ),
				'Category': str( row[ 3 ] or '' ),
				'Caption': str( row[ 1 ] or '' ),
				'Name': str( row[ 2 ] or '' ),
				'Text': str( row[ 4 ] or '' )
			} )
		df_prompt_table = pd.DataFrame(
			table_rows, columns=[ 'Selected', 'ID', 'Category', 'Caption', 'Name', 'Text' ] )
		df_edited_prompts = st.data_editor( df_prompt_table, hide_index=True,
			use_container_width=True, key='prompt_table' )
		
		# ------------------------------------------------------------------
		# Selection processing
		# ------------------------------------------------------------------
		selected_rows = df_edited_prompts.loc[ df_edited_prompts[ 'Selected' ] == True ].to_dict(
			orient='records' ) if not df_edited_prompts.empty else [ ]
		if len( selected_rows ) == 1:
			pid = int( selected_rows[ 0 ][ 'ID' ] )
			if pid != st.session_state.pe_selected_id:
				load_prompt_for_edit( pid )
				if bool( st.session_state.get( 'pe_cascade_enabled', False ) ):
					apply_prompt_to_text_generation( st.session_state.pe_text )
					apply_prompt_metadata_to_shared_state(
						category=st.session_state.pe_category,
						task_type=st.session_state.get( 'prompt_task', 'Chat' ),
						response_format=st.session_state.get( 'prompt_response_format', 'Markdown' ),
						language=st.session_state.get( 'pe_language', 'English' ) )
		elif len( selected_rows ) > 1:
			st.warning( 'Select exactly one prompt row.' )
		
		# ------------------------------------------------------------------
		# Paging
		# ------------------------------------------------------------------
		p1, p2, p3 = st.columns( [ 0.25, 3.5, 0.25 ] )
		with p1:
			if st.button( '◀ Prev' ) and st.session_state.pe_page > 1:
				st.session_state.pe_page -= 1
				st.rerun( )
		with p2:
			st.markdown( f'Page **{st.session_state.pe_page}** of **{total_pages}**' )
		with p3:
			if st.button( 'Next ▶' ) and st.session_state.pe_page < total_pages:
				st.session_state.pe_page += 1
				st.rerun( )
		
		st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
		
		# ------------------------------------------------------------------
		# Prompt actions
		# ------------------------------------------------------------------
		with st.expander( '⚙️ Prompt Actions', expanded=False ):
			app_c1, app_c2, app_c3 = st.columns( [ 0.34, 0.33, 0.33 ], border=True, gap='medium' )
			with app_c1:
				st.selectbox( 'Task Type', get_prompt_task_types( ), key='pe_task_type_edit' )
			with app_c2:
				st.selectbox( 'Response Format', get_response_formats( ), key='prompt_response_format' )
			with app_c3:
				st.selectbox( 'Response Language', get_spoken_languages( False ), key='pe_language_edit' )
			act_c1, act_c2, act_c3, act_c4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ] )
			with act_c1:
				if st.button( 'Apply to Text Generation', width='stretch' ):
					selected_category = str(
						st.session_state.get( 'pe_category', '' ) or '' )
					if not is_prompt_category_allowed_for_workflow( selected_category, 'Text Generation' ):
						st.error(
							f'The "{selected_category}" prompt category is not supported by the active '
							'Text Generation workflow and cannot be applied to Text Generation.' )
					else:
						apply_prompt_to_text_generation( st.session_state.get( 'pe_text', '' ) )
						apply_prompt_metadata_to_shared_state(
							category=selected_category,
							task_type=st.session_state.get( 'pe_task_type_edit', 'Chat' ),
							response_format=st.session_state.get( 'prompt_response_format', 'Markdown' ),
							language=st.session_state.get( 'pe_language_edit', 'English' ) )
						st.success( 'Applied to shared Text Generation settings.' )
			with act_c2:
				if st.button( 'Apply to Document Q&A', width='stretch' ):
					selected_category = str(
						st.session_state.get( 'pe_category', '' ) or '' )
					if not is_prompt_category_allowed_for_workflow( selected_category, 'Document Q&A' ):
						st.error(
							f'The "{selected_category}" prompt category is not supported by the active '
							'Document Q&A workflow and cannot be applied to Document Q&A.' )
					else:
						apply_prompt_to_document_qna( st.session_state.get( 'pe_text', '' ) )
						apply_prompt_metadata_to_shared_state(
							category=selected_category,
							task_type=st.session_state.get( 'pe_task_type_edit', 'Chat' ),
							response_format=st.session_state.get( 'prompt_response_format', 'Markdown' ),
							language=st.session_state.get( 'pe_language_edit', 'English' ) )
						st.success( 'Applied to shared Document Q&A settings.' )
			with act_c3:
				if st.button( 'Clone as New Template', width='stretch' ):
					source_prompt = {
						'ID': st.session_state.get( 'pe_selected_id' ),
						'Caption': st.session_state.get( 'pe_caption', '' ),
						'Name': st.session_state.get( 'pe_name', '' ),
						'Category': st.session_state.get( 'pe_category', '' ),
						'Text': st.session_state.get( 'pe_text', '' )
					}
					clone_prompt_record( source_prompt )
					st.success( 'Prompt cloned into a new editable draft.' )
			with act_c4:
				if st.button( 'Generate Starter Prompt', width='stretch' ):
					st.session_state.pe_text = build_starter_prompt_template(
						category=st.session_state.get( 'pe_category', '' ),
						task_type=st.session_state.get( 'pe_task_type_edit', 'Chat' ),
						response_format=st.session_state.get( 'prompt_response_format', 'Markdown' ),
						language=st.session_state.get( 'pe_language_edit', 'English' ) )
					st.success( 'Starter prompt generated into the edit surface.' )
		
			if st.button( label='Reset', key='prompt_actions_reset', width='stretch', icon='🔄' ):
				for key in [ 'pe_task_type_edit', 'prompt_response_format', 'pe_language_edit' ]:
					if key in st.session_state:
						del st.session_state[ key ]
				st.rerun( )

		# ------------------------------------------------------------------
		# Prompt generator
		# ------------------------------------------------------------------
		with st.expander( '🧪 Prompt Generator', expanded=False ):
			model_categories = fetch_prompt_categories( cfg.DB_PATH )
			if st.session_state.get( 'prompt_category_draft' ) not in model_categories:
				st.session_state[ 'prompt_category_draft' ] = (
					model_categories[ 0 ] if model_categories else '' )
			gen_c1, gen_c2, gen_c3, gen_c4, gen_c5 = st.columns(
				[ 0.2, 0.2, 0.2, 0.2, 0.2 ], border=True )
			with gen_c1:
				st.selectbox( 'Category', model_categories, key='prompt_category_draft' )
			with gen_c2:
				st.selectbox( 'Task Type', get_prompt_task_types( ), key='prompt_task_generator' )
			with gen_c3:
				st.selectbox( 'Response Format', get_response_formats( ), key='prompt_format' )
			with gen_c4:
				st.selectbox( 'Language', get_spoken_languages( include_auto_detect=False ),
					key='pe_language' )
			with gen_c5:
				st.selectbox( 'Generator Style', [ 'Practical', 'Formal', 'Analytical', 'Concise' ],
					key='pe_generator_style' )
			st.text_input( 'Goal', key='pe_generator_goal' )
			st.text_area( 'Constraints', height=120, key='pe_generator_constraints' )
			if st.button( 'Generate Template Draft', width='stretch' ):
				draft = generate_prompt_template_draft(
					goal=st.session_state.get( 'pe_generator_goal', '' ),
					constraints=st.session_state.get( 'pe_generator_constraints', '' ),
					style=st.session_state.get( 'pe_generator_style', 'Practical' ),
					category=st.session_state.get( 'prompt_category_draft', '' ),
					task_type=st.session_state.get( 'prompt_task_generator', 'Chat' ),
					response_format=st.session_state.get( 'prompt_format', 'Markdown' ),
					language=st.session_state.get( 'pe_language', 'English' ) )
				st.session_state[ 'pe_generated_template' ] = draft
				st.session_state.pe_text = draft
			if st.session_state.get( 'pe_generated_template', '' ):
				st.text_area( 'Generated Draft', value=st.session_state.get( 'pe_generated_template', '' ),
					height=180, disabled=True )
			if st.button( label='Reset', key='prompt_generator_reset', width='stretch', icon='🔄' ):
				for key in [ 'prompt_category_draft', 'prompt_task_generator', 'prompt_format', 'pe_language',
						'pe_generator_style', 'pe_generator_goal', 'pe_generator_constraints', 'pe_generated_template' ]:
					if key in st.session_state:
						del st.session_state[ key ]
				st.rerun( )
		
		# ------------------------------------------------------------------
		# Edit Prompt
		# ------------------------------------------------------------------
		with st.expander( '🖊️ Edit Prompt', expanded=False ):
			meta_c1, meta_c2, meta_c3, meta_c4 = st.columns(
				[ 0.25, 0.25, 0.25, 0.25 ], border=True, gap='medium' )
			with meta_c1:
				st.text_input( 'ID', value=st.session_state.pe_selected_id or '', disabled=True )
			with meta_c2:
				st.selectbox( 'Category', available_categories, key='pe_category' )
			with meta_c3:
				st.text_input( 'Caption', key='pe_caption' )
			with meta_c4:
				st.text_input( 'Name', key='pe_name' )
			st.text_area( 'Text', key='pe_text', height=260 )
			
			c1, c2, c3 = st.columns( 3 )
			with c1:
				save_label = '💾 Save Changes' if st.session_state.pe_selected_id else '➕ Create Prompt'
				if st.button( save_label ):
					prompt_data = {
						'Caption': st.session_state.pe_caption,
						'Name': st.session_state.pe_name,
						'Category': st.session_state.pe_category,
						'Text': st.session_state.pe_text
					}
					if st.session_state.pe_selected_id:
						update_prompt( int( st.session_state.pe_selected_id ), prompt_data )
					else:
						st.session_state.pe_selected_id = insert_prompt( prompt_data )
					st.success( 'Saved.' )
			with c2:
				if st.session_state.pe_selected_id and st.button( 'Delete' ):
					delete_prompt( int( st.session_state.pe_selected_id ) )
					reset_prompt_selection( )
					st.success( 'Deleted.' )
			with c3:
				st.button( '🧹 Clear Selection', on_click=reset_prompt_selection )


# ==============================================================================
# DATA MANAGEMENT MODE
# ==============================================================================
elif mode == 'Data Management':
	left, center, right = st.columns( [ 0.025, 0.95, 0.025 ] )
	with center:
		st.subheader( '🏛️ Data Management', help=cfg.DATA_MANAGEMENT )
		tabs = st.tabs( [ '📥 Import', '🗂 Browse', '💉 CRUD', '📊 Explore', '🔎 Filter',
		                  '🧮 Aggregate', '📈 Visualize', '⚙ Admin', '🧠 SQL' ] )
		
		tables = list_tables( )
		if not tables:
			st.info( 'No tables available.' )
		else:
			table = st.selectbox( 'Table', tables )
			df_full = read_table( table )
		
		# ----------------------------------------------------------------------
		# IMPORT TAB
		# ----------------------------------------------------------------------
		with tabs[ 0 ]:
			st.markdown( '#### Structured Data Import' )
			uploaded_file = st.file_uploader( 'Upload Excel File', type=[ 'xlsx' ] )
			overwrite = st.checkbox( 'Overwrite existing tables', value=True )
			
			if uploaded_file:
				try:
					sheets = pd.read_excel( uploaded_file, sheet_name=None )
					with create_connection( ) as conn:
						conn.execute( 'BEGIN' )
						for sheet_name, df in sheets.items( ):
							table_name = create_identifier( sheet_name )
							if overwrite:
								conn.execute( f'DROP TABLE IF EXISTS "{table_name}"' )
							
							columns = [ ]
							df.columns = [ create_identifier( c ) for c in df.columns ]
							for col in df.columns:
								sql_type = get_sqlite_type( df[ col ].dtype )
								columns.append( f'"{col}" {sql_type}' )
							
							create_stmt = (f'CREATE TABLE "{table_name}" '
							               f'({", ".join( columns )});')
							
							conn.execute( create_stmt )
							
							placeholders = ", ".join( [ "?" ] * len( df.columns ) )
							insert_stmt = (f'INSERT INTO "{table_name}" '
							               f'VALUES ({placeholders});')
							
							conn.executemany( insert_stmt,
								df.where( pd.notnull( df ), None ).values.tolist( ) )
						
						conn.commit( )
					
					st.success( 'Import completed successfully (transaction committed).' )
					st.rerun( )
				
				except Exception as e:
					exception = Error( e )
					exception.module = 'app'
					exception.cause = 'Streamlit UI'
					exception.method = 'streamlit_ui_block() -> None'
					Logger( ).write( exception )
					try:
						conn.rollback( )
					except Exception as e:
						exception = Error( e )
						exception.module = 'app'
						exception.cause = 'Streamlit UI'
						exception.method = 'streamlit_ui_block() -> None'
						Logger( ).write( exception )
						pass
					st.error( f'Import failed — transaction rolled back.\n\n{e}' )
			
			st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
			st.markdown( '#### AI Asset Registration' )
			asset_c1, asset_c2 = st.columns( [ 0.5, 0.5 ], border=True )
			with asset_c1:
				if st.button( 'Register Active Documents', width='stretch' ):
					doc_result = register_session_documents( )
					chunk_result = register_session_chunks( )
					embed_result = register_session_embeddings( )
					
					st.session_state[ 'dm_asset_sync_status' ] = (
							f'Documents inserted: {doc_result[ "inserted" ]}, '
							f'updated: {doc_result[ "updated" ]}, '
							f'chunks inserted: {chunk_result[ "inserted" ]}, '
							f'embeddings inserted: {embed_result[ "inserted" ]}')
					st.success( st.session_state[ 'dm_asset_sync_status' ] )
			
			with asset_c2:
				image_uploads = st.file_uploader( 'Upload images for metadata registration',
					type=[ 'png', 'jpg', 'jpeg', 'webp' ], accept_multiple_files=True,
					key='dm_image_uploads' )
				
				if st.button( 'Register Uploaded Images', width='stretch' ):
					if image_uploads:
						image_result = register_upload_images( image_uploads )
						st.session_state[ 'dm_asset_sync_status' ] = (
								f'Images inserted: {image_result[ "inserted" ]}, '
								f'updated: {image_result[ "updated" ]}')
						st.success( st.session_state[ 'dm_asset_sync_status' ] )
					else:
						st.info( 'Upload one or more images first.' )
			
			if st.session_state.get( 'dm_asset_sync_status', '' ):
				st.caption( st.session_state.get( 'dm_asset_sync_status', '' ) )
		
		# ----------------------------------------------------------------------
		# BROWSE TAB
		# ----------------------------------------------------------------------
		with tabs[ 1 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='table_name' )
				df = read_table( table )
				st.dataframe( df, use_container_width=True )
			else:
				st.info( 'No tables available.' )
		
		# ----------------------------------------------------------------------
		# CRUD TAB
		# ----------------------------------------------------------------------
		with tabs[ 2 ]:
			tables = list_tables( )
			if not tables:
				st.info( 'No tables available.' )
			else:
				table = st.selectbox( 'Select Table', tables, key='crud_table' )
				df = read_table( table )
				schema = create_schema( table )
				
				type_map = { col[ 1 ]: col[ 2 ].upper( ) for col in schema if col[ 1 ] != 'rowid' }
				
				st.subheader( 'Insert Row' )
				insert_data = { }
				for column, col_type in type_map.items( ):
					if 'INT' in col_type:
						insert_data[ column ] = st.number_input( column, step=1,
							key=f'ins_{column}' )
					elif 'REAL' in col_type:
						insert_data[ column ] = st.number_input( column, format='%.6f',
							key=f'ins_{column}' )
					elif 'BOOL' in col_type:
						insert_data[ column ] = 1 if st.checkbox( column,
							key=f'ins_{column}' ) else 0
					else:
						insert_data[ column ] = st.text_input( column, key=f'ins_{column}' )
				
				if st.button( 'Insert Row' ):
					cols = list( insert_data.keys( ) )
					placeholders = ', '.join( [ '?' ] * len( cols ) )
					stmt = f'INSERT INTO "{table}" ({", ".join( cols )}) VALUES ({placeholders});'
					
					with create_connection( ) as conn:
						conn.execute( stmt, list( insert_data.values( ) ) )
						conn.commit( )
					
					st.success( 'Row inserted.' )
					st.rerun( )
				
				st.subheader( 'Update Row' )
				rowid = st.number_input( 'Row ID', min_value=1, step=1 )
				update_data = { }
				for column, col_type in type_map.items( ):
					if 'INT' in col_type:
						val = st.number_input( column, step=1, key=f'upd_{column}' )
						update_data[ column ] = val
					elif 'REAL' in col_type:
						val = st.number_input( column, format='%.6f', key=f'upd_{column}' )
						update_data[ column ] = val
					elif 'BOOL' in col_type:
						val = 1 if st.checkbox( column, key=f'upd_{column}' ) else 0
						update_data[ column ] = val
					else:
						val = st.text_input( column, key=f'upd_{column}' )
						update_data[ column ] = val
				
				if st.button( 'Update Row' ):
					set_clause = ', '.join( [ f'{c}=?' for c in update_data ] )
					stmt = f'UPDATE {table} SET {set_clause} WHERE rowid=?;'
					
					with create_connection( ) as conn:
						conn.execute( stmt, list( update_data.values( ) ) + [ rowid ] )
						conn.commit( )
					
					st.success( 'Row updated.' )
					st.rerun( )
				
				st.subheader( 'Delete Row' )
				delete_id = st.number_input( 'Row ID to Delete', min_value=1, step=1 )
				if st.button( 'Delete Row' ):
					with create_connection( ) as conn:
						conn.execute( f'DELETE FROM {table} WHERE rowid=?;', (delete_id,) )
						conn.commit( )
					
					st.success( 'Row deleted.' )
					st.rerun( )
		
		# ----------------------------------------------------------------------
		# EXPLORE TAB
		# ----------------------------------------------------------------------
		with tabs[ 3 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='explore_table' )
				page_size = st.slider( 'Rows per page', 10, 500, 50 )
				page = st.number_input( 'Page', min_value=1, step=1 )
				offset = (page - 1) * page_size
				df_page = read_table( table, page_size, offset )
				st.dataframe( df_page, use_container_width=True )
		
		# ----------------------------------------------------------------------
		# FILTER TAB
		# ----------------------------------------------------------------------
		with tabs[ 4 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='filter_table' )
				df = read_table( table )
				column = st.selectbox( 'Column', df.columns )
				value = st.text_input( 'Contains' )
				if value:
					df = df[ df[ column ].astype( str ).str.contains( value ) ]
				st.dataframe( df, use_container_width=True )
		
		# ----------------------------------------------------------------------
		# AGGREGATE TAB
		# ----------------------------------------------------------------------
		with tabs[ 5 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='agg_table' )
				df = read_table( table )
				numeric_cols = df.select_dtypes( include=[ 'number' ] ).columns.tolist( )
				if numeric_cols:
					col = st.selectbox( 'Column', numeric_cols )
					agg = st.selectbox( 'Function', [ 'SUM', 'AVG', 'COUNT' ] )
					if agg == 'SUM':
						st.metric( 'Result', df[ col ].sum( ) )
					elif agg == 'AVG':
						st.metric( 'Result', df[ col ].mean( ) )
					elif agg == 'COUNT':
						st.metric( 'Result', df[ col ].count( ) )
		
		# ----------------------------------------------------------------------
		# VISUALIZE TAB
		# ----------------------------------------------------------------------
		with tabs[ 6 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='viz_table' )
				df = read_table( table )
				numeric_cols = df.select_dtypes( include=[ 'number' ] ).columns.tolist( )
				if numeric_cols:
					col = st.selectbox( 'Column', numeric_cols, key='viz_column' )
					fig = px.histogram( df, x=col )
					st.plotly_chart( fig, use_container_width=True )
		
		# ----------------------------------------------------------------------
		# ADMIN TAB
		# ----------------------------------------------------------------------
		with tabs[ 7 ]:
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Table', tables, key='admin_table' )
			
			st.divider( )
			
			st.subheader( 'AI Asset Governance' )
			
			if st.button( 'Refresh AI Asset Counts', width='stretch' ):
				st.session_state[ 'dm_asset_counts' ] = get_ai_asset_counts( )
			
			asset_counts = st.session_state.get( 'dm_asset_counts', { } )
			if asset_counts:
				ac1, ac2, ac3, ac4 = st.columns( [ 0.25, 0.25, 0.25, 0.25 ] )
				with ac1:
					st.metric( 'Documents', int( asset_counts.get( 'documents', 0 ) ) )
				with ac2:
					st.metric( 'Document Chunks',
						int( asset_counts.get( 'document_chunks', 0 ) ) )
				with ac3:
					st.metric( 'Document Embeddings',
						int( asset_counts.get( 'document_embeddings', 0 ) ) )
				with ac4:
					st.metric( 'Images', int( asset_counts.get( 'images', 0 ) ) )
			
			asset_admin_c1, asset_admin_c2 = st.columns( [ 0.5, 0.5 ], border=True )
			
			with asset_admin_c1:
				if st.button( 'Rebuild Active Document Asset Rows', width='stretch' ):
					doc_result = register_session_documents( )
					chunk_result = register_session_chunks( )
					embed_result = register_session_embeddings( )
					
					st.success(
						f'Documents inserted: {doc_result[ "inserted" ]}, '
						f'updated: {doc_result[ "updated" ]}, '
						f'chunks inserted: {chunk_result[ "inserted" ]}, '
						f'embeddings inserted: {embed_result[ "inserted" ]}' )
			
			with asset_admin_c2:
				if st.button( 'Purge Orphaned AI Assets', width='stretch' ):
					purge_result = purge_orphaned_ai_assets( )
					st.success( f'Deleted chunks: {purge_result[ "deleted_chunks" ]}, '
					            f'deleted embeddings: {purge_result[ "deleted_embeddings" ]}' )
			
			st.markdown( cfg.BLUE_DIVIDER, unsafe_allow_html=True )
			
			st.subheader( 'Data Profiling' )
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Select Table', tables, key='profile_table' )
				if st.button( 'Generate Profile' ):
					profile_df = create_profile_table( table )
					st.dataframe( profile_df, use_container_width=True )
			
			st.subheader( 'Drop Table' )
			
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Select Table to Drop', tables, key='admin_drop_table' )
				
				if 'dm_confirm_drop' not in st.session_state:
					st.session_state.dm_confirm_drop = False
				
				if st.button( 'Drop Table', key='admin_drop_button' ):
					st.session_state.dm_confirm_drop = True
				
				if st.session_state.dm_confirm_drop:
					st.warning( f'You are about to permanently delete table {table}. '
					            'This action cannot be undone.' )
					
					col1, col2 = st.columns( 2 )
					
					if col1.button( 'Confirm Drop', key='admin_confirm_drop' ):
						try:
							drop_table( table )
							st.success( f'Table {table} dropped successfully.' )
						except Exception as e:
							exception = Error( e )
							exception.module = 'app'
							exception.cause = 'Streamlit UI'
							exception.method = 'streamlit_ui_block() -> None'
							Logger( ).write( exception )
							st.error( f'Drop failed: {e}' )
						
						st.session_state.dm_confirm_drop = False
						st.rerun( )
					
					if col2.button( 'Cancel', key='admin_cancel_drop' ):
						st.session_state.dm_confirm_drop = False
						st.rerun( )
				
				df = read_table( table )
				col = st.selectbox( 'Create Index On', df.columns )
				
				if st.button( 'Create Index' ):
					create_index( table, col )
					st.success( 'Index created.' )
			
			st.divider( )
			
			st.subheader( 'Create Custom Table' )
			new_table_name = st.text_input( 'Table Name' )
			column_count = st.number_input( 'Number of Columns', min_value=1, max_value=20,
				value=1 )
			columns = [ ]
			for i in range( column_count ):
				st.markdown( f'### Column {i + 1}' )
				col_name = st.text_input( 'Column Name', key=f'col_name_{i}' )
				col_type = st.selectbox( 'Column Type', [ 'INTEGER', 'REAL', 'TEXT' ],
					key=f'col_type_{i}' )
				
				not_null = st.checkbox( 'NOT NULL', key=f'not_null_{i}' )
				primary_key = st.checkbox( 'PRIMARY KEY', key=f'pk_{i}' )
				auto_inc = st.checkbox( 'AUTOINCREMENT (INTEGER only)', key=f'ai_{i}' )
				
				columns.append( {
						'name': col_name,
						'type': col_type,
						'not_null': not_null,
						'primary_key': primary_key,
						'auto_increment': auto_inc
				} )
			
			if st.button( 'Create Table' ):
				try:
					create_custom_table( new_table_name, columns )
					st.success( 'Table created successfully.' )
					st.rerun( )
				except Exception as e:
					exception = Error( e )
					exception.module = 'app'
					exception.cause = 'Streamlit UI'
					exception.method = 'streamlit_ui_block() -> None'
					Logger( ).write( exception )
					st.error( f'Error: {e}' )
			
			st.divider( )
			st.subheader( 'Schema Viewer' )
			
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Select Table', tables, key='schema_view_table' )
				schema = create_schema( table )
				schema_df = pd.DataFrame( schema,
					columns=[ 'cid', 'name', 'type', 'notnull', 'default', 'pk' ] )
				
				st.markdown( '### Columns' )
				st.dataframe( schema_df, use_container_width=True )
				with create_connection( ) as conn:
					count = conn.execute( f'SELECT COUNT(*) FROM "{table}"' ).fetchone( )[ 0 ]
				
				st.metric( 'Row Count', f'{count:,}' )
				indexes = get_indexes( table )
				if indexes:
					idx_df = pd.DataFrame( indexes,
						columns=[ 'seq', 'name', 'unique', 'origin', 'partial' ] )
					st.markdown( '### Indexes' )
					st.dataframe( idx_df, use_container_width=True )
				else:
					st.info( 'No indexes defined.' )
			
			st.divider( )
			st.subheader( 'ALTER TABLE Operations' )
			tables = list_tables( )
			if tables:
				table = st.selectbox( 'Select Table', tables, key='alter_table_select' )
				operation = st.selectbox( 'Operation',
					[ 'Add Column', 'Rename Column', 'Rename Table', 'Drop Column' ] )
				
				if operation == 'Add Column':
					new_col = st.text_input( 'Column Name' )
					col_type = st.selectbox( 'Column Type', [ 'INTEGER', 'REAL', 'TEXT' ] )
					
					if st.button( 'Add Column' ):
						add_column( table, new_col, col_type )
						st.success( 'Column added.' )
						st.rerun( )
				
				elif operation == 'Rename Column':
					schema = create_schema( table )
					col_names = [ col[ 1 ] for col in schema ]
					old_col = st.selectbox( 'Column to Rename', col_names )
					new_col = st.text_input( 'New Column Name' )
					
					if st.button( 'Rename Column' ):
						rename_column( table, old_col, new_col )
						st.success( 'Column renamed.' )
						st.rerun( )
				
				elif operation == 'Rename Table':
					new_name = st.text_input( 'New Table Name' )
					
					if st.button( 'Rename Table' ):
						rename_table( table, new_name )
						st.success( 'Table renamed.' )
						st.rerun( )
				
				elif operation == 'Drop Column':
					schema = create_schema( table )
					col_names = [ col[ 1 ] for col in schema ]
					drop_col = st.selectbox( 'Column to Drop', col_names )
					
					if st.button( 'Drop Column' ):
						drop_column( table, drop_col )
						st.success( 'Column dropped.' )
						st.rerun( )
		
		# ----------------------------------------------------------------------
		# SQL TAB
		# ----------------------------------------------------------------------
		with tabs[ 8 ]:
			st.subheader( 'SQL Console' )
			query = st.text_area( 'Enter SQL Query' )
			if st.button( 'Run Query' ):
				if not is_safe_query( query ):
					st.error( 'Query blocked: Only read-only SELECT statements are allowed.' )
				else:
					try:
						start_time = time.perf_counter( )
						with create_connection( ) as conn:
							result = pd.read_sql_query( query, conn )
						
						end_time = time.perf_counter( )
						elapsed = end_time - start_time
						st.dataframe( result, use_container_width=True )
						row_count = len( result )
						col1, col2 = st.columns( 2 )
						col1.metric( 'Rows Returned', f'{row_count:,}' )
						col2.metric( 'Execution Time (seconds)', f'{elapsed:.6f}' )
						
						if elapsed > 2.0:
							st.warning( 'Slow query detected (> 2 seconds). Consider indexing.' )
						
						if not result.empty:
							csv = result.to_csv( index=False ).encode( 'utf-8' )
							st.download_button( 'Download CSV', csv, 'query_results.csv',
								'text/csv' )
					
					except Exception as e:
						exception = Error( e )
						exception.module = 'app'
						exception.cause = 'Streamlit UI'
						exception.method = 'streamlit_ui_block() -> None'
						Logger( ).write( exception )
						st.error( f'Execution failed: {e}' )

# ==============================================================================
# FOOTER — SECTION
# ==============================================================================
st.markdown(
	"""
	<style>
	.block-container {
		padding-bottom: 3rem;
	}
	</style>
	""",
	unsafe_allow_html=True )

# ---- Fixed Container
st.markdown( """
	<style>
	.boo-status-bar {
		position: fixed;
		bottom: 0;
		left: 0;
		width: 100%;
		background-color: rgba(27, 27, 27, 0.95);
		border-top: 1px solid #4d4d4d;
		padding: 10px 16px;
		font-size: 0.80rem;
		color: #4aa2f7;
		z-index: 1000;
	}
	.boo-status-inner {
		display: flex;
		justify-content: space-between;
		align-items: center;
		max-width: 100%;
	}
	</style>
	""", unsafe_allow_html=True, )

# ======================================================================================
# FOOTER RENDERING
# ======================================================================================

right_parts: List[ str ] = [ ]
model = 'Bro'

mode_value = mode if mode is not None else st.session_state.get( 'mode' )
if mode_value:
	right_parts.append( f'Mode: {mode_value}' )

temperature = st.session_state.get( 'temperature' )
top_p = st.session_state.get( 'top_percent' )
top_k = st.session_state.get( 'top_k' )
frequency = st.session_state.get( 'frequency_penalty' )
presense = st.session_state.get( 'presense_penalty' )
repeat_penalty = st.session_state.get( 'repeat_penalty' )
max_tokens = st.session_state.get( 'max_tokens' )
context_window = st.session_state.get( 'context_window' )
cpu_threads = st.session_state.get( 'cpu_threads' )
repeat_window = st.session_state.get( 'repeat_window' )
use_semantic = st.session_state.get( 'use_semantic' )
basic_docs = st.session_state.get( 'basic_docs' )

# ------------------------------------------------------------------
# Parameter summary (show 0 values; suppress only when None)
# ------------------------------------------------------------------
if temperature is not None:
	right_parts.append( f'Temp: {float( temperature ):0.2f}' )

if top_p is not None:
	right_parts.append( f'Top-P: {float( top_p ):0.2f}' )

if top_k is not None:
	right_parts.append( f'Top-K: {int( top_k )}' )

if frequency is not None:
	right_parts.append( f'Freq: {float( frequency ):0.2f}' )

if presense is not None:
	right_parts.append( f'Presence: {float( presense ):0.2f}' )

if repeat_penalty is not None:
	right_parts.append( f'Repeat: {float( repeat_penalty ):0.2f}' )

if repeat_window is not None:
	right_parts.append( f'Repeat Window: {int( repeat_window )}' )

if max_tokens is not None:
	right_parts.append( f'Max Tokens: {int( max_tokens )}' )

if context_window is not None:
	right_parts.append( f'Context: {int( context_window )}' )

if cpu_threads is not None:
	right_parts.append( f'Threads: {int( cpu_threads )}' )

# ------------------------------------------------------------------
# Context flags (optional but useful)
# ------------------------------------------------------------------
if use_semantic is not None:
	right_parts.append( f'Semantic: {"On" if use_semantic else "Off"}' )

if isinstance( basic_docs, list ):
	right_parts.append( f'Docs: {len( basic_docs )}' )

right_text = ' ◽ '.join( right_parts ) if right_parts else '—'

# ---- Rendering Method
st.markdown(
	f"""
    <div class="boo-status-bar">
        <div class="boo-status-inner">
            <span>{model}</span>
            <span>{right_text}</span>
        </div>
    </div>
    """,
	unsafe_allow_html=True, )